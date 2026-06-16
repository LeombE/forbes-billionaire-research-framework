"""Legacy 2025 baseline DOCX report generation.

Use ``src.enriched_reports --year <target_year>`` for year-aware enriched
reports. This module intentionally remains tied to the preserved 2025 flat
processed outputs for backward compatibility.
"""

from __future__ import annotations

import argparse
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .config import PEOPLE_REPORTS_DIR, PROCESSED_DIR, REPORT_SECTIONS, WEALTH_ENGINE_CATEGORIES, ensure_project_dirs


PLACEHOLDER_SOURCE_NEEDS = [
    "ownership stake / voting control filing",
    "annual report or 10-K/20-F financial statements",
    "segment revenue and profit source",
    "market capitalization or private valuation bridge",
    "risk-factor source from primary filings",
]


@dataclass(frozen=True)
class ReportContext:
    person: pd.Series
    history: pd.DataFrame
    metrics: pd.Series
    citations: pd.DataFrame
    person_quality: pd.Series | None
    peers: pd.DataFrame


def slugify(value: str) -> str:
    """Return a filesystem-safe ASCII slug."""
    normalized = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", normalized).strip("-").lower()
    return slug or "person"


def report_filename(person: pd.Series) -> str:
    """Return the required report filename for a person."""
    return f"{int(person['rank_2025']):03d}_{slugify(str(person['name']))}_business_analysis.docx"


def _load_processed(processed_dir: Path = PROCESSED_DIR) -> dict[str, pd.DataFrame]:
    return {
        "top100": pd.read_csv(processed_dir / "top100_2025.csv"),
        "history": pd.read_csv(processed_dir / "billionaire_wealth_history_long.csv"),
        "metrics": pd.read_csv(processed_dir / "billionaire_growth_metrics.csv"),
        "citations": pd.read_csv(processed_dir / "source_citations.csv"),
        "person_quality": pd.read_csv(processed_dir / "person_data_quality_scores.csv"),
    }


def _pick_person(top100: pd.DataFrame, *, rank: int | None, name: str | None, uri: str | None) -> pd.Series:
    if uri:
        matches = top100[top100["forbes_uri"].astype(str).str.casefold() == uri.casefold()]
    elif rank is not None:
        matches = top100[top100["rank_2025"].astype(int) == int(rank)]
        if len(matches) > 1:
            matches = matches.sort_values("position_2025").head(1)
    elif name:
        matches = top100[top100["name"].astype(str).str.casefold() == name.casefold()]
        if matches.empty:
            matches = top100[top100["name"].astype(str).str.contains(name, case=False, na=False, regex=False)]
    else:
        raise ValueError("Provide one selector: rank, name, or uri.")
    if matches.empty:
        raise ValueError("No matching person found in top100_2025.csv.")
    return matches.iloc[0]


def build_report_context(
    *,
    rank: int | None = None,
    name: str | None = None,
    uri: str | None = None,
    processed_dir: Path = PROCESSED_DIR,
) -> ReportContext:
    """Load structured data needed for one person report."""
    data = _load_processed(processed_dir)
    person = _pick_person(data["top100"], rank=rank, name=name, uri=uri)
    forbes_uri = str(person["forbes_uri"])
    history = data["history"][data["history"]["forbes_uri"].astype(str) == forbes_uri].sort_values("year")
    metric_matches = data["metrics"][data["metrics"]["forbes_uri"].astype(str) == forbes_uri]
    if metric_matches.empty:
        raise ValueError(f"No growth metrics row for {person['name']}.")
    metrics = metric_matches.iloc[0]
    citations = data["citations"][
        (data["citations"]["forbes_uri"].astype(str) == forbes_uri)
        | (data["citations"]["citation_scope"].astype(str) == "dataset_year")
    ]
    quality_matches = data["person_quality"][data["person_quality"]["forbes_uri"].astype(str) == forbes_uri]
    person_quality = quality_matches.iloc[0] if not quality_matches.empty else None
    peer_pool = data["metrics"].merge(
        data["top100"][["forbes_uri", "net_worth_2025_usd_b", "industry", "country_or_territory"]], on="forbes_uri"
    )
    peers = peer_pool[
        (peer_pool["wealth_engine_category"] == metrics["wealth_engine_category"])
        & (peer_pool["forbes_uri"].astype(str) != forbes_uri)
    ].sort_values("rank_2025").head(8)
    return ReportContext(person=person, history=history, metrics=metrics, citations=citations, person_quality=person_quality, peers=peers)


def _fmt_b(value: Any) -> str:
    if pd.isna(value):
        return "Not available"
    return f"${float(value):,.1f}B"


def _fmt_pct(value: Any) -> str:
    if pd.isna(value):
        return "Not available"
    return f"{float(value) * 100:,.1f}%"


def _fmt_num(value: Any, decimals: int = 2) -> str:
    if pd.isna(value):
        return "Not available"
    return f"{float(value):,.{decimals}f}"


def _citation_ids(citations: pd.DataFrame, table_name: str, field_names: list[str]) -> str:
    if citations.empty or "source_id" not in citations:
        return "[NO_SOURCE_ID]"
    matches = citations[
        citations["table_name"].astype(str).eq(table_name)
        & citations["field_name"].astype(str).isin(field_names)
    ]["source_id"].dropna().astype(str)
    ids = list(dict.fromkeys(matches.tolist()))[:4]
    return "[" + "; ".join(ids) + "]" if ids else "[SOURCE_NEEDED]"


def _add_key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True


def _section(document: Document, title: str) -> None:
    document.add_heading(title, level=1)


def _add_history_table(document: Document, history: pd.DataFrame) -> None:
    latest = history.tail(10)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Year", "Rank", "Net worth", "Source"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for _, row in latest.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(int(row["year"]))
        cells[1].text = "" if pd.isna(row.get("rank")) else str(int(row["rank"]))
        cells[2].text = _fmt_b(row["net_worth_usd_b"])
        cells[3].text = str(row.get("source_url", ""))


def _add_evidence_table(document: Document, citations: pd.DataFrame) -> None:
    cols = ["source_id", "field_name", "source_title", "publisher", "source_url", "evidence_note"]
    table = document.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for idx, col in enumerate(cols):
        table.rows[0].cells[idx].text = col
    for _, row in citations.drop_duplicates("source_id").head(25).iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(cols):
            cells[idx].text = str(row.get(col, ""))[:350]


def create_person_report(
    context: ReportContext,
    output_dir: Path = PEOPLE_REPORTS_DIR,
) -> Path:
    """Create one structured person DOCX report from processed data and citations."""
    ensure_project_dirs()
    output_dir.mkdir(parents=True, exist_ok=True)
    person = context.person
    metrics = context.metrics
    citations = context.citations
    citation_snapshot = _citation_ids(
        citations,
        "top100_2025",
        ["rank_2025", "net_worth_2025_usd_b", "source_of_wealth", "industry", "primary_company_or_asset"],
    )
    citation_history = _citation_ids(citations, "billionaire_wealth_history_long", ["net_worth_usd_b"])
    citation_metrics = _citation_ids(citations, "billionaire_growth_metrics", ["CAGR_nominal", "exponential_fit_r2"])

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{person['name']} Business Empire Analysis")
    title_run.bold = True
    title_run.font.size = Pt(18)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        f"Forbes 2025 rank {int(person['rank_2025'])} | Structured research draft | Source-backed where collected"
    )

    _section(document, REPORT_SECTIONS[0])
    document.add_paragraph(
        f"{person['name']} appears in the Forbes 2025 Top 100 with estimated net worth of "
        f"{_fmt_b(person['net_worth_2025_usd_b'])}. The current structured evidence ties the fortune to "
        f"{person['primary_company_or_asset']} / {person['source_of_wealth']} and classifies the dominant wealth engine as "
        f"{metrics['wealth_engine_category']}. {citation_snapshot}"
    )

    _section(document, REPORT_SECTIONS[1])
    _add_key_value_table(
        document,
        [
            ("Forbes 2025 net worth", _fmt_b(person["net_worth_2025_usd_b"])),
            ("Key asset/company", str(metrics.get("key_asset_or_company", person["primary_company_or_asset"]))),
            ("Wealth engine", str(metrics["wealth_engine_category"])),
            ("Classification confidence", str(metrics.get("classification_confidence", "Not available"))),
            ("Public equity dependency", str(metrics["public_equity_dependency_flag"])),
        ],
    )
    _add_note(document, f"Wealth equation source note: Forbes annual-list estimate and project classification. {citation_snapshot}")

    _section(document, REPORT_SECTIONS[2])
    document.add_paragraph(
        f"Current asset map: source of wealth is listed as {person['source_of_wealth']}; primary company or asset is "
        f"{person['primary_company_or_asset']}; self-made/inherited flag is {person['self_made_or_inherited_if_available']}. "
        f"Precise ownership percentages, voting control, family vehicles, debt, and discounts are not available from the collected Forbes-only sources. {citation_snapshot}"
    )

    _section(document, REPORT_SECTIONS[3])
    _add_key_value_table(
        document,
        [
            ("First year observed", str(metrics.get("first_year_observed", "Not available"))),
            ("First net worth", _fmt_b(metrics.get("first_net_worth_usd_b", np.nan))),
            ("Years observed", str(int(metrics.get("years_observed", 0)))),
            ("Wealth multiple", _fmt_num(metrics.get("wealth_multiple_first_to_2025", np.nan))),
            ("Nominal CAGR", _fmt_pct(metrics.get("CAGR_nominal", np.nan))),
            ("Peak net worth", _fmt_b(metrics.get("peak_net_worth_usd_b", np.nan))),
        ],
    )
    _add_history_table(document, context.history)
    _add_note(document, f"Wealth history source note: annual Forbes snapshots. {citation_history}")

    _section(document, REPORT_SECTIONS[4])
    document.add_paragraph(
        f"Log-linear fit slope: {_fmt_num(metrics.get('log_linear_growth_slope', np.nan), 4)}; "
        f"R^2: {_fmt_num(metrics.get('exponential_fit_r2', np.nan), 3)}; "
        f"estimated doubling time: {_fmt_num(metrics.get('estimated_doubling_time_years', np.nan), 2)} years. "
        "These values describe the observed Forbes annual estimates and should not be described as true exponential growth unless the data coverage and fit support that interpretation. "
        f"{citation_metrics}"
    )

    _section(document, REPORT_SECTIONS[5])
    if len(context.history) >= 2:
        changes = context.history[["year", "net_worth_usd_b"]].copy()
        changes["change"] = changes["net_worth_usd_b"].diff()
        top_changes = changes.dropna().sort_values("change", key=lambda s: s.abs(), ascending=False).head(5)
        for _, row in top_changes.iterrows():
            document.add_paragraph(
                f"{int(row['year'])}: annual Forbes net worth changed by {_fmt_b(row['change'])}. {citation_history}",
                style=None,
            )
    else:
        document.add_paragraph("Not available from collected sources; at least two annual observations are needed.")

    _section(document, REPORT_SECTIONS[6])
    document.add_paragraph(
        f"Collected structured data identifies {person['primary_company_or_asset']} as the core company/asset. "
        "Detailed operating-company descriptions, subsidiaries, and asset-level economics require primary company sources. "
        f"Needed sources: {', '.join(PLACEHOLDER_SOURCE_NEEDS[:3])}."
    )

    _section(document, REPORT_SECTIONS[7])
    document.add_paragraph(
        f"The current industry label is {person['industry']}. Industry structure, market share, profit pools, supplier/customer power, and regulation are not fully sourced in the current Forbes-only dataset. {citation_snapshot}"
    )

    _section(document, REPORT_SECTIONS[8])
    document.add_paragraph(
        "Not available from collected sources. Add annual reports, audited statements, investor presentations, exchange filings, or prospectuses before making revenue, margin, ROIC, cash-flow, or valuation-bridge claims."
    )

    _section(document, REPORT_SECTIONS[9])
    document.add_paragraph(
        f"Working moat hypothesis from structured fields: {metrics['evidence_summary']} This is a heuristic classification, not a final primary-source-supported moat conclusion."
    )

    _section(document, REPORT_SECTIONS[10])
    document.add_paragraph(
        "Capital allocation pattern is not available from collected sources beyond observed wealth history. Add buyback, dividend, acquisition, reinvestment, fund performance, or asset-sale sources before final analysis."
    )

    _section(document, REPORT_SECTIONS[11])
    document.add_paragraph(
        f"Country/territory: {person['country_or_territory']}; citizenship: {person.get('country_of_citizenship', 'Not available')}. "
        "Macro and regulatory implications require country/industry primary or official sources before final claims. "
        f"{citation_snapshot}"
    )

    _section(document, REPORT_SECTIONS[12])
    document.add_paragraph(
        "Key strategic decisions are not fully source-backed in the current dataset. Add company histories, filing narratives, acquisition releases, IPO/prospectus documents, and founder/board-control sources."
    )

    _section(document, REPORT_SECTIONS[13])
    document.add_paragraph(
        f"Quantitative risk indicators from Forbes annual estimates: largest drawdown {_fmt_pct(metrics.get('largest_drawdown_pct', np.nan))}; "
        f"annual growth volatility {_fmt_pct(metrics.get('volatility_of_annual_growth', np.nan))}; public equity dependency {metrics['public_equity_dependency_flag']}. "
        "Business-specific risks need primary risk-factor sources."
    )

    _section(document, REPORT_SECTIONS[14])
    if context.peers.empty:
        document.add_paragraph("No same-category peer rows available from the current Top 100 metrics table.")
    else:
        peer_names = ", ".join(context.peers["name"].astype(str).head(8))
        document.add_paragraph(
            f"Comparable Top 100 people by wealth engine category ({metrics['wealth_engine_category']}): {peer_names}. "
            "Peer financial comparison requires company-level source collection."
        )

    _section(document, REPORT_SECTIONS[15])
    document.add_paragraph(
        "Transferable lesson placeholder: ownership concentration, asset quality, and timing matter more than salary income. Treat this as an inference from the structured Top 100 dataset until expanded with primary-source business evidence."
    )

    _section(document, REPORT_SECTIONS[16])
    _add_evidence_table(document, citations)
    quality_note = "Person-level quality scoring not available."
    if context.person_quality is not None:
        quality_note = (
            f"Data quality score: {context.person_quality.get('data_quality_score')}. "
            f"Report ready from Forbes structured data: {context.person_quality.get('docx_report_ready')}. "
            f"Limitations: {context.person_quality.get('missing_or_limited_evidence')}."
        )
    document.add_paragraph(quality_note)

    output_path = output_dir / report_filename(person)
    document.save(output_path)
    return output_path


def validate_docx_report(path: Path) -> dict[str, Any]:
    """Validate that a generated DOCX contains required sections, citations, and quality notes."""
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    missing_sections = [section for section in REPORT_SECTIONS if section not in text]
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    has_citations = "source_id" in table_text and "src-" in table_text
    has_quality_notes = "Data quality score" in text or "Person-level quality scoring" in text
    return {
        "path": str(path),
        "missing_sections": missing_sections,
        "has_citations": has_citations,
        "has_quality_notes": has_quality_notes,
        "passed": not missing_sections and has_citations and has_quality_notes,
    }


def generate_report(
    *,
    rank: int | None = None,
    name: str | None = None,
    uri: str | None = None,
    output_dir: Path = PEOPLE_REPORTS_DIR,
    processed_dir: Path = PROCESSED_DIR,
) -> Path:
    """Generate and validate one legacy 2025 baseline person report."""
    context = build_report_context(rank=rank, name=name, uri=uri, processed_dir=processed_dir)
    path = create_person_report(context, output_dir=output_dir)
    validation = validate_docx_report(path)
    if not validation["passed"]:
        raise RuntimeError(f"DOCX validation failed: {validation}")
    return path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate legacy 2025 structured Forbes billionaire DOCX reports.")
    selector = parser.add_mutually_exclusive_group(required=True)
    selector.add_argument("--rank", type=int, help="Forbes 2025 rank to generate.")
    selector.add_argument("--name", help="Exact or partial person name to generate.")
    selector.add_argument("--uri", help="Forbes profile URI to generate.")
    parser.add_argument("--output-dir", type=Path, default=PEOPLE_REPORTS_DIR)
    parser.add_argument("--processed-dir", type=Path, default=PROCESSED_DIR)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    path = generate_report(
        rank=args.rank,
        name=args.name,
        uri=args.uri,
        output_dir=args.output_dir,
        processed_dir=args.processed_dir,
    )
    print(f"Generated DOCX report: {path}")


if __name__ == "__main__":
    main()
