"""Reusable enriched DOCX report generator for Forbes Top 100 people.

The baseline generator in ``src.docx_reports`` creates a structured draft from
project data. This module adds the richer "business empire" report path:
short citation keys in the body, grouped evidence appendices, chart insertion,
and optional person-specific evidence packs.
"""

from __future__ import annotations

import argparse
import math
import re
import unicodedata
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .config import (
    CHARTS_DIR,
    DEFAULT_TARGET_YEAR,
    PEOPLE_REPORTS_DIR,
    PROCESSED_DIR,
    WEALTH_ENGINE_CATEGORIES,
    YearConfig,
    ensure_project_dirs,
    ensure_year_dirs,
    get_year_config,
)


EVIDENCE_PACK_COLUMNS = [
    "forbes_uri",
    "person_name",
    "rank_2025",
    "entity_or_asset",
    "wealth_engine_archetype",
    "evidence_category",
    "report_section",
    "citation_key",
    "claim_supported",
    "ownership_stake_pct",
    "voting_or_control_note",
    "valuation_or_market_cap",
    "valuation_date",
    "revenue",
    "profit_or_cash_flow",
    "capex_or_capital_intensity",
    "funding_round_or_transaction",
    "asset_sale_or_acquisition",
    "segment_financials",
    "regulatory_context",
    "risk_factors",
    "business_model_notes",
    "source_title",
    "source_url",
    "source_file",
    "publisher",
    "author",
    "publication_date",
    "accessed_at",
    "source_type",
    "reliability_tier",
    "confidence_level",
    "evidence_note",
    "limitations",
]


EVIDENCE_CATEGORIES = [
    "Forbes / project data",
    "Public company filings",
    "Investor relations / company sources",
    "Private company valuation sources",
    "Company announcements",
    "Reputable secondary sources",
    "Manual evidence pack",
]

REPORT_STATUS_VALUES = [
    "baseline",
    "enriched_draft",
    "enriched_v2",
    "enriched_v3",
    "research_grade_candidate",
    "final_review_ready",
]


ARCHETYPE_GUIDANCE = {
    "Founder/operator public equity": {
        "wealth_equation": "ownership stake x listed market value +/- options, debt, taxes, pledges, and liquidity constraints",
        "required_evidence": "proxy statement, annual report, beneficial ownership table, market-cap history, risk factors",
        "analysis_focus": "stock-price sensitivity, voting control, operating margins, reinvestment, multiple expansion",
    },
    "Founder/operator private company": {
        "wealth_equation": "founder stake x latest credible private valuation +/- debt, preferences, dilution, and illiquidity discount",
        "required_evidence": "funding rounds, tender offers, audited statements if available, customer/contracts evidence",
        "analysis_focus": "private valuation quality, unit economics, growth runway, governance, liquidity path",
    },
    "Early employee/executive equity": {
        "wealth_equation": "equity grants/options/retained shares x company value +/- sales, taxes, vesting, and dilution",
        "required_evidence": "IPO prospectus, proxy ownership tables, option disclosures, employment history",
        "analysis_focus": "timing, retained ownership, platform growth, executive leverage versus founder control",
    },
    "Investor/capital allocator": {
        "wealth_equation": "investment partnership/holding-company economics + carried interest + retained ownership",
        "required_evidence": "fund letters, public holdings, annual reports, regulatory filings, performance records",
        "analysis_focus": "capital allocation process, compounding record, fee structure, leverage, drawdown control",
    },
    "Inherited/family-controlled business": {
        "wealth_equation": "family stake x controlled company value +/- succession structures, trusts, dividends, and debt",
        "required_evidence": "family ownership filings, annual reports, trust disclosures where available, succession materials",
        "analysis_focus": "control, governance, dividend policy, professional management, succession risk",
    },
    "Luxury/retail brand ownership": {
        "wealth_equation": "brand group ownership x luxury/retail earnings power and valuation multiple",
        "required_evidence": "annual reports, brand portfolio data, segment margins, store/channel data",
        "analysis_focus": "brand scarcity, pricing power, distribution control, creative direction, resilience through cycles",
    },
    "Technology/platform monopoly/network effects": {
        "wealth_equation": "platform equity x user/network scale, data advantage, switching costs, and monetization",
        "required_evidence": "annual reports, user metrics, revenue segments, competition/risk disclosures",
        "analysis_focus": "network effects, data loops, ecosystem lock-in, regulatory risk, AI/platform optionality",
    },
    "Real estate/land/infrastructure": {
        "wealth_equation": "asset ownership x appraised/market value + cash yield - leverage and development risk",
        "required_evidence": "property records, REIT/holding-company reports, project approvals, debt disclosures",
        "analysis_focus": "location scarcity, leverage, cap rates, development pipeline, regulatory/permit exposure",
    },
    "Commodities/energy/resources": {
        "wealth_equation": "resource/control stake x commodity price/cost curve/reserves +/- political and cycle risk",
        "required_evidence": "reserve reports, production data, filings, permits, commodity exposure disclosures",
        "analysis_focus": "cost curve position, reserves, capital intensity, cycle timing, geopolitical/regulatory risk",
    },
    "Diversified holding company": {
        "wealth_equation": "look-through value of operating subsidiaries and securities portfolio +/- holding discount",
        "required_evidence": "holding-company reports, segment data, public holdings, acquisition/sale history",
        "analysis_focus": "capital allocation, decentralization, tax efficiency, succession, portfolio concentration",
    },
    "Other/unclear": {
        "wealth_equation": "known assets x sourced values; unknown assets remain unvalued until evidence is available",
        "required_evidence": "Forbes notes, company filings, credible profiles, manual evidence pack",
        "analysis_focus": "separate verified wealth drivers from unsourced biography or speculation",
    },
}


def infer_report_status(path: Path) -> str:
    """Infer the report status label from the filename convention."""
    name = path.name.casefold()
    if "final_review_ready" in name:
        return "final_review_ready"
    if "research_grade_candidate" in name or "research-grade-candidate" in name:
        return "research_grade_candidate"
    if "enriched_v3" in name:
        return "enriched_v3"
    if "enriched_v2" in name:
        return "enriched_v2"
    if "enriched_draft" in name:
        return "enriched_draft"
    if "business_analysis" in name:
        return "baseline"
    return "baseline"


@dataclass(frozen=True)
class EvidenceItem:
    key: str
    category: str
    title: str
    publisher: str
    date: str
    supports: str
    reliability: str
    locator: str
    limitation: str
    confidence: str = "Medium"
    entity: str = ""
    accessed_at: str = "2026-06-11"


@dataclass(frozen=True)
class EnrichedContext:
    person: pd.Series
    history: pd.DataFrame
    metrics: pd.Series
    citations: pd.DataFrame
    evidence_pack: pd.DataFrame


AMAZON_FACTS = {
    "net_sales_2024_b": 637.959,
    "net_product_sales_2024_b": 272.311,
    "net_service_sales_2024_b": 365.648,
    "operating_income_2024_b": 68.593,
    "net_income_2024_b": 59.248,
    "operating_cash_flow_2024_b": 115.877,
    "ppe_purchases_net_2024_b": 77.658,
    "ppe_purchases_gross_2024_b": 82.999,
    "free_cash_flow_2024_b": 38.219,
    "fulfillment_expense_2024_b": 98.505,
    "technology_infrastructure_expense_2024_b": 88.544,
    "north_america_revenue_2024_b": 387.497,
    "north_america_operating_income_2024_b": 24.967,
    "international_revenue_2024_b": 142.906,
    "international_operating_income_2024_b": 3.792,
    "aws_revenue_2024_b": 107.556,
    "aws_operating_income_2024_b": 39.834,
    "online_stores_revenue_2024_b": 247.029,
    "physical_stores_revenue_2024_b": 21.215,
    "third_party_seller_services_revenue_2024_b": 156.146,
    "advertising_services_revenue_2024_b": 56.214,
    "subscription_services_revenue_2024_b": 44.374,
    "other_revenue_2024_b": 5.425,
    "diluted_weighted_average_shares_2024_m": 10_721,
    "bezos_beneficial_shares_2025": 1_021_742_026,
    "bezos_beneficial_ownership_pct_2025": 9.6,
    "bezos_sole_voting_no_investment_shares_2025": 112_032_131,
    "directors_and_officers_shares_2025": 1_025_004_331,
    "directors_and_officers_pct_2025": 9.7,
}


META_FACTS = {
    "revenue_2024_b": 164.501,
    "costs_expenses_2024_b": 95.121,
    "operating_income_2024_b": 69.380,
    "net_income_2024_b": 62.360,
    "operating_cash_flow_2024_b": 91.328,
    "capex_2024_b": 37.256,
    "capex_including_finance_lease_principal_2024_b": 39.230,
    "free_cash_flow_2024_b": 54.072,
    "meta_defined_free_cash_flow_2024_b": 52.103,
    "share_repurchases_2024_b": 30.125,
    "repurchased_and_retired_2024_b": 29.750,
    "dividends_2024_b": 5.072,
    "assets_2024_b": 276.054,
    "cash_2024_b": 43.889,
    "family_revenue_2024_b": 162.355,
    "family_operating_income_2024_b": 87.109,
    "family_costs_expenses_2024_b": 75.246,
    "reality_labs_revenue_2024_b": 2.146,
    "reality_labs_operating_loss_2024_b": -17.729,
    "reality_labs_costs_expenses_2024_b": 19.875,
    "advertising_revenue_2024_b": 160.633,
    "other_family_revenue_2024_b": 1.722,
    "family_daily_active_people_dec_2024_b": 3.35,
    "ad_impressions_growth_2024": 0.11,
    "average_price_per_ad_growth_2024": 0.10,
    "class_a_outstanding_2025": 2_181_270_402,
    "class_b_outstanding_2025": 343_179_151,
    "zuckerberg_class_a_2025": 141_000,
    "zuckerberg_class_b_2025": 342_606_985,
    "zuckerberg_class_b_pct_2025": 99.8,
    "zuckerberg_voting_power_pct_2025": 61.0,
    "zuckerberg_pledged_class_b_2025": 12_000_000,
    "zuckerberg_pledged_pct_beneficial_2025": 3.5,
    "zuckerberg_pledged_pct_total_common_2025": 0.5,
    "zuckerberg_pledged_voting_power_pct_2025": 2.1,
}


LVMH_FACTS = {
    "revenue_2024_b": 84.683,
    "gross_margin_2024_b": 56.765,
    "profit_from_recurring_operations_2024_b": 19.571,
    "net_profit_2024_b": 12.958,
    "group_share_net_profit_2024_b": 12.550,
    "net_cash_from_operating_activities_2024_b": 18.924,
    "operating_investments_2024_b": 7.478,
    "operating_free_cash_flow_2024_b": 10.478,
    "net_financial_debt_2024_b": 9.228,
    "equity_2024_b": 69.287,
    "stores_2024": 6307,
    "family_group_share_capital_pct_2024": 49.00,
    "family_group_exercisable_voting_rights_pct_2024": 64.81,
    "family_group_shares_2024": 245_173_934,
    "christian_dior_share_capital_pct_2024": 41.87,
    "christian_dior_exercisable_voting_rights_pct_2024": 56.77,
    "arnault_other_share_capital_pct_2024": 7.13,
    "arnault_other_exercisable_voting_rights_pct_2024": 8.04,
    "fashion_retail_share_2024": 0.95,
    "fashion_stores_2024_min": 2300,
    "segment_rows_2024": [
        ("Wines and Spirits", 5.862, 1.356),
        ("Fashion and Leather Goods", 41.060, 15.230),
        ("Perfumes and Cosmetics", 8.418, 0.671),
        ("Watches and Jewelry", 10.577, 1.546),
        ("Selective Retailing", 18.262, 1.385),
        ("Other activities and eliminations", 0.504, -0.617),
    ],
    "revenue_region_pct_2024": {
        "Asia excl. Japan": 28,
        "United States": 25,
        "Europe excl. France": 17,
        "Other markets": 13,
        "Japan": 9,
        "France": 8,
    },
}


ORACLE_FACTS = {
    "total_revenue_2024_b": 52.961,
    "total_operating_expenses_2024_b": 37.608,
    "total_operating_margin_2024_b": 15.353,
    "total_operating_margin_pct_2024": 0.29,
    "net_income_2024_b": 10.467,
    "operating_cash_flow_2024_b": 18.673,
    "capital_expenditures_2024_b": 6.866,
    "free_cash_flow_2024_b": 11.807,
    "common_stock_repurchases_2024_b": 1.202,
    "tax_withholding_share_repurchases_2024_b": 2.040,
    "dividends_paid_2024_b": 4.391,
    "cash_equivalents_2024_b": 10.454,
    "cloud_and_license_revenue_2024_b": 44.464,
    "hardware_revenue_2024_b": 3.066,
    "services_revenue_2024_b": 5.431,
    "cloud_and_license_margin_2024_b": 28.514,
    "hardware_margin_2024_b": 1.915,
    "services_margin_2024_b": 0.916,
    "cloud_services_revenue_2024_b": 19.774,
    "license_support_revenue_2024_b": 19.609,
    "cloud_license_on_prem_revenue_2024_b": 5.081,
    "applications_cloud_support_revenue_2024_b": 18.172,
    "infrastructure_cloud_support_revenue_2024_b": 21.211,
    "cloud_services_license_support_total_2024_b": 39.383,
    "americas_revenue_2024_b": 33.122,
    "emea_revenue_2024_b": 13.030,
    "asia_pacific_revenue_2024_b": 6.809,
    "ellison_beneficial_shares_2024": 1_153_232_353,
    "ellison_beneficial_ownership_pct_2024": 41.6,
    "ellison_exercisable_options_included_2024": 7_500_000,
    "ellison_pledged_shares_2024": 277_000_000,
}


def slugify(value: str) -> str:
    """Return an ASCII slug safe for filenames."""
    normalized = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", normalized).strip("-").lower()
    return slug or "person"


def clean(value: Any) -> str:
    """Return a readable string for cells that may be blank/NaN."""
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except (TypeError, ValueError):
        pass
    return str(value).strip()


def money(value: Any, suffix: str = "B") -> str:
    """Format a USD billions value when available."""
    try:
        if pd.isna(value):
            return "not disclosed"
        number = float(value)
        if number < 0:
            return f"-${abs(number):,.1f}{suffix}"
        return f"${number:,.1f}{suffix}"
    except (TypeError, ValueError):
        text = clean(value)
        return text or "not disclosed"


def money_precise_b(value: Any) -> str:
    """Format USD billions to three decimals for extracted filing values."""
    try:
        if pd.isna(value):
            return "not disclosed"
        return f"${float(value):,.3f}B"
    except (TypeError, ValueError):
        text = clean(value)
        return text or "not disclosed"


def eur_b(value: Any) -> str:
    """Format EUR billions to three decimals for extracted LVMH filing values."""
    try:
        if pd.isna(value):
            return "not disclosed"
        return f"EUR {float(value):,.3f}B"
    except (TypeError, ValueError):
        text = clean(value)
        return text or "not disclosed"


def pct(value: Any, *, already_percent: bool = False) -> str:
    """Format percentage values."""
    try:
        if pd.isna(value):
            return "not disclosed"
        number = float(value)
        if not already_percent:
            number *= 100
        return f"{number:,.1f}%"
    except (TypeError, ValueError):
        text = clean(value)
        return text or "not disclosed"


def amazon_beneficial_ownership_pct() -> float:
    """Return Bezos's 2025 proxy beneficial ownership percentage as a decimal."""
    return AMAZON_FACTS["bezos_beneficial_ownership_pct_2025"] / 100


def amazon_aws_operating_margin() -> float:
    """Return AWS 2024 operating margin from the Amazon 2024 10-K segment table."""
    return AMAZON_FACTS["aws_operating_income_2024_b"] / AMAZON_FACTS["aws_revenue_2024_b"]


def lvmh_family_group_ownership_pct() -> float:
    """Return Arnault family group 2024 LVMH share-capital percentage as a decimal."""
    return LVMH_FACTS["family_group_share_capital_pct_2024"] / 100


def oracle_ellison_ownership_pct() -> float:
    """Return Ellison's 2024 proxy beneficial ownership percentage as a decimal."""
    return ORACLE_FACTS["ellison_beneficial_ownership_pct_2024"] / 100


def output_filename(person: pd.Series, variant: str, config: YearConfig | None = None) -> str:
    """Return an enriched report filename for a person and variant."""
    config = config or get_year_config(DEFAULT_TARGET_YEAR)
    rank = int(person[config.rank_col])
    slug = clean(person.get("name_slug")) or slugify(str(person["name"]))
    if config.legacy_layout:
        return f"{rank:03d}_{slug}_business_analysis_{variant}.docx"
    return f"{rank:03d}_{slug}_business_analysis_{config.year}_{variant}.docx"


def person_year(person: pd.Series) -> int:
    """Infer the annual-list year from a person row."""
    value = clean(person.get("year"))
    return int(float(value)) if value else 2025


def person_rank(person: pd.Series) -> int:
    """Return the target-year Forbes rank from a legacy or multi-year row."""
    field = "rank" if "rank" in person.index else "rank_2025"
    return int(float(person[field]))


def person_net_worth(person: pd.Series) -> Any:
    """Return the target-year net worth value from a legacy or multi-year row."""
    return person.get("net_worth_usd_b") if "net_worth_usd_b" in person.index else person.get("net_worth_2025_usd_b")


def metric_multiple(metric: pd.Series) -> Any:
    """Return the first-to-target-year multiple from a legacy or multi-year metric row."""
    return (
        metric.get("wealth_multiple_first_to_target_year")
        if "wealth_multiple_first_to_target_year" in metric.index
        else metric.get("wealth_multiple_first_to_2025")
    )


def load_evidence_pack(path: Path | None, forbes_uri: str | None = None) -> pd.DataFrame:
    """Load an optional person-specific evidence pack and normalize columns."""
    if path is None or not path.exists():
        return pd.DataFrame(columns=EVIDENCE_PACK_COLUMNS)
    pack = pd.read_csv(path)
    if "rank" in pack.columns and "rank_2025" not in pack.columns:
        pack["rank_2025"] = pack["rank"]
    if "person_name" not in pack.columns and "name" in pack.columns:
        pack["person_name"] = pack["name"]
    for column in EVIDENCE_PACK_COLUMNS:
        if column not in pack.columns:
            pack[column] = ""
    extra_columns = [
        column
        for column in ["report_year", "rank", "person_slug", "source_as_of_date", "claim_year"]
        if column in pack.columns and column not in EVIDENCE_PACK_COLUMNS
    ]
    pack = pack[EVIDENCE_PACK_COLUMNS + extra_columns]
    if forbes_uri:
        pack = pack[pack["forbes_uri"].astype(str).str.casefold().eq(str(forbes_uri).casefold())]
    return pack.fillna("")


def load_context(
    *,
    rank: int | None = None,
    name: str | None = None,
    uri: str | None = None,
    processed_dir: Path | None = None,
    evidence_pack_path: Path | None = None,
    year: int = DEFAULT_TARGET_YEAR,
) -> EnrichedContext:
    """Load top100, wealth history, growth metrics, citations, and optional evidence pack."""
    config = get_year_config(year)
    processed_dir = processed_dir or config.processed_dir
    top100 = pd.read_csv(processed_dir / config.top100_filename)
    history = pd.read_csv(processed_dir / config.history_filename)
    metrics = pd.read_csv(processed_dir / config.metrics_filename)
    citations = pd.read_csv(processed_dir / config.citations_filename)

    if rank is not None:
        matches = top100[top100[config.rank_col].astype(int).eq(int(rank))]
    elif uri:
        matches = top100[top100["forbes_uri"].astype(str).str.casefold().eq(uri.casefold())]
    elif name:
        matches = top100[top100["name"].astype(str).str.contains(name, case=False, regex=False, na=False)]
    else:
        raise ValueError("Provide --rank, --name, or --uri.")
    if matches.empty:
        raise ValueError(f"No matching person found in {config.top100_filename}.")

    person = matches.sort_values(config.rank_col).iloc[0]
    forbes_uri = clean(person.get("forbes_uri"))
    person_history = history[history["forbes_uri"].astype(str).eq(forbes_uri)].sort_values("year")
    metric_matches = metrics[metrics["forbes_uri"].astype(str).eq(forbes_uri)]
    if metric_matches.empty:
        raise ValueError(f"No growth metrics row found for {person['name']}.")
    metric = metric_matches.iloc[0]
    person_citations = citations[
        citations.get("forbes_uri", pd.Series(dtype=str)).astype(str).eq(forbes_uri)
        | citations.get("citation_scope", pd.Series(dtype=str)).astype(str).eq("dataset_year")
    ].copy()
    evidence_pack = load_evidence_pack(evidence_pack_path, forbes_uri=forbes_uri)
    return EnrichedContext(person, person_history, metric, person_citations, evidence_pack)


def project_evidence(context: EnrichedContext) -> list[EvidenceItem]:
    """Create short project-data evidence entries for any person."""
    person = context.person
    metric = context.metrics
    year = person_year(person)
    rank = person_rank(person)
    net_worth = person_net_worth(person)
    canonical_url = clean(person.get("canonical_source_url")) or clean(person.get("canonical_annual_list_source_url")) or f"data/raw/forbes/{year}/forbes_billionaires_{year}.json"
    profile_url = clean(person.get("forbes_profile_url"))
    return [
        EvidenceItem(
            key="F1",
            category="Forbes / project data",
            title=f"Forbes World's Billionaires {year} annual list record",
            publisher="Forbes / local project cache",
            date=f"{year} annual list",
            supports=(
                f"Rank {rank}, {year} net worth {money(net_worth)}, "
                f"country, source of wealth, industry, and profile URL for {person['name']}."
            ),
            reliability="Canonical for Forbes rank/net worth estimate",
            locator=canonical_url if len(canonical_url) < 120 else "Forbes annual API URL retained in source_citations.csv",
            limitation="Forbes net worth is an estimate, not an audited personal balance sheet.",
            confidence="High",
        ),
        EvidenceItem(
            key="F2",
            category="Forbes / project data",
            title="Forbes profile URL retained by project",
            publisher="Forbes",
            date="profile URL captured in processed data",
            supports=f"Profile locator for background audit: {profile_url or 'not available'}.",
            reliability="Background profile locator",
            locator=profile_url or "not available",
            limitation="Profile pages can change after the annual-list snapshot.",
            confidence="Medium",
        ),
        EvidenceItem(
            key="D1",
            category="Forbes / project data",
            title="Project processed Forbes annual wealth history and growth metrics",
            publisher="Local project data derived from Forbes annual files",
            date=f"2001-{year} snapshots where available",
            supports=(
                f"First observed year {clean(metric.get('first_year_observed'))}; "
                f"CAGR {pct(metric.get('CAGR_nominal'))}; R^2 {clean(round(float(metric.get('exponential_fit_r2')), 3)) if clean(metric.get('exponential_fit_r2')) else 'not available'}."
            ),
            reliability="Derived from canonical Forbes annual-list history",
            locator=f"data/processed/{year}/billionaire_wealth_history_long_{year}.csv; data/processed/{year}/billionaire_growth_metrics_{year}.csv" if year != 2025 else "data/processed/billionaire_wealth_history_long.csv; data/processed/billionaire_growth_metrics.csv",
            limitation="Annual snapshots miss intra-year volatility and private valuation timing.",
            confidence="High",
        ),
    ]


def bezos_gold_evidence() -> list[EvidenceItem]:
    """Curated official source set for the Phase 5B Jeff Bezos report."""
    sec_10k = "https://www.sec.gov/Archives/edgar/data/1018724/000101872425000004/amzn-20241231.htm"
    sec_proxy = "https://www.sec.gov/Archives/edgar/data/1018724/000110465925033442/tm252295-1_def14a.htm"
    return [
        EvidenceItem(
            "AMZ1",
            "Public company filings",
            "Amazon.com, Inc. 2024 Form 10-K consolidated financials and cash flow",
            "U.S. SEC / Amazon.com, Inc.",
            "2025-02-07",
            "Amazon 2024 consolidated net sales of $637.959B, operating income of $68.593B, net income of $59.248B, operating cash flow of $115.877B, net PPE purchases of $77.658B, free cash flow of $38.219B, fulfillment expense of $98.505B, and technology/infrastructure expense of $88.544B.",
            "Primary filing",
            sec_10k,
            "Company filing supports Amazon economics, not Bezos's complete personal balance sheet.",
            "High",
            entity="Amazon",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "AMZ2",
            "Public company filings",
            "Amazon.com, Inc. 2024 Form 10-K segment and revenue-mix tables",
            "U.S. SEC / Amazon.com, Inc.",
            "2025-02-07",
            "Amazon 2024 segment net sales and operating income for North America, International, and AWS; revenue lines for online stores, physical stores, third-party seller services, advertising services, subscription services, AWS, and other.",
            "Primary filing",
            sec_10k,
            "Segment data is company-reported and does not assign separate market values to each business line.",
            "High",
            entity="Amazon segments",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "AMZ3",
            "Public company filings",
            "Amazon.com, Inc. 2024 Form 10-K business description and risk factors",
            "U.S. SEC / Amazon.com, Inc.",
            "2025-02-07",
            "Business model evidence for retail, third-party sellers, AWS, machine learning/cloud services, advertising, subscriptions, fulfillment/logistics, competition, regulation, infrastructure, and operating risk.",
            "Primary filing",
            sec_10k,
            "Risk disclosures identify exposures but do not quantify every probability or Bezos-specific wealth effect.",
            "High",
            entity="Amazon risk factors",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "AMZ4",
            "Public company filings",
            "Amazon.com, Inc. 2025 definitive proxy statement",
            "U.S. SEC / Amazon.com, Inc.",
            "2025-04-10",
            "Bezos beneficial ownership of 1,021,742,026 shares, 9.6% of class, one-share/one-vote common stock governance context, founder/executive chair biography, Blue Origin founder reference, and Washington Post ownership reference.",
            "Primary filing",
            sec_proxy,
            "Beneficial ownership is not a full personal wealth bridge; proxy does not disclose Bezos's taxes, trusts, debt, liquidity discounts, share-sale plan, or full private holdings.",
            "High",
            entity="Jeff Bezos / Amazon ownership",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "AIR1",
            "Investor relations / company sources",
            "Amazon Investor Relations overview",
            "Amazon Investor Relations",
            "accessed 2026-06-16",
            "Official Amazon investor-relations hub for annual reports, proxies, SEC filings, quarterly results, governance, officers/directors, and investor materials.",
            "Company investor-relations source",
            "https://ir.aboutamazon.com/",
            "IR hub is a source locator; exact financial claims in this report are anchored to the SEC 10-K and proxy.",
            "High",
            entity="Amazon investor relations",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "BO1",
            "Company announcements",
            "About Blue Origin",
            "Blue Origin",
            "accessed 2026-06-16",
            "Blue Origin mission, reusable-rocket objective, New Shepard/New Glenn/lunar/engine/in-space mobility product map, and long-horizon space infrastructure framing.",
            "Company source",
            "https://www.blueorigin.com/about-blue",
            "Company page does not disclose Bezos ownership percentage, valuation, revenue, profit, capex, or liquidity path.",
            "Medium",
            entity="Blue Origin",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "BO2",
            "Company announcements",
            "New Glenn",
            "Blue Origin",
            "accessed 2026-06-16",
            "New Glenn reusable orbital launch vehicle description, first-stage reuse design, payload capability, engines, and customer/mission framing.",
            "Company source",
            "https://www.blueorigin.com/new-glenn",
            "Company page supports product/strategy context, not private valuation or Bezos wealth attribution.",
            "Medium",
            entity="Blue Origin New Glenn",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "BE1",
            "Company announcements",
            "Bezos Expeditions selected investments page",
            "Bezos Expeditions",
            "accessed 2026-06-16",
            "Official Bezos Expeditions page listing selected investments and links to Blue Origin, The Washington Post, Bezos Day One Fund, Bezos Family Foundation, and other holdings/investments.",
            "Personal/company source",
            "https://www.bezosexpeditions.com/",
            "Selected-investment page does not disclose ownership percentages, entry prices, current values, liquidity, or whether each holding is material to Forbes net worth.",
            "Medium",
            entity="Bezos Expeditions",
            accessed_at="2026-06-16",
        ),
    ]


def elon_gold_evidence() -> list[EvidenceItem]:
    """Curated source set for the Phase 3 Elon Musk gold-standard sample."""
    return [
        EvidenceItem("T1", "Public company filings", "Tesla Form 10-K for fiscal year 2024", "U.S. SEC / Tesla, Inc.", "2025-01-30", "Tesla revenue, gross profit, operating income, net income, operating cash flow, capex, cash, risk factors, and market sensitivity.", "Primary filing", "SEC accession 0001628280-25-003063", "Tesla filing reports company economics, not Musk's complete personal balance sheet.", "High", "Tesla"),
        EvidenceItem("T2", "Public company filings", "Tesla 2024 definitive proxy statement", "U.S. SEC / Tesla, Inc.", "2024-04-29", "Musk beneficial ownership of 715,022,706 Tesla shares, 20.5% beneficial ownership, pledged-share disclosure, and 2018 CEO award context.", "Primary filing", "SEC accession 0001104659-24-053333", "Ownership table is as of March 31, 2024, before the Forbes 2025 list date.", "High", "Tesla"),
        EvidenceItem("T3", "Public company filings", "Tesla 2025 definitive proxy statement", "U.S. SEC / Tesla, Inc.", "2025-09-17", "Later ownership context: 717,323,438 beneficial shares and 19.8% ownership as of September 15, 2025.", "Primary filing", "SEC accession 0001104659-25-090866", "Post-dates Forbes 2025 annual list and is used only as later context.", "Medium-high", "Tesla"),
        EvidenceItem("S1", "Private company valuation sources", "SpaceX valued at $350bn as company agrees to buy shares from employees", "The Guardian, citing Bloomberg/internal tender reporting", "2024-12-11", "SpaceX private valuation estimate around $350B, tender offer mechanics, and reported Musk stake context.", "Reputable secondary", "theguardian.com/science/2024/dec/11/spacex-valued-at-350bn-as-company-agrees-to-buy-shares-from-employees", "Private valuation and ownership figures are secondary estimates, not company-filed financials.", "Medium", "SpaceX"),
        EvidenceItem("S2", "Company announcements", "NASA Commercial Crew Program", "NASA", "accessed 2026-06-11", "Commercial crew public-private model and NASA reliance on commercial providers for ISS transportation.", "Primary government source", "nasa.gov/humans-in-space/commercial-space/commercial-crew-program/", "Program page describes missions, not SpaceX margins or profitability.", "High", "SpaceX"),
        EvidenceItem("S3", "Company announcements", "NASA picks SpaceX to land next Americans on Moon", "NASA", "2021-04-16", "SpaceX HLS Starship award, $2.89B contract, and reusable Starship architecture.", "Primary government source", "nasa.gov/news-release/as-artemis-moves-forward-nasa-picks-spacex-to-land-next-americans-on-moon/", "Contract value does not disclose SpaceX margin or internal development cost.", "High", "SpaceX"),
        EvidenceItem("X1", "Company announcements", "Elon Musk to Acquire Twitter", "Twitter, Inc. press release via PR Newswire", "2022-04-25", "$44B Twitter acquisition price, $54.20 per share, debt/margin loan financing, and equity commitment.", "Company announcement", "prnewswire.com/news-releases/elon-musk-to-acquire-twitter-301532245.html", "Announcement is pre-close and does not report later X operating performance.", "High", "X / Twitter"),
        EvidenceItem("X2", "Private company valuation sources", "Elon Musk says xAI bought X in all-stock deal", "Business Insider", "2025-03-28", "Reported xAI-X all-stock values: xAI at $80B and X at $33B equity value.", "Reputable secondary", "businessinsider.com/elon-musk-says-xai-acquired-x-in-all-stock-deal-2025-3", "Exact ownership allocation, debt, and minority investor economics remain undisclosed.", "Medium-low", "X / xAI"),
        EvidenceItem("A1", "Company announcements", "Series B funding round", "xAI", "2024-05-26", "xAI $6B Series B, investors, product infrastructure, and mission.", "Company announcement", "x.ai/news/series-b", "No valuation table, revenue, profit, or Musk ownership percentage disclosed.", "Medium", "xAI"),
        EvidenceItem("A2", "Company announcements", "xAI raises $6B Series C", "xAI", "2024-12-23", "xAI $6B Series C, Colossus 100,000 Nvidia Hopper GPUs, planned 200,000 GPU expansion, and Grok/X product link.", "Company announcement", "x.ai/news/series-c", "Company announcement does not disclose audited financials.", "Medium", "xAI"),
        EvidenceItem("A3", "Company announcements", "Grok 3 Beta - The Age of Reasoning Agents", "xAI", "2025-02-19", "Grok 3 model positioning, reasoning emphasis, infrastructure, and company-reported benchmark claims.", "Company announcement", "x.ai/news/grok-3", "Company benchmarks should be independently verified before investment-grade conclusions.", "Medium-low", "xAI"),
        EvidenceItem("N1", "Reputable secondary sources", "Neuralink's first implant partly detached from patient's brain", "The Guardian", "2024-05-09", "Neuralink first-human implant issue, restored functionality claim, and approximate $5B valuation context.", "Reputable secondary", "theguardian.com/technology/article/2024/may/09/neuralink-brain-chip-implant", "Valuation and clinical details are incomplete; not a regulatory filing or audited disclosure.", "Low-medium", "Neuralink"),
        EvidenceItem("B1", "Company announcements", "The Boring Company homepage, Vegas Loop, and Prufrock pages", "The Boring Company", "accessed 2026-06-11", "Loop operating claims, Vegas Loop approvals/passenger count, and Prufrock speed/cost targets.", "Company source", "boringcompany.com; /vegas-loop; /prufrock", "Company claims are not independently audited and do not disclose revenue or Musk ownership.", "Low-medium", "The Boring Company"),
    ]


def meta_economic_ownership_pct() -> float:
    """Return Zuckerberg's approximate economic share exposure from proxy counts."""
    owned = META_FACTS["zuckerberg_class_a_2025"] + META_FACTS["zuckerberg_class_b_2025"]
    outstanding = META_FACTS["class_a_outstanding_2025"] + META_FACTS["class_b_outstanding_2025"]
    return owned / outstanding


def mark_gold_evidence() -> list[EvidenceItem]:
    """Curated official source set for the Phase 4B Mark Zuckerberg report."""
    return [
        EvidenceItem(
            "M1",
            "Public company filings",
            "Meta Platforms 2024 Form 10-K",
            "U.S. SEC / Meta Platforms, Inc.",
            "2025-01-30",
            "2024 revenue, operating income, net income, operating cash flow, purchases of property and equipment, repurchases, dividends, Family of Apps and Reality Labs segment economics, DAP, ad impressions, ad pricing, AI/infrastructure, and risk factors.",
            "Primary filing",
            "https://www.sec.gov/Archives/edgar/data/1326801/000132680125000017/meta-20241231.htm",
            "Company filing supports Meta economics, not Zuckerberg's complete personal balance sheet.",
            "High",
            "Meta Platforms",
        ),
        EvidenceItem(
            "M2",
            "Public company filings",
            "Meta Platforms 2025 definitive proxy statement",
            "U.S. SEC / Meta Platforms, Inc.",
            "2025-04-17",
            "Zuckerberg beneficial ownership, 99.8% of Class B shares, 61.0% total voting power, common share counts, dual-class voting rights, controlled-company status, and pledged-share disclosures.",
            "Primary filing",
            "https://www.sec.gov/Archives/edgar/data/1326801/000132680125000040/meta-20250417.htm",
            "Proxy ownership is as of April 1, 2025 and does not disclose Zuckerberg's full taxes, trusts, or personal debt.",
            "High",
            "Meta Platforms",
        ),
        EvidenceItem(
            "M3",
            "Investor relations / company sources",
            "Meta fourth quarter and full year 2024 results release",
            "Meta Investor Relations",
            "2025-01-29",
            "Official investor-relations release for Q4 and full-year 2024 results; used as company-source support alongside the 10-K.",
            "Company investor-relations source",
            "https://s21.q4cdn.com/399680738/files/doc_financials/2024/q4/Meta-12-31-2024-Exhibit-99-1-Final.pdf",
            "Financial values in this report are anchored to the 10-K where possible.",
            "High",
            "Meta Platforms",
        ),
        EvidenceItem(
            "M4",
            "Public company filings",
            "Meta 2024 Form 10-K risk-factor and business overview sections",
            "U.S. SEC / Meta Platforms, Inc.",
            "2025-01-30",
            "Official risk language for privacy, targeting, mobile platform changes, regulation, AI risk, user engagement, ad demand, and dual-class voting control.",
            "Primary filing",
            "https://www.sec.gov/Archives/edgar/data/1326801/000132680125000017/meta-20241231.htm",
            "Risk-factor language identifies exposures but cannot quantify every probability or personal wealth effect.",
            "High",
            "Meta Platforms",
        ),
        EvidenceItem(
            "M5",
            "Public company filings",
            "Meta 2024 Form 10-K segment table",
            "U.S. SEC / Meta Platforms, Inc.",
            "2025-01-30",
            "Family of Apps 2024 revenue of $162.355B and operating income of $87.109B; Reality Labs 2024 revenue of $2.146B and operating loss of $17.729B.",
            "Primary filing",
            "https://www.sec.gov/Archives/edgar/data/1326801/000132680125000017/meta-20241231.htm",
            "Segment data is company-reported and does not attribute a separate market value to each segment.",
            "High",
            "Meta Platforms",
        ),
    ]


def arnault_gold_evidence() -> list[EvidenceItem]:
    """Curated official source set for the Phase 5C Bernard Arnault report."""
    urd_2024 = "https://lvmh-com.cdn.prismic.io/lvmh-com/Z-PY3HdAxsiBv6wN_UniversalRegistrationDocument2024.pdf"
    key_figures = "https://www.lvmh.com/en/investors/key-figures"
    key_documents = "https://www.lvmh.com/en/press/key-documents"
    lvmh_share = "https://www.lvmh.com/en/investors/lvmh_share"
    return [
        EvidenceItem(
            "LVMH1",
            "Public company filings",
            "LVMH 2024 Universal Registration Document - financial highlights",
            "LVMH Moet Hennessy Louis Vuitton SE",
            "2025-03-26",
            "2024 revenue of EUR 84.683B, profit from recurring operations of EUR 19.571B, operating free cash flow of EUR 10.478B, operating investments of EUR 7.478B, net financial debt of EUR 9.228B, 6,307 stores, and regional revenue mix.",
            "Primary filing",
            urd_2024,
            "Company filing supports LVMH economics, not the Arnault family's complete personal balance sheet.",
            "High",
            entity="LVMH financials",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "LVMH2",
            "Public company filings",
            "LVMH 2024 Universal Registration Document - business-group segment tables",
            "LVMH Moet Hennessy Louis Vuitton SE",
            "2025-03-26",
            "2024 revenue and profit from recurring operations by business group: Wines and Spirits, Fashion and Leather Goods, Perfumes and Cosmetics, Watches and Jewelry, Selective Retailing, and Other activities/eliminations.",
            "Primary filing",
            urd_2024,
            "Segment data is company-reported and does not assign separate market values to each brand or Maison.",
            "High",
            entity="LVMH business groups",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "LVMH3",
            "Public company filings",
            "LVMH 2024 Universal Registration Document - share capital and voting rights",
            "LVMH Moet Hennessy Louis Vuitton SE",
            "2025-03-26",
            "Arnault family group 2024 LVMH position: 245,173,934 shares, 49.00% of share capital, and 64.81% of voting rights exercisable at shareholder meetings; Christian Dior SE represented 41.87% of share capital and 56.77% of exercisable voting rights.",
            "Primary filing",
            urd_2024,
            "The disclosure describes the family group and controlled companies, not each person's taxes, trusts, debt, liquidity, or intra-family allocation.",
            "High",
            entity="Arnault family group / LVMH control",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "LVMH4",
            "Public company filings",
            "LVMH 2024 Universal Registration Document - Fashion and Leather Goods business model",
            "LVMH Moet Hennessy Louis Vuitton SE",
            "2025-03-26",
            "Fashion and Leather Goods 2024 revenue of EUR 41.060B, profit from recurring operations of EUR 15.230B, operating margin of 37.1%, 95% retail revenue mix, more than 2,300 stores, Louis Vuitton workshop footprint, and distribution-control strategy.",
            "Primary filing",
            urd_2024,
            "Business group data does not disclose brand-level revenue/profit for Louis Vuitton or Dior separately.",
            "High",
            entity="Fashion and Leather Goods",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "LVMH5",
            "Investor relations / company sources",
            "LVMH key figures",
            "LVMH Moet Hennessy Louis Vuitton SE",
            "accessed 2026-06-16",
            "Official investor-relations key figures for income statement, business group revenue/profit, geography, operating investments, net debt, free cash flow, and non-financial indicators.",
            "Company investor-relations source",
            key_figures,
            "The online key figures page can update after the Forbes 2025 annual-list date; the report anchors 2024 filing claims to the URD where possible.",
            "High",
            entity="LVMH investor relations",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "LVMH6",
            "Investor relations / company sources",
            "LVMH key documents page",
            "LVMH Moet Hennessy Louis Vuitton SE",
            "accessed 2026-06-16",
            "Official locator for the 2024 Universal Registration Document, 2025 Annual Report, and other LVMH reports.",
            "Company investor-relations source",
            key_documents,
            "This is a source locator; financial claims use the URD/key-figures entries.",
            "High",
            entity="LVMH publications",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "LVMH7",
            "Investor relations / company sources",
            "LVMH share page",
            "LVMH Moet Hennessy Louis Vuitton SE",
            "accessed 2026-06-16",
            "Official investor page for LVMH share information, dividend trends, and shareholder-structure locator.",
            "Company investor-relations source",
            lvmh_share,
            "The web table can update after the 2024 URD; ownership/control claims in this report use the URD snapshot.",
            "Medium-high",
            entity="LVMH share",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "LVMH8",
            "Company announcements",
            "LVMH group and Maisons pages",
            "LVMH Moet Hennessy Louis Vuitton SE",
            "accessed 2026-06-16",
            "Official LVMH site evidence that the group is family-run and has more than 75 Maisons rooted in six sectors, with example Maisons including Louis Vuitton, Christian Dior, Fendi, Tiffany, Sephora, Moet & Chandon, Hennessy, Bulgari, and Guerlain.",
            "Company source",
            "https://www.lvmh.com/en",
            "Company pages support portfolio structure and brand examples, not standalone brand valuations.",
            "High",
            entity="LVMH Maisons",
            accessed_at="2026-06-16",
        ),
    ]


def oracle_gold_evidence() -> list[EvidenceItem]:
    """Curated official source set for the Phase 5D Larry Ellison report."""
    sec_10k = "https://www.sec.gov/Archives/edgar/data/1341439/000095017024075605/orcl-20240531.htm"
    sec_proxy = "https://www.sec.gov/Archives/edgar/data/1341439/000119312524226139/d860928ddef14a.htm"
    ir_home = "https://investor.oracle.com/"
    return [
        EvidenceItem(
            "ORCL1",
            "Public company filings",
            "Oracle Corporation fiscal 2024 Form 10-K consolidated financials and cash flow",
            "U.S. SEC / Oracle Corporation",
            "2024-06-20",
            "Fiscal 2024 total revenues of $52.961B, total operating margin of $15.353B, net income of $10.467B, operating cash flow of $18.673B, capex of $6.866B, free cash flow of $11.807B, common-stock repurchases of $1.202B, and dividends paid of $4.391B.",
            "Primary filing",
            sec_10k,
            "Company filing supports Oracle economics, not Ellison's complete personal balance sheet.",
            "High",
            entity="Oracle financials",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "ORCL2",
            "Public company filings",
            "Oracle Corporation fiscal 2024 Form 10-K revenue and margin tables",
            "U.S. SEC / Oracle Corporation",
            "2024-06-20",
            "Fiscal 2024 revenue mix: cloud and license $44.464B, hardware $3.066B, services $5.431B; cloud services $19.774B, license support $19.609B, cloud license and on-premise license $5.081B; cloud services and license support total $39.383B.",
            "Primary filing",
            sec_10k,
            "Revenue categories are company-reported and do not assign separate market values to database, OCI, applications, Java, NetSuite, or Cerner.",
            "High",
            entity="Oracle revenue categories",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "ORCL3",
            "Public company filings",
            "Oracle Corporation fiscal 2024 Form 10-K business description and risk factors",
            "U.S. SEC / Oracle Corporation",
            "2024-06-20",
            "Official support for Oracle cloud, license support, hardware, services, data-center capacity risk, AI product risk, competition, renewal importance, and technology-transition risk.",
            "Primary filing",
            sec_10k,
            "Risk factors identify exposures but do not quantify every probability or Ellison-specific wealth effect.",
            "High",
            entity="Oracle business and risk factors",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "ORCL4",
            "Public company filings",
            "Oracle Corporation 2024 definitive proxy statement - Ellison ownership",
            "U.S. SEC / Oracle Corporation",
            "2024-09-25",
            "Ellison beneficial ownership of 1,153,232,353 shares, approximately 41.6% of Oracle common stock as of September 16, 2024; proxy also identifies him as founder, chairman, CTO, and largest stockholder.",
            "Primary filing",
            sec_proxy,
            "Beneficial ownership is not a full personal wealth bridge and does not disclose taxes, trusts, liquidity discounts, private assets, or complete debt terms.",
            "High",
            entity="Larry Ellison / Oracle ownership",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "ORCL5",
            "Public company filings",
            "Oracle Corporation 2024 definitive proxy statement - pledging policy and arrangements",
            "U.S. SEC / Oracle Corporation",
            "2024-09-25",
            "As of September 16, 2024, Ellison had pledged 277,000,000 Oracle shares as collateral for certain personal indebtedness; proxy states these were term-loan collateral, not margin-account collateral, and subject to board/committee review.",
            "Primary filing",
            sec_proxy,
            "Proxy describes pledged shares and governance review, not all loan balances, terms, covenants, trusts, taxes, or personal liquidity.",
            "High",
            entity="Larry Ellison pledged shares",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "ORCL6",
            "Investor relations / company sources",
            "Oracle Investor Relations",
            "Oracle Corporation",
            "accessed 2026-06-16",
            "Official investor-relations source locator for annual reports, proxies, SEC filings, quarterly results, governance, and investor materials.",
            "Company investor-relations source",
            ir_home,
            "IR page is a source locator; exact financial claims in this report are anchored to SEC filings where possible.",
            "High",
            entity="Oracle investor relations",
            accessed_at="2026-06-16",
        ),
        EvidenceItem(
            "ORCL7",
            "Public company filings",
            "Oracle Corporation fiscal 2024 Form 10-K segment margin table",
            "U.S. SEC / Oracle Corporation",
            "2024-06-20",
            "Fiscal 2024 margin by business: cloud and license margin $28.514B, hardware margin $1.915B, and services margin $0.916B.",
            "Primary filing",
            sec_10k,
            "Margins are business-category measures, not GAAP operating income by product line or separate valuations.",
            "High",
            entity="Oracle segment margins",
            accessed_at="2026-06-16",
        ),
    ]


def evidence_from_pack(pack: pd.DataFrame, existing_keys: set[str]) -> list[EvidenceItem]:
    """Convert rows from a manual/sourced evidence pack into EvidenceItem objects."""
    items: list[EvidenceItem] = []
    for index, row in pack.iterrows():
        key = clean(row.get("citation_key"))
        if not key:
            key = f"E{index + 1}"
        while key in existing_keys:
            key = f"{key}X"
        existing_keys.add(key)
        supports = clean(row.get("claim_supported")) or clean(row.get("evidence_note")) or "Manual evidence-pack row."
        limitation = clean(row.get("limitations")) or "Manual evidence should be reviewed before final use."
        if key == "BA3":
            supports = (
                "LVMH publications page is retained as the official locator for reports and investor publications. "
                "In this enriched v2 report, exact LVMH financial, ownership/control, segment, and risk claims are supported by the extracted LVMH1-LVMH8 source cards."
            )
            limitation = (
                "Locator row only. Use the extracted 2024 Universal Registration Document and official LVMH source cards for exact report claims."
            )
        title = clean(row.get("source_title")) or f"Evidence pack row for {clean(row.get('entity_or_asset')) or clean(row.get('person_name'))}"
        category = clean(row.get("evidence_category")) or "Manual evidence pack"
        if category not in EVIDENCE_CATEGORIES:
            category = "Manual evidence pack"
        items.append(
            EvidenceItem(
                key=key,
                category=category,
                title=title,
                publisher=clean(row.get("publisher")),
                date=clean(row.get("publication_date")) or clean(row.get("valuation_date")),
                supports=supports,
                reliability=clean(row.get("reliability_tier")) or clean(row.get("source_type")) or "Manual evidence pack",
                locator=clean(row.get("source_url")) or clean(row.get("source_file")) or "manual evidence pack",
                limitation=limitation,
                confidence=clean(row.get("confidence_level")) or "Medium",
                entity=clean(row.get("entity_or_asset")),
            )
        )
    return items


def build_evidence_items(context: EnrichedContext) -> list[EvidenceItem]:
    """Build evidence items from project data, optional pack, and gold-sample additions."""
    items = project_evidence(context)
    existing = {item.key for item in items}
    forbes_uri = clean(context.person.get("forbes_uri"))
    if forbes_uri == "jeff-bezos":
        for item in bezos_gold_evidence():
            if item.key not in existing:
                items.append(item)
                existing.add(item.key)
    if forbes_uri == "elon-musk":
        for item in elon_gold_evidence():
            if item.key not in existing:
                items.append(item)
                existing.add(item.key)
    if forbes_uri == "mark-zuckerberg":
        for item in mark_gold_evidence():
            if item.key not in existing:
                items.append(item)
                existing.add(item.key)
    if forbes_uri == "bernard-arnault":
        for item in arnault_gold_evidence():
            if item.key not in existing:
                items.append(item)
                existing.add(item.key)
    if forbes_uri == "larry-ellison":
        for item in oracle_gold_evidence():
            if item.key not in existing:
                items.append(item)
                existing.add(item.key)
        return items
    items.extend(evidence_from_pack(context.evidence_pack, existing))
    return items


def add_para(doc: Document, text: str = "", style: str | None = None, italic: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        run.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def set_table_font(table, size: float = 8.5) -> None:
    """Use compact table text for appendix readability."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)


def add_matrix_table(doc: Document, columns: list[str], rows: list[list[str]], *, font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = clean(value)[:900]
    set_table_font(table, font_size)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], *, font_size: float = 9) -> None:
    add_matrix_table(doc, ["Field", "Value"], [[key, value] for key, value in rows], font_size=font_size)


def add_confidence_note(doc: Document, confidence: str, note: str, keys: str) -> None:
    add_para(doc, f"Confidence note: {confidence}. {note} Evidence: {keys}.", italic=True)


def evidence_pack_rows_for_body(pack: pd.DataFrame) -> list[list[str]]:
    """Return compact, citation-keyed evidence rows for draft body sections."""
    if pack.empty:
        return []
    rows: list[list[str]] = []
    for _, row in pack.iterrows():
        key = clean(row.get("citation_key"))
        if not key:
            continue
        claim = clean(row.get("claim_supported")) or clean(row.get("business_model_notes")) or "Evidence-pack row available."
        rows.append(
            [
                f"[{key}]",
                clean(row.get("entity_or_asset")) or clean(row.get("person_name")),
                clean(row.get("evidence_category")) or "Manual evidence pack",
                claim,
                clean(row.get("confidence_level")) or "Medium",
            ]
        )
    return rows


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 14.5), ("Heading 2", 12), ("Heading 3", 10.5)]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.bold = True
        style.font.size = Pt(size)
        style.paragraph_format.keep_with_next = True
    if "Caption" in styles:
        styles["Caption"].font.name = "Aptos"
        styles["Caption"].font.size = Pt(8.5)
        styles["Caption"].font.italic = True


def create_charts(context: EnrichedContext, evidence: list[EvidenceItem], charts_dir: Path, variant: str) -> list[Path]:
    """Create reusable report charts and return paths inserted into the DOCX."""
    charts_dir.mkdir(parents=True, exist_ok=True)
    person = context.person
    slug = slugify(str(person["name"]))
    history = context.history
    created: list[Path] = []

    if not history.empty:
        wealth_path = charts_dir / f"{slug}_wealth_history_{variant}.png"
        fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=180)
        ax.plot(history["year"], history["net_worth_usd_b"], marker="o", linewidth=2.2, color="#1f77b4")
        ax.fill_between(history["year"], history["net_worth_usd_b"], color="#1f77b4", alpha=0.12)
        ax.set_title(f"{person['name']} Forbes Annual Net Worth History")
        ax.set_xlabel("Forbes annual list year")
        ax.set_ylabel("Net worth, USD billions")
        ax.grid(True, alpha=0.25)
        fig.tight_layout()
        fig.savefig(wealth_path, bbox_inches="tight")
        plt.close(fig)
        created.append(wealth_path)

    if clean(person.get("forbes_uri")) == "jeff-bezos":
        exposure = amazon_beneficial_ownership_pct()
        sens_path = charts_dir / f"{slug}_amazon_sensitivity_{variant}.png"
        deltas = list(range(-500, 501, 100))
        fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=180)
        ax.axhline(0, color="#222222", linewidth=0.8)
        ax.axvline(0, color="#222222", linewidth=0.8)
        ax.plot(
            deltas,
            [delta * exposure for delta in deltas],
            marker="o",
            color="#ff9900",
            label=f"{AMAZON_FACTS['bezos_beneficial_ownership_pct_2025']:.1f}% proxy beneficial ownership",
        )
        ax.set_title("Illustrative Amazon Market-Cap Sensitivity to Bezos Wealth")
        ax.set_xlabel("Change in Amazon market capitalization, USD billions")
        ax.set_ylabel("Approx. change in Amazon-linked wealth, USD billions")
        ax.grid(True, alpha=0.25)
        ax.legend(frameon=False, fontsize=8)
        ax.text(
            -490,
            max(delta * exposure for delta in deltas) * 0.72,
            "Illustrative only: beneficial ownership is not a full\npersonal balance-sheet bridge. Excludes taxes, sales,\ntrusts, liquidity, debt, and private assets.",
            fontsize=8,
            bbox={"boxstyle": "round,pad=0.35", "facecolor": "#f7f7f7", "edgecolor": "#bbbbbb"},
        )
        fig.tight_layout()
        fig.savefig(sens_path, bbox_inches="tight")
        plt.close(fig)
        created.append(sens_path)

        segment_path = charts_dir / f"{slug}_amazon_segment_economics_{variant}.png"
        labels = ["North America", "International", "AWS"]
        revenue = [
            AMAZON_FACTS["north_america_revenue_2024_b"],
            AMAZON_FACTS["international_revenue_2024_b"],
            AMAZON_FACTS["aws_revenue_2024_b"],
        ]
        operating = [
            AMAZON_FACTS["north_america_operating_income_2024_b"],
            AMAZON_FACTS["international_operating_income_2024_b"],
            AMAZON_FACTS["aws_operating_income_2024_b"],
        ]
        x_positions = range(len(labels))
        fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=180)
        width = 0.35
        ax.bar([x - width / 2 for x in x_positions], revenue, width=width, label="Net sales", color="#2563eb")
        ax.bar([x + width / 2 for x in x_positions], operating, width=width, label="Operating income", color="#f97316")
        ax.axhline(0, color="#222222", linewidth=0.8)
        ax.set_xticks(list(x_positions))
        ax.set_xticklabels(labels)
        ax.set_ylabel("USD billions")
        ax.set_title("Amazon 2024 Segment Economics")
        ax.legend(frameon=False, fontsize=8)
        ax.grid(True, axis="y", alpha=0.25)
        fig.tight_layout()
        fig.savefig(segment_path, bbox_inches="tight")
        plt.close(fig)
        created.append(segment_path)

    if clean(person.get("forbes_uri")) == "elon-musk":
        sens_path = charts_dir / f"{slug}_tesla_sensitivity_{variant}.png"
        deltas = list(range(-500, 501, 100))
        exposure_2024 = 20.5
        exposure_2025 = 19.8
        fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=180)
        ax.axhline(0, color="#222222", linewidth=0.8)
        ax.axvline(0, color="#222222", linewidth=0.8)
        ax.plot(deltas, [delta * exposure_2024 / 100 for delta in deltas], marker="o", label="20.5% 2024 proxy exposure", color="#d62728")
        ax.plot(deltas, [delta * exposure_2025 / 100 for delta in deltas], marker="s", label="19.8% later 2025 proxy exposure", color="#9467bd")
        ax.set_title("Illustrative Tesla Market-Cap Sensitivity to Musk Wealth")
        ax.set_xlabel("Change in Tesla market capitalization, USD billions")
        ax.set_ylabel("Approx. change in Tesla-linked wealth, USD billions")
        ax.grid(True, alpha=0.25)
        ax.legend(frameon=False, fontsize=8)
        ax.text(
            -490,
            82,
            "Illustrative only: excludes taxes, pledge terms,\noption exercise prices, discounts, debt, and liquidity.",
            fontsize=8,
            bbox={"boxstyle": "round,pad=0.35", "facecolor": "#f7f7f7", "edgecolor": "#bbbbbb"},
        )
        fig.tight_layout()
        fig.savefig(sens_path, bbox_inches="tight")
        plt.close(fig)
        created.append(sens_path)

    if clean(person.get("forbes_uri")) == "mark-zuckerberg":
        exposure = meta_economic_ownership_pct()
        sens_path = charts_dir / f"{slug}_meta_sensitivity_{variant}.png"
        deltas = list(range(-500, 501, 100))
        fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=180)
        ax.axhline(0, color="#222222", linewidth=0.8)
        ax.axvline(0, color="#222222", linewidth=0.8)
        ax.plot(deltas, [delta * exposure for delta in deltas], marker="o", color="#0f766e", label=f"Approx. {exposure * 100:.1f}% economic exposure")
        ax.set_title("Illustrative Meta Market-Cap Sensitivity to Zuckerberg Wealth")
        ax.set_xlabel("Change in Meta market capitalization, USD billions")
        ax.set_ylabel("Approx. change in Zuckerberg Meta-linked wealth, USD billions")
        ax.grid(True, alpha=0.25)
        ax.legend(frameon=False, fontsize=8)
        ax.text(
            -490,
            max(delta * exposure for delta in deltas) * 0.72,
            "Illustrative only: voting control is not economic ownership.\nExcludes taxes, trusts, pledges, liquidity, debt, and sales.",
            fontsize=8,
            bbox={"boxstyle": "round,pad=0.35", "facecolor": "#f7f7f7", "edgecolor": "#bbbbbb"},
        )
        fig.tight_layout()
        fig.savefig(sens_path, bbox_inches="tight")
        plt.close(fig)
        created.append(sens_path)

        segment_path = charts_dir / f"{slug}_segment_economics_{variant}.png"
        labels = ["Family of Apps", "Reality Labs"]
        revenue = [META_FACTS["family_revenue_2024_b"], META_FACTS["reality_labs_revenue_2024_b"]]
        operating = [META_FACTS["family_operating_income_2024_b"], META_FACTS["reality_labs_operating_loss_2024_b"]]
        x_positions = range(len(labels))
        fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=180)
        width = 0.35
        ax.bar([x - width / 2 for x in x_positions], revenue, width=width, label="Revenue", color="#2563eb")
        ax.bar([x + width / 2 for x in x_positions], operating, width=width, label="Operating income (loss)", color="#f97316")
        ax.axhline(0, color="#222222", linewidth=0.8)
        ax.set_xticks(list(x_positions))
        ax.set_xticklabels(labels)
        ax.set_ylabel("USD billions")
        ax.set_title("Meta 2024 Segment Economics")
        ax.legend(frameon=False, fontsize=8)
        ax.grid(True, axis="y", alpha=0.25)
        fig.tight_layout()
        fig.savefig(segment_path, bbox_inches="tight")
        plt.close(fig)
        created.append(segment_path)

    if clean(person.get("forbes_uri")) == "bernard-arnault":
        exposure = lvmh_family_group_ownership_pct()
        sens_path = charts_dir / f"{slug}_lvmh_family_group_sensitivity_{variant}.png"
        deltas = list(range(-200, 201, 50))
        fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=180)
        ax.axhline(0, color="#222222", linewidth=0.8)
        ax.axvline(0, color="#222222", linewidth=0.8)
        ax.plot(
            deltas,
            [delta * exposure for delta in deltas],
            marker="o",
            color="#8b5a2b",
            label=f"{LVMH_FACTS['family_group_share_capital_pct_2024']:.1f}% family group share-capital exposure",
        )
        ax.set_title("Illustrative LVMH Market-Cap Sensitivity to Family Group Value")
        ax.set_xlabel("Change in LVMH market capitalization, EUR billions")
        ax.set_ylabel("Approx. family group look-through change, EUR billions")
        ax.grid(True, alpha=0.25)
        ax.legend(frameon=False, fontsize=8)
        ax.text(
            -195,
            max(delta * exposure for delta in deltas) * 0.65,
            "Illustrative only: family group exposure is not a\ncomplete personal net-worth bridge. Excludes taxes,\ntrusts, debt, liquidity, discounts, and intra-family allocation.",
            fontsize=8,
            bbox={"boxstyle": "round,pad=0.35", "facecolor": "#f7f7f7", "edgecolor": "#bbbbbb"},
        )
        fig.tight_layout()
        fig.savefig(sens_path, bbox_inches="tight")
        plt.close(fig)
        created.append(sens_path)

        segment_path = charts_dir / f"{slug}_lvmh_segment_economics_{variant}.png"
        rows = LVMH_FACTS["segment_rows_2024"]
        labels = [row[0].replace(" and ", " & ").replace(" activities", "") for row in rows]
        revenue = [row[1] for row in rows]
        operating = [row[2] for row in rows]
        x_positions = range(len(labels))
        fig, ax = plt.subplots(figsize=(9.2, 4.8), dpi=180)
        width = 0.35
        ax.bar([x - width / 2 for x in x_positions], revenue, width=width, label="Revenue", color="#8b5a2b")
        ax.bar([x + width / 2 for x in x_positions], operating, width=width, label="Profit from recurring operations", color="#1f7a4d")
        ax.axhline(0, color="#222222", linewidth=0.8)
        ax.set_xticks(list(x_positions))
        ax.set_xticklabels(labels, rotation=22, ha="right")
        ax.set_ylabel("EUR billions")
        ax.set_title("LVMH 2024 Segment Economics")
        ax.legend(frameon=False, fontsize=8)
        ax.grid(True, axis="y", alpha=0.25)
        fig.tight_layout()
        fig.savefig(segment_path, bbox_inches="tight")
        plt.close(fig)
        created.append(segment_path)

    if clean(person.get("forbes_uri")) == "larry-ellison":
        exposure = oracle_ellison_ownership_pct()
        sens_path = charts_dir / f"{slug}_oracle_sensitivity_{variant}.png"
        deltas = list(range(-500, 501, 100))
        fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=180)
        ax.axhline(0, color="#222222", linewidth=0.8)
        ax.axvline(0, color="#222222", linewidth=0.8)
        ax.plot(
            deltas,
            [delta * exposure for delta in deltas],
            marker="o",
            color="#c74634",
            label=f"{ORACLE_FACTS['ellison_beneficial_ownership_pct_2024']:.1f}% proxy beneficial ownership",
        )
        ax.set_title("Illustrative Oracle Market-Cap Sensitivity to Ellison Wealth")
        ax.set_xlabel("Change in Oracle market capitalization, USD billions")
        ax.set_ylabel("Approx. change in Oracle-linked wealth, USD billions")
        ax.grid(True, alpha=0.25)
        ax.legend(frameon=False, fontsize=8)
        ax.text(
            -490,
            max(delta * exposure for delta in deltas) * 0.68,
            "Illustrative only: beneficial ownership is not a full\npersonal balance-sheet bridge. Excludes taxes, pledges,\ndebt terms, trusts, liquidity, and private assets.",
            fontsize=8,
            bbox={"boxstyle": "round,pad=0.35", "facecolor": "#f7f7f7", "edgecolor": "#bbbbbb"},
        )
        fig.tight_layout()
        fig.savefig(sens_path, bbox_inches="tight")
        plt.close(fig)
        created.append(sens_path)

        segment_path = charts_dir / f"{slug}_oracle_segment_economics_{variant}.png"
        labels = ["Cloud & license", "Hardware", "Services"]
        revenue = [
            ORACLE_FACTS["cloud_and_license_revenue_2024_b"],
            ORACLE_FACTS["hardware_revenue_2024_b"],
            ORACLE_FACTS["services_revenue_2024_b"],
        ]
        margin = [
            ORACLE_FACTS["cloud_and_license_margin_2024_b"],
            ORACLE_FACTS["hardware_margin_2024_b"],
            ORACLE_FACTS["services_margin_2024_b"],
        ]
        x_positions = range(len(labels))
        fig, ax = plt.subplots(figsize=(8.2, 4.4), dpi=180)
        width = 0.35
        ax.bar([x - width / 2 for x in x_positions], revenue, width=width, label="Revenue", color="#2563eb")
        ax.bar([x + width / 2 for x in x_positions], margin, width=width, label="Business margin", color="#c74634")
        ax.axhline(0, color="#222222", linewidth=0.8)
        ax.set_xticks(list(x_positions))
        ax.set_xticklabels(labels)
        ax.set_ylabel("USD billions")
        ax.set_title("Oracle FY2024 Revenue and Business Margin")
        ax.legend(frameon=False, fontsize=8)
        ax.grid(True, axis="y", alpha=0.25)
        fig.tight_layout()
        fig.savefig(segment_path, bbox_inches="tight")
        plt.close(fig)
        created.append(segment_path)

    map_path = charts_dir / f"{slug}_business_empire_map_{variant}.png"
    if clean(person.get("forbes_uri")) == "jeff-bezos":
        entities = [
            "Amazon retail\nmarketplace",
            "Third-party seller\nservices",
            "AWS cloud + AI\ninfrastructure",
            "Advertising\ncommerce media",
            "Prime subscriptions\nloyalty flywheel",
            "Logistics + fulfillment\ncapital moat",
            "Blue Origin\nprivate optionality",
            "Bezos Expeditions\nselected investments",
        ]
        positions = [(0.16, 0.61), (0.39, 0.61), (0.63, 0.61), (0.86, 0.61), (0.25, 0.33), (0.50, 0.33), (0.73, 0.33), (0.88, 0.33)]
        colors = ["#f59e0b", "#2563eb", "#0f766e", "#7c3aed", "#e11d48", "#0891b2", "#111827", "#6b7280"]
        fig, ax = plt.subplots(figsize=(10.8, 5.9), dpi=180)
        ax.axis("off")
        ax.text(
            0.5,
            0.86,
            "Bezos wealth engine\nAmazon public equity + operating platform + private optionality",
            ha="center",
            va="center",
            color="white",
            fontsize=10,
            fontweight="bold",
            bbox={"boxstyle": "round,pad=0.55", "facecolor": "#111827", "edgecolor": "#111827"},
        )
        for idx, entity in enumerate(entities):
            x, y = positions[idx]
            ax.text(
                x,
                y,
                entity,
                ha="center",
                va="center",
                color="white",
                fontsize=8.4,
                fontweight="bold",
                bbox={"boxstyle": "round,pad=0.42", "facecolor": colors[idx], "edgecolor": colors[idx]},
            )
            ax.annotate("", xy=(x, y + 0.06), xytext=(0.5, 0.78), arrowprops={"arrowstyle": "->", "color": "#4b5563", "lw": 1.0})
        ax.text(0.5, 0.10, "Forbes wealth link: Amazon beneficial ownership is the visible core; private assets are not precisely valued here.", ha="center", fontsize=9)
    elif clean(person.get("forbes_uri")) == "mark-zuckerberg":
        entities = [
            "Facebook\nsocial graph",
            "Instagram\nvisual attention",
            "WhatsApp + Messenger\nprivate messaging",
            "Threads\ntext/social graph extension",
            "Ads auction + data\nmonetization engine",
            "Meta AI + infra\nrecommendation/creation",
            "Reality Labs\nQuest/wearables option",
        ]
        positions = [(0.18, 0.58), (0.42, 0.58), (0.68, 0.58), (0.88, 0.58), (0.27, 0.30), (0.55, 0.30), (0.80, 0.30)]
        colors = ["#1877f2", "#c13584", "#25d366", "#111827", "#f59e0b", "#7c3aed", "#0f766e"]
        fig, ax = plt.subplots(figsize=(10.4, 5.8), dpi=180)
        ax.axis("off")
        ax.text(
            0.5,
            0.84,
            "Meta platform wealth engine\nnetwork effects + attention + ad auction + AI infrastructure",
            ha="center",
            va="center",
            color="white",
            fontsize=10,
            fontweight="bold",
            bbox={"boxstyle": "round,pad=0.55", "facecolor": "#111827", "edgecolor": "#111827"},
        )
        for idx, entity in enumerate(entities):
            x, y = positions[idx]
            ax.text(
                x,
                y,
                entity,
                ha="center",
                va="center",
                color="white",
                fontsize=8.8,
                fontweight="bold",
                bbox={"boxstyle": "round,pad=0.45", "facecolor": colors[idx], "edgecolor": colors[idx]},
            )
            ax.annotate("", xy=(x, y + 0.06), xytext=(0.5, 0.76), arrowprops={"arrowstyle": "->", "color": "#4b5563", "lw": 1.1})
        ax.text(0.5, 0.08, "Forbes wealth link: Meta equity exposure; governance link: dual-class voting control", ha="center", fontsize=9)
    elif clean(person.get("forbes_uri")) == "bernard-arnault":
        entities = [
            "Family control\nAgache + Christian Dior",
            "Fashion & Leather\nLouis Vuitton, Dior, Fendi",
            "Selective Retailing\nSephora + travel retail",
            "Wines & Spirits\nMoet, Hennessy, champagne",
            "Watches & Jewelry\nTiffany, Bulgari, TAG Heuer",
            "Perfumes & Cosmetics\nDior, Guerlain, Fenty",
            "Other activities\nBelmond, Cheval Blanc, media",
        ]
        positions = [(0.50, 0.64), (0.20, 0.43), (0.43, 0.43), (0.67, 0.43), (0.86, 0.43), (0.33, 0.20), (0.67, 0.20)]
        colors = ["#111827", "#8b5a2b", "#0f766e", "#7c2d12", "#475569", "#be123c", "#6b7280"]
        fig, ax = plt.subplots(figsize=(10.8, 5.9), dpi=180)
        ax.axis("off")
        ax.text(
            0.5,
            0.86,
            "Arnault wealth engine\nfamily-controlled LVMH portfolio + luxury scarcity + public equity compounding",
            ha="center",
            va="center",
            color="white",
            fontsize=10,
            fontweight="bold",
            bbox={"boxstyle": "round,pad=0.55", "facecolor": "#111827", "edgecolor": "#111827"},
        )
        for idx, entity in enumerate(entities):
            x, y = positions[idx]
            ax.text(
                x,
                y,
                entity,
                ha="center",
                va="center",
                color="white",
                fontsize=8.6,
                fontweight="bold",
                bbox={"boxstyle": "round,pad=0.45", "facecolor": colors[idx], "edgecolor": colors[idx]},
            )
            ax.annotate("", xy=(x, y + 0.06), xytext=(0.5, 0.78), arrowprops={"arrowstyle": "->", "color": "#4b5563", "lw": 1.0})
        ax.text(0.5, 0.07, "Forbes wealth link: family group control of listed LVMH; exact personal liquidity, trusts, debt, and intra-family allocation remain undisclosed.", ha="center", fontsize=8.6)
    elif clean(person.get("forbes_uri")) == "larry-ellison":
        entities = [
            "Oracle Database\nmission-critical systems",
            "Cloud services\nOCI + SaaS/PaaS/IaaS",
            "License support\nrecurring maintenance",
            "Enterprise apps\nFusion, NetSuite, verticals",
            "Java + middleware\nsoftware ecosystem",
            "Hardware + engineered systems\ninfrastructure stack",
            "Cerner / health\nenterprise data optionality",
            "Founder ownership\npublic equity sensitivity",
        ]
        positions = [(0.18, 0.58), (0.43, 0.58), (0.67, 0.58), (0.88, 0.58), (0.25, 0.31), (0.50, 0.31), (0.73, 0.31), (0.88, 0.31)]
        colors = ["#c74634", "#2563eb", "#0f766e", "#7c3aed", "#f59e0b", "#475569", "#0891b2", "#111827"]
        fig, ax = plt.subplots(figsize=(10.8, 5.9), dpi=180)
        ax.axis("off")
        ax.text(
            0.5,
            0.85,
            "Ellison wealth engine\nOracle enterprise software lock-in + OCI/AI infrastructure + founder public equity",
            ha="center",
            va="center",
            color="white",
            fontsize=10,
            fontweight="bold",
            bbox={"boxstyle": "round,pad=0.55", "facecolor": "#111827", "edgecolor": "#111827"},
        )
        for idx, entity in enumerate(entities):
            x, y = positions[idx]
            ax.text(
                x,
                y,
                entity,
                ha="center",
                va="center",
                color="white",
                fontsize=8.4,
                fontweight="bold",
                bbox={"boxstyle": "round,pad=0.42", "facecolor": colors[idx], "edgecolor": colors[idx]},
            )
            ax.annotate("", xy=(x, y + 0.06), xytext=(0.5, 0.77), arrowprops={"arrowstyle": "->", "color": "#4b5563", "lw": 1.0})
        ax.text(0.5, 0.08, "Forbes wealth link: large Oracle beneficial ownership; personal debt, pledges, taxes, trusts, and private assets are not fully modeled.", ha="center", fontsize=8.6)
    else:
        entities = [clean(item.entity) for item in evidence if clean(item.entity)]
        entities = list(dict.fromkeys(entities))[:6]
        if not entities:
            entities = [clean(context.metrics.get("key_asset_or_company")) or clean(person.get("primary_company_or_asset")) or "Core asset"]
        fig, ax = plt.subplots(figsize=(9.4, 5.4), dpi=180)
        ax.axis("off")
        ax.text(
            0.5,
            0.83,
            f"{person['name']}\nwealth engine: {clean(context.metrics.get('wealth_engine_category'))}",
            ha="center",
            va="center",
            color="white",
            fontsize=10,
            fontweight="bold",
            bbox={"boxstyle": "round,pad=0.55", "facecolor": "#111827", "edgecolor": "#111827"},
        )
        colors = ["#e11d48", "#2563eb", "#7c3aed", "#059669", "#d97706", "#0891b2"]
        positions = [(0.17, 0.55), (0.50, 0.55), (0.83, 0.55), (0.28, 0.27), (0.60, 0.27), (0.83, 0.27)]
        for idx, entity in enumerate(entities):
            x, y = positions[idx]
            ax.text(
                x,
                y,
                entity,
                ha="center",
                va="center",
                color="white",
                fontsize=9,
                fontweight="bold",
                bbox={"boxstyle": "round,pad=0.5", "facecolor": colors[idx % len(colors)], "edgecolor": colors[idx % len(colors)]},
            )
            ax.annotate("", xy=(x, y + 0.06), xytext=(0.5, 0.76), arrowprops={"arrowstyle": "->", "color": "#4b5563", "lw": 1.2})
        ax.text(0.5, 0.08, "Evidence-driven asset map. Unknown or unsourced assets remain excluded.", ha="center", fontsize=9)
    if clean(person.get("forbes_uri")) != "mark-zuckerberg":
        pass
    fig.tight_layout()
    fig.savefig(map_path, bbox_inches="tight")
    plt.close(fig)
    created.append(map_path)
    return created


def add_picture_with_caption(doc: Document, path: Path, caption: str) -> None:
    if path.exists():
        picture = doc.add_picture(str(path), width=Inches(6.7))
        picture._inline.docPr.set("descr", caption)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = doc.add_paragraph(style="Caption" if "Caption" in doc.styles else None)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(f"Figure: {caption} (file: {path.name}).")
        run.italic = True


def add_title_and_reader_guide(doc: Document, context: EnrichedContext, variant: str) -> None:
    person = context.person
    year = person_year(person)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{person['name']}: Business Empire and Wealth Engine Analysis")
    run.bold = True
    run.font.size = Pt(20)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        f"Forbes {year} rank {person_rank(person)} | {money(person_net_worth(person))} estimated net worth | {variant}"
    )
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "Body citations use short keys such as [F1]. Full source details, limitations, and confidence notes are grouped in the evidence appendix."
    ).italic = True

    doc.add_heading("How to read this report", level=1)
    add_bullets(
        doc,
        [
            "Read short citation keys in the body as pointers to the grouped evidence appendix, not as raw URLs.",
            "Treat Forbes net worth as an annual estimate and private-company valuations as marked evidence, not audited liquidity.",
            "Separate operating engines from option-value engines: evidence-backed cash flow deserves higher confidence than private or long-horizon optionality.",
            "Confidence notes describe how much primary evidence supports each section.",
        ],
    )
    add_confidence_note(doc, "High", "This section describes methodology and citation conventions.", "[F1][D1]")

    doc.add_heading("Report Map", level=1)
    add_para(
        doc,
        "This report map is a reader-facing guide, not a raw Word TOC placeholder. The document still uses Word Heading styles so a formal table of contents can be inserted later if needed.",
    )
    add_bullets(
        doc,
        [
            "Executive thesis",
            "Wealth equation and asset map",
            "Wealth history, CAGR, and exponential-fit result",
            "Major operating businesses and option-value assets",
            "First-principles analysis, risks, lessons, and peer patterns",
            "Grouped evidence appendix, evidence gaps, confidence table, and data limitations",
        ],
    )
    add_confidence_note(doc, "High", "Heading structure is generated by the DOCX renderer.", "[F1]")


def add_generic_core_sections(doc: Document, context: EnrichedContext, evidence_keys: str, charts: list[Path]) -> None:
    person = context.person
    metric = context.metrics
    year = person_year(person)
    multiple = metric_multiple(metric)
    archetype = clean(metric.get("wealth_engine_category")) or "Other/unclear"
    display_archetype = archetype
    if clean(person.get("forbes_uri")) == "bernard-arnault":
        archetype = "Luxury/retail brand ownership"
        display_archetype = "Luxury/retail brand ownership + family-controlled holding company + LVMH portfolio compounding"
    if clean(person.get("forbes_uri")) == "larry-ellison":
        archetype = "Founder/operator public equity"
        display_archetype = "Enterprise software/database lock-in + Oracle cloud infrastructure + founder public-equity wealth engine"
    guidance = ARCHETYPE_GUIDANCE.get(archetype, ARCHETYPE_GUIDANCE["Other/unclear"])

    doc.add_heading("1. Executive thesis", level=1)
    taxonomy_note = " The label is a project taxonomy category, not a legal monopoly finding." if "monopoly" in display_archetype.casefold() else ""
    add_para(
        doc,
        f"{person['name']} appears in the Forbes {year} Top 100 with estimated net worth of {money(person_net_worth(person))}. "
        f"The structured project classifies the primary wealth engine for this report as {display_archetype}. The analytical starting point is not biography; it is the wealth equation: "
        f"{guidance['wealth_equation']}.{taxonomy_note}",
    )
    add_confidence_note(doc, "Medium-high", "Forbes rank and wealth estimate are sourced; deeper operating analysis depends on person-specific evidence packs.", "[F1][D1]")

    doc.add_heading("2. Wealth equation and asset map", level=1)
    add_kv_table(
        doc,
        [
            (f"Forbes {year} net worth", money(person_net_worth(person))),
            ("Forbes source of wealth", clean(person.get("source_of_wealth"))),
            ("Primary company or asset", clean(person.get("primary_company_or_asset"))),
            ("Project wealth engine", display_archetype),
            ("Equation lens", guidance["wealth_equation"]),
            ("Required evidence before final report", guidance["required_evidence"]),
        ],
    )
    add_confidence_note(doc, "Medium", "Top-level asset fields are sourced; ownership, debt, trusts, and discounts require additional evidence.", "[F1]")

    doc.add_heading("3. Wealth history, CAGR, and exponential-fit result", level=1)
    add_kv_table(
        doc,
        [
            ("First year observed", clean(metric.get("first_year_observed"))),
            ("First net worth", money(metric.get("first_net_worth_usd_b"))),
            ("Years observed", clean(metric.get("years_observed"))),
            (f"Wealth multiple first-to-{year}", clean(round(float(multiple), 2)) if clean(multiple) else "not available"),
            ("Nominal CAGR", pct(metric.get("CAGR_nominal"))),
            ("Log-linear slope", clean(round(float(metric.get("log_linear_growth_slope")), 4)) if clean(metric.get("log_linear_growth_slope")) else "not available"),
            ("Exponential fit R^2", clean(round(float(metric.get("exponential_fit_r2")), 3)) if clean(metric.get("exponential_fit_r2")) else "not available"),
            ("Estimated doubling time", f"{float(metric.get('estimated_doubling_time_years')):,.2f} years" if clean(metric.get("estimated_doubling_time_years")) else "not available"),
            ("Largest drawdown", pct(metric.get("largest_drawdown_pct"))),
        ],
    )
    if charts:
        add_picture_with_caption(doc, charts[0], "Forbes annual net worth history from project processed data")
    add_confidence_note(doc, "High", "Metrics are derived from annual Forbes observations; do not claim true exponential growth unless the fit and data coverage support it.", "[D1]")

    doc.add_heading("4. Business empire timeline", level=1)
    if context.history.empty:
        add_para(doc, "No annual history rows are available for this person.")
    else:
        changes = context.history[["year", "net_worth_usd_b"]].copy()
        changes["annual_change"] = changes["net_worth_usd_b"].diff()
        top_changes = changes.dropna().sort_values("annual_change", key=lambda series: series.abs(), ascending=False).head(5)
        for _, row in top_changes.iterrows():
            add_para(doc, f"{int(row['year'])}: Forbes annual net worth changed by {money(row['annual_change'])}.")
    add_confidence_note(doc, "Medium", "Timeline is wealth-snapshot based until company event sources are added.", "[D1]")

    doc.add_heading("5. Archetype-specific analysis lens", level=1)
    add_kv_table(
        doc,
        [
            ("Wealth engine archetype", display_archetype),
            ("Analysis focus", guidance["analysis_focus"]),
            ("Source standard", guidance["required_evidence"]),
        ],
    )
    add_confidence_note(doc, "Medium", "Archetype guidance is a reusable analytical rubric; final claims need person-specific citations.", evidence_keys)


def add_bezos_specific_sections(doc: Document, context: EnrichedContext, charts: list[Path]) -> None:
    """Add a deep public-equity/platform report for Jeff Bezos."""
    person = context.person
    metric = context.metrics
    ownership_pct = amazon_beneficial_ownership_pct()
    aws_margin = amazon_aws_operating_margin()
    aws_share_of_operating_income = AMAZON_FACTS["aws_operating_income_2024_b"] / AMAZON_FACTS["operating_income_2024_b"]

    doc.add_heading("6. Bezos wealth equation: Amazon public equity first", level=1)
    add_para(
        doc,
        f"The cleanest Bezos wealth equation is: Forbes net worth of {money(person.get('net_worth_2025_usd_b'))} ≈ Amazon beneficial/economic exposure x Amazon public market value, plus or minus taxes, share sales, philanthropy, trusts, liquidity discounts, debt, and private assets. "
        f"The 2025 Amazon proxy reports Bezos beneficial ownership of {AMAZON_FACTS['bezos_beneficial_shares_2025']:,} shares, or {AMAZON_FACTS['bezos_beneficial_ownership_pct_2025']:.1f}% of the class. "
        f"The same proxy footnote says {AMAZON_FACTS['bezos_sole_voting_no_investment_shares_2025']:,} of those shares had sole voting power and no investment power, so the report treats the proxy figure as beneficial ownership evidence, not a complete personal balance-sheet model.",
    )
    add_kv_table(
        doc,
        [
            ("Forbes 2025 estimate", money(person.get("net_worth_2025_usd_b"))),
            ("Visible core asset", "Amazon common stock exposure"),
            ("Proxy beneficial shares", f"{AMAZON_FACTS['bezos_beneficial_shares_2025']:,}"),
            ("Proxy percent of class", f"{AMAZON_FACTS['bezos_beneficial_ownership_pct_2025']:.1f}%"),
            ("Shares with sole voting/no investment power", f"{AMAZON_FACTS['bezos_sole_voting_no_investment_shares_2025']:,}"),
            ("What remains outside the equation", "Taxes, trusts, debt, liquidity discounts, share sales, philanthropy, Blue Origin, Bezos Expeditions, real estate, media ownership, and other private assets."),
        ],
        font_size=8.2,
    )
    add_confidence_note(doc, "High for Amazon proxy ownership; medium for personal net-worth bridge", "The proxy supports beneficial ownership and governance context; Forbes supplies the net-worth estimate; the exact personal balance sheet is not public.", "[F1][AMZ4]")

    doc.add_heading("7. Amazon business empire map", level=1)
    add_para(
        doc,
        "The Bezos fortune is not best understood as a generic e-commerce biography. Amazon became a multi-engine operating platform: first-party retail, third-party marketplace services, Prime/subscriptions, fulfillment/logistics, advertising, and AWS. "
        "The operating platform is public and filing-visible; Blue Origin and Bezos Expeditions are treated separately as private optionality because valuation, ownership percentage, and liquidity evidence are incomplete.",
    )
    if len(charts) > 3:
        add_picture_with_caption(doc, charts[3], "Amazon platform and Bezos private-optionality map")
    add_confidence_note(doc, "High for Amazon platform components; medium-low for private optionality", "Amazon platform components are filing-supported; private assets are acknowledged only where official pages or proxy biography support their existence.", "[AMZ2][AMZ3][AMZ4][BO1][BE1]")

    doc.add_heading("8. Amazon revenue mix and platform economics", level=1)
    add_para(
        doc,
        f"Amazon's 2024 10-K shows a scale profile that is much broader than retail. Consolidated net sales were {money_precise_b(AMAZON_FACTS['net_sales_2024_b'])}. "
        f"Revenue lines included online stores at {money_precise_b(AMAZON_FACTS['online_stores_revenue_2024_b'])}, third-party seller services at {money_precise_b(AMAZON_FACTS['third_party_seller_services_revenue_2024_b'])}, advertising services at {money_precise_b(AMAZON_FACTS['advertising_services_revenue_2024_b'])}, subscription services at {money_precise_b(AMAZON_FACTS['subscription_services_revenue_2024_b'])}, physical stores at {money_precise_b(AMAZON_FACTS['physical_stores_revenue_2024_b'])}, AWS at {money_precise_b(AMAZON_FACTS['aws_revenue_2024_b'])}, and other revenue at {money_precise_b(AMAZON_FACTS['other_revenue_2024_b'])}. "
        "That mix explains why Amazon can compound through multiple layers: customer demand, seller monetization, infrastructure, advertising intent, and enterprise cloud.",
    )
    add_matrix_table(
        doc,
        ["Revenue line", "2024 net sales", "Business interpretation"],
        [
            ["Online stores", money_precise_b(AMAZON_FACTS["online_stores_revenue_2024_b"]), "First-party retail demand and selection engine."],
            ["Third-party seller services", money_precise_b(AMAZON_FACTS["third_party_seller_services_revenue_2024_b"]), "Marketplace services, seller tools, fulfillment, and transaction-linked fees."],
            ["Advertising services", money_precise_b(AMAZON_FACTS["advertising_services_revenue_2024_b"]), "Commerce-media layer tied to high-intent shopping surfaces."],
            ["Subscription services", money_precise_b(AMAZON_FACTS["subscription_services_revenue_2024_b"]), "Prime and other recurring membership/media/service revenue."],
            ["AWS", money_precise_b(AMAZON_FACTS["aws_revenue_2024_b"]), "Enterprise cloud infrastructure and AI/machine-learning service base."],
        ],
        font_size=8.0,
    )
    add_confidence_note(doc, "High", "Revenue lines are extracted from Amazon's 2024 10-K table for net sales by groups of similar products and services.", "[AMZ2]")

    doc.add_heading("9. AWS as valuation-quality engine", level=1)
    add_para(
        doc,
        f"AWS changes the quality of Amazon's valuation story because it converts infrastructure capability into a high-income enterprise platform. In 2024, AWS reported {money_precise_b(AMAZON_FACTS['aws_revenue_2024_b'])} of net sales and {money_precise_b(AMAZON_FACTS['aws_operating_income_2024_b'])} of operating income, an operating margin of {pct(aws_margin)}. "
        f"AWS represented about {pct(aws_share_of_operating_income)} of Amazon consolidated operating income. This does not mean AWS is the whole company, but it explains why Amazon's wealth engine is different from a pure retailer: a large part of operating income comes from cloud infrastructure, developer tools, data services, and machine-learning services.",
    )
    add_para(
        doc,
        "The AI optionality belongs here rather than in a separate speculative startup story. Amazon's 10-K describes AWS as offering compute, storage, database, analytics, machine learning, and other services. The evidence supports AI/cloud infrastructure optionality; it does not support a separate Bezos AI valuation bridge.",
    )
    add_confidence_note(doc, "High for AWS figures; medium for valuation-quality interpretation", "AWS revenue and operating income are filing data; the valuation interpretation is analytical and should be read as a business model inference.", "[AMZ2][AMZ3]")

    doc.add_heading("10. Marketplace, Prime, advertising, and logistics flywheel", level=1)
    add_para(
        doc,
        "Amazon's compounding mechanism is a set of reinforcing loops. Marketplace selection attracts customers; customer demand attracts sellers; sellers buy services and advertising; Prime raises frequency and delivery expectations; fulfillment infrastructure supports selection, speed, and reliability; data and search surfaces create advertising demand. "
        "This is not the same as a social-network ad auction. Amazon's advertising layer is closer to commerce media: sellers and brands pay for placement near shopping intent.",
    )
    add_para(
        doc,
        f"The moat is capital-intensive. In 2024, Amazon reported fulfillment expense of {money_precise_b(AMAZON_FACTS['fulfillment_expense_2024_b'])}, technology and infrastructure expense of {money_precise_b(AMAZON_FACTS['technology_infrastructure_expense_2024_b'])}, and net purchases of property and equipment of {money_precise_b(AMAZON_FACTS['ppe_purchases_net_2024_b'])}. "
        "Those figures are why the Bezos archetype differs from a lighter software platform: much of the advantage is embedded in physical logistics, data centers, engineering systems, supplier relationships, and customer habit formation.",
    )
    add_confidence_note(doc, "High for expense/capex figures; medium for flywheel interpretation", "The 10-K supports the cost and revenue scale; flywheel language is an analytical synthesis of those filing-supported business lines.", "[AMZ1][AMZ2][AMZ3]")

    doc.add_heading("11. Financial statement linkage", level=1)
    add_para(
        doc,
        f"Amazon 2024 financials connect directly to Bezos's wealth engine because public markets value Amazon shares based on a mix of revenue growth, operating income, cash flow, capital intensity, and long-term optionality. The 2024 10-K reports consolidated net sales of {money_precise_b(AMAZON_FACTS['net_sales_2024_b'])}, operating income of {money_precise_b(AMAZON_FACTS['operating_income_2024_b'])}, net income of {money_precise_b(AMAZON_FACTS['net_income_2024_b'])}, operating cash flow of {money_precise_b(AMAZON_FACTS['operating_cash_flow_2024_b'])}, and free cash flow of {money_precise_b(AMAZON_FACTS['free_cash_flow_2024_b'])}. "
        "This report does not invent an AMZN market cap/date bridge; the sensitivity chart below is illustrative and keyed to the proxy beneficial-ownership percentage.",
    )
    if len(charts) > 2:
        add_picture_with_caption(doc, charts[2], "Amazon 2024 segment net sales and operating income from Form 10-K")
    add_matrix_table(
        doc,
        ["Segment", "2024 net sales", "2024 operating income", "Interpretation"],
        [
            ["North America", money_precise_b(AMAZON_FACTS["north_america_revenue_2024_b"]), money_precise_b(AMAZON_FACTS["north_america_operating_income_2024_b"]), "Large commerce and logistics engine with improving operating income."],
            ["International", money_precise_b(AMAZON_FACTS["international_revenue_2024_b"]), money_precise_b(AMAZON_FACTS["international_operating_income_2024_b"]), "Large but thinner global commerce engine; reached positive operating income in 2024."],
            ["AWS", money_precise_b(AMAZON_FACTS["aws_revenue_2024_b"]), money_precise_b(AMAZON_FACTS["aws_operating_income_2024_b"]), "Cloud profit pool and AI infrastructure platform."],
        ],
        font_size=8.0,
    )
    add_confidence_note(doc, "High", "Consolidated and segment financials are extracted from Amazon's 2024 Form 10-K.", "[AMZ1][AMZ2]")

    doc.add_heading("12. Amazon market-cap sensitivity to Bezos wealth", level=1)
    add_para(
        doc,
        f"With {AMAZON_FACTS['bezos_beneficial_ownership_pct_2025']:.1f}% proxy beneficial ownership, a $100B change in Amazon market capitalization would mechanically imply about ${100 * ownership_pct:,.1f}B of change in Bezos Amazon-linked beneficial-ownership value before taxes, sales, trusts, debt, liquidity, and private assets. "
        "That sensitivity is useful for intuition, not a final net-worth model. Forbes may include or exclude private assets and discounts differently, and the proxy footnote means beneficial ownership should not be collapsed into direct liquid ownership.",
    )
    if len(charts) > 1:
        add_picture_with_caption(doc, charts[1], "Illustrative Amazon market-cap sensitivity to Bezos Amazon-linked wealth")
    add_confidence_note(doc, "Medium", "The ownership percentage is sourced; the market-cap sensitivity is a linear illustration, not a sourced Forbes valuation model.", "[AMZ4]")

    doc.add_heading("13. Blue Origin, Bezos Expeditions, and media/private optionality", level=1)
    add_para(
        doc,
        "Blue Origin matters strategically but is not treated as a precise wealth bridge in this report. The Amazon proxy identifies Bezos as Blue Origin's founder and as owner of The Washington Post; Blue Origin's own pages describe a reusable-rocket mission, New Shepard, New Glenn, lunar systems, engines, and space infrastructure objectives. "
        "Bezos Expeditions' official page lists selected investments and links to Blue Origin, The Washington Post, philanthropy, and other investments. None of those sources discloses a current Bezos ownership percentage, audited financials, private valuation, or liquidity terms.",
    )
    add_matrix_table(
        doc,
        ["Asset / vehicle", "Evidence-supported claim", "What remains unknown"],
        [
            ["Blue Origin", "Official pages support mission, reusable-rocket/product map, and New Glenn context.", "Bezos ownership percentage, valuation, revenue, cash burn, profitability, debt, and liquidity."],
            ["Bezos Expeditions", "Official page lists selected investments and links to Bezos-related assets/initiatives.", "Current values, ownership percentages, realized gains, and materiality to Forbes net worth."],
            ["The Washington Post", "Amazon proxy states Bezos owns The Washington Post.", "Current value, profitability, debt, and whether it materially affects Forbes net worth."],
        ],
        font_size=8.0,
    )
    add_confidence_note(doc, "Medium for existence; low for valuation", "Official sources establish the existence/strategic context of private assets but do not provide valuation or ownership mechanics.", "[AMZ4][BO1][BO2][BE1]")

    doc.add_heading("14. Wealth history, drawdowns, and exponential-style fit", level=1)
    add_para(
        doc,
        f"The Forbes annual sequence has 25 valid observations from 2001 to 2025. Bezos moves from {money(metric.get('first_net_worth_usd_b'))} in {int(metric.get('first_year_observed'))} to {money(person.get('net_worth_2025_usd_b'))} in 2025, a {float(metric.get('wealth_multiple_first_to_2025')):,.1f}x multiple and nominal endpoint CAGR of {pct(metric.get('CAGR_nominal'))}. "
        f"The log-linear fit has slope {float(metric.get('log_linear_growth_slope')):,.4f}, R^2 {float(metric.get('exponential_fit_r2')):,.3f}, and model-implied doubling time of {float(metric.get('estimated_doubling_time_years')):,.2f} years. This is a strong descriptive fit to annual Forbes observations, not proof of true exponential growth or a forecast.",
    )
    add_para(
        doc,
        f"The volatility is material: the project metrics show annual growth volatility of {pct(metric.get('volatility_of_annual_growth'))}, max one-year gain of {money(metric.get('max_one_year_gain_usd_b'))}, max one-year loss of {money(metric.get('max_one_year_loss_usd_b'))}, and largest drawdown of {pct(metric.get('largest_drawdown_pct'))}. The 2021-2023 drawdown and 2024-2025 rebound should be read as public-equity valuation sensitivity, not cash profit or loss.",
    )
    add_confidence_note(doc, "High for metrics; medium for interpretation", "Metrics are computed from Forbes annual snapshots; causality requires Amazon stock/market data and event evidence.", "[D1][AMZ4]")

    doc.add_heading("15. First-principles business analysis", level=1)
    add_para(
        doc,
        "The scarce resource Bezos organized was not only capital. It was a long-horizon operating system for customer demand: broad selection, low friction, delivery reliability, seller participation, cloud infrastructure, and willingness to reinvest cash flow into scale before short-term margins looked attractive. "
        "Marketplace plus logistics created commerce frequency; AWS created a separate enterprise infrastructure profit pool; advertising monetized purchase intent; Prime converted service reliability into habit and recurring demand.",
    )
    add_para(
        doc,
        "This differs from a Meta-style social-attention platform because Amazon's moat is more operational and capital-intensive: warehouses, sort centers, delivery, data centers, procurement, seller systems, and cloud infrastructure. It differs from a Tesla-style manufacturing/energy story because Amazon's physical assets support a multi-sided commerce and cloud platform rather than a narrower hardware-product adoption curve.",
    )
    add_confidence_note(doc, "Medium-high", "First-principles claims synthesize filing-supported revenue, segment, expense, and risk data; they remain analytical interpretations.", "[AMZ1][AMZ2][AMZ3]")

    doc.add_heading("16. Capital allocation and reinvestment pattern", level=1)
    add_para(
        doc,
        f"Amazon's public filings show a high-reinvestment engine. In 2024, operating cash flow was {money_precise_b(AMAZON_FACTS['operating_cash_flow_2024_b'])} and free cash flow was {money_precise_b(AMAZON_FACTS['free_cash_flow_2024_b'])} after net PPE purchases of {money_precise_b(AMAZON_FACTS['ppe_purchases_net_2024_b'])}. "
        "That pattern matters for the wealth engine: Amazon's value is not only today's earnings power, but the market's belief that reinvestment in logistics, cloud infrastructure, advertising systems, and AI services can keep expanding future profit pools.",
    )
    add_para(
        doc,
        "The report does not assert Bezos personally controls current capital allocation in the same way a CEO or dual-class controller would. The 2025 proxy identifies him as founder and Executive Chair, while Amazon has one class of common stock with equal voting rights. Strategic influence, governance role, and voting economics should therefore be kept separate.",
    )
    add_confidence_note(doc, "High for cash-flow and governance facts; medium for capital-allocation interpretation", "Cash-flow and governance claims are filing-supported; strategic influence is interpreted conservatively.", "[AMZ1][AMZ4]")

    doc.add_heading("17. Risks, counter-thesis, and fragility", level=1)
    add_para(
        doc,
        "The counter-thesis is straightforward: a Bezos fortune dominated by Amazon public equity is exposed to Amazon market-cap drawdowns. Amazon also carries capital-intensity risk, cloud-demand and AI-infrastructure risk, retail competition, seller/regulatory pressure, labor and fulfillment risk, cybersecurity and service reliability risk, international execution risk, and advertising demand cyclicality. "
        "The private-asset side adds opacity rather than certainty: Blue Origin may have enormous option value, but the current evidence does not disclose a financial bridge.",
    )
    add_confidence_note(doc, "High for public-company risk categories; medium for wealth impact", "Risk categories are filing-supported, but exact probability and Bezos-specific value impact are not disclosed.", "[AMZ3][BO1][BO2]")

    doc.add_heading("18. Comparable billionaire patterns", level=1)
    add_para(
        doc,
        "Bezos is the third archetype template for this project: founder public-equity wealth tied to a platform operating company. The comparison set should not copy the Elon/Mark templates. The Amazon pattern is a hybrid of commerce marketplace, logistics network, cloud infrastructure, and advertising monetization. "
        "Its key distinction is that public-market wealth is attached to a company that combines physical capital intensity with software/platform economics.",
    )
    add_confidence_note(doc, "Medium", "Comparable pattern is a project taxonomy interpretation grounded in Amazon filing data and Forbes wealth history.", "[D1][AMZ1][AMZ2]")

    doc.add_heading("19. Transferable business lessons", level=1)
    add_numbered(
        doc,
        [
            "Concentrated ownership of a public operating platform can create immense wealth when the platform keeps expanding profit pools.",
            "A marketplace flywheel becomes more defensible when it is backed by fulfillment infrastructure and seller tooling.",
            "Cloud infrastructure can transform the valuation quality of a retail-heavy company by adding enterprise recurring demand and higher operating income.",
            "Advertising is more powerful when tied to purchase intent, seller demand, and closed-loop commerce data.",
            "Capital intensity can be a moat when it improves service levels and raises entry barriers, but it also creates utilization and financing risk.",
            "A strong log-linear wealth fit is descriptive evidence of compounding, not proof that wealth grows smoothly or predictably.",
        ],
    )
    add_confidence_note(doc, "Medium-high", "Lessons are drawn from cited ownership, segment, cash-flow, and wealth-history evidence.", "[AMZ1][AMZ2][AMZ4][D1]")

    doc.add_heading("20. Bezos-specific evidence gaps", level=1)
    add_bullets(
        doc,
        [
            "Exact personal taxes, debt, trust structure, liquidity discounts, and current share-sale plans are not fully public.",
            "A precise Amazon market-cap bridge to the Forbes $215B estimate still needs a dated AMZN market cap/share price and Forbes methodology assumptions.",
            "Blue Origin ownership percentage, valuation, revenue, profit, debt, cash burn, and liquidity path are not disclosed in the sources used here.",
            "Bezos Expeditions and private investment values are listed as selected investments only; current values and ownership percentages are not disclosed.",
            "The Washington Post is identified as owned by Bezos in the proxy, but current value and financial contribution are not sourced here.",
        ],
    )
    add_confidence_note(doc, "High", "These gaps are explicit limitations and should prevent the report from becoming an unsupported personal balance-sheet model.", "[AMZ4][BO1][BE1]")


def add_elon_specific_sections(doc: Document, context: EnrichedContext, charts: list[Path]) -> None:
    metric = context.metrics
    doc.add_heading("6. Tesla public equity wealth engine", level=1)
    add_para(
        doc,
        "Tesla is the clearest public-equity engine in Musk's Forbes-level wealth. The company reports real operating scale, and the proxy evidence ties Musk's personal exposure to a large beneficial ownership position. "
        "Tesla's 2024 results included $97.7B revenue, $17.5B gross profit, $7.1B operating income, $7.1B net income, $14.9B operating cash flow, and $11.3B of property and equipment capex. "
        "That combination means the wealth driver is not just an EV brand; it is a public market capitalization attached to manufacturing scale, software/autonomy expectations, energy storage, charging, and robotics optionality.",
    )
    add_para(
        doc,
        "The proxy evidence is crucial. Tesla disclosed Musk beneficial ownership of 715.0M shares, or 20.5%, as of March 31, 2024, including exercisable options and pledged-share disclosure. "
        "A later 2025 proxy showed 717.3M beneficial shares and 19.8% ownership as of September 15, 2025, after the Forbes annual-list date. The report therefore uses the 2024 table as a pre-2025 anchor and the 2025 table as later context.",
    )
    if len(charts) > 1:
        add_picture_with_caption(doc, charts[1], "Illustrative Tesla market-cap sensitivity to Musk's Tesla-linked wealth")
    add_confidence_note(doc, "High for company financials and ownership table; medium for personal net-worth bridge", "SEC filings support Tesla economics and beneficial ownership, but taxes, debt, pledge mechanics, and liquidity discounts remain outside the filing data.", "[T1][T2][T3]")

    doc.add_heading("7. SpaceX and Starlink private-company valuation engine", level=1)
    add_para(
        doc,
        "SpaceX is the private-company engine. It attacks a physical bottleneck: the cost, cadence, and reliability of access to orbit. Reusable rockets lower launch cost and increase cadence; Starlink then turns launch capability into a satellite-broadband network with recurring service potential. "
        "NASA evidence supports the operating importance: commercial crew is a public-private transport model, and NASA awarded SpaceX a $2.89B Human Landing System contract tied to Starship architecture.",
    )
    add_para(
        doc,
        "The valuation bridge is less certain than Tesla because SpaceX is private. Late-2024 tender reporting put SpaceX around a $350B valuation. That mark matters because a large founder stake multiplied by a very large private mark can rival a public-equity fortune. "
        "But it should be read as a private-market price signal, not audited intrinsic value or guaranteed liquidity.",
    )
    add_confidence_note(doc, "High for NASA contracts/program context; medium for private valuation", "Operating evidence is primary/government-sourced; valuation and ownership are secondary and opaque.", "[S1][S2][S3]")

    doc.add_heading("8. X / Twitter ownership and strategic optionality", level=1)
    add_para(
        doc,
        "X is best treated as strategic optionality rather than a clean cash-flow engine in the current evidence set. Twitter's acquisition announcement put the deal at about $44B, with $54.20 per share consideration, debt and margin loan financing, and equity commitment. "
        "After the acquisition, the strategic thesis shifted from public social-media valuation to private platform control: distribution, identity, payments, ads, creator economics, and data for AI products.",
    )
    add_para(
        doc,
        "A later reported xAI-X all-stock transaction valued xAI at $80B and X at $33B equity value, but the exact Musk ownership allocation, X debt, minority investor economics, and valuation discounts are not fully disclosed. This section should therefore be read as optionality, not a precise wealth bridge.",
    )
    add_confidence_note(doc, "Medium-low", "The acquisition announcement is primary; later transaction values and current economics remain opaque.", "[X1][X2]")

    doc.add_heading("9. xAI valuation and AI infrastructure optionality", level=1)
    add_para(
        doc,
        "xAI is a compute, talent, and data option. It announced a $6B Series B in May 2024 and a $6B Series C in December 2024. The Series C announcement described Colossus as a 100,000 Nvidia Hopper GPU supercomputer with plans to double capacity to 200,000 GPUs. "
        "The economic thesis is that scarce frontier AI compute plus X distribution can become products, subscriptions, APIs, or enterprise/government systems.",
    )
    add_para(
        doc,
        "The risk is that infrastructure spending and benchmark claims are not the same as durable revenue. Ownership, revenue, gross margin, inference costs, customer retention, and governance economics are not disclosed enough for a final valuation bridge.",
    )
    add_confidence_note(doc, "Medium", "Funding and infrastructure claims come from company announcements; business model conversion remains uncertain.", "[A1][A2][A3]")

    doc.add_heading("10. Neuralink long-horizon biotech/neurotechnology optionality", level=1)
    add_para(
        doc,
        "Neuralink attacks a different scarce interface: high-bandwidth communication between neural signals and digital systems. The upside case is medical-device approval and eventual broader assistive or augmentation markets. "
        "The current evidence remains early and regulated, including first-human progress and reported technical setbacks around implant thread retraction/detachment followed by software adjustments.",
    )
    add_confidence_note(doc, "Low-medium", "Evidence supports early clinical activity and risk, but valuation, ownership, revenue, and regulatory path remain incomplete.", "[N1]")

    doc.add_heading("11. The Boring Company infrastructure optionality", level=1)
    add_para(
        doc,
        "The Boring Company follows the same constraint-removal pattern: reduce the cost and time of tunneling so cities can add transportation capacity without consuming surface land. Company materials describe Loop as an all-electric underground public transportation system, LVCC Loop as commercially operating, Vegas Loop approvals, and Prufrock targets for speed and cost. "
        "The wealth contribution is option-like because financials, ownership, and independently verified unit economics are not public.",
    )
    add_confidence_note(doc, "Low-medium", "Company operating claims are available, but audited economics and ownership are missing.", "[B1]")

    doc.add_heading("12. First-principles business analysis", level=1)
    add_para(
        doc,
        "The scarce resource Musk controls is not simply capital. It is a bundle of founder control, technical recruiting power, public attention, private-market access, risk tolerance, and willingness to organize teams around hard constraints before the economics are obvious. "
        "Tesla and SpaceX are operating engines with real products and customers. xAI, X, Neuralink, and Boring add increasingly uncertain option value. Capital markets rewarded the strategy because proof in one hard domain made investors more willing to underwrite adjacent hard-domain bets.",
    )
    add_matrix_table(
        doc,
        ["Asset", "Constraint attacked", "Cash-flow versus option-value character", "Capital-market reward logic"],
        [
            ["Tesla", "EV cost, manufacturing scale, battery/charging/software integration", "Operating engine plus autonomy/robotics option", "Public market rewards growth, brand, margin potential, and founder-led optionality"],
            ["SpaceX", "Launch cost/reliability and global broadband deployment", "Private operating engine plus Starship/Starlink option", "Private markets reward scarcity, technical lead, government validation, and Starlink TAM"],
            ["X/xAI", "Distribution, data, AI compute, products", "Mostly option-value at current evidence depth", "Frontier AI scarcity and X distribution can support high private valuations"],
            ["Neuralink/Boring", "Neural interface and tunneling cost", "Long-duration regulated options", "Potentially enormous outcomes, but high evidence uncertainty"],
        ],
    )
    add_confidence_note(doc, "Medium", "This is interpretive synthesis grounded in cited evidence; precise dollar attribution remains unavailable.", "[T1][S1][A2][B1][N1]")

    doc.add_heading("13. Transferable business lessons", level=1)
    add_numbered(
        doc,
        [
            "Ownership concentration is the compounding engine: large retained stakes matter more than salary.",
            "High-convexity bets can dominate a portfolio when the upside is tied to major bottlenecks.",
            "Public equity can become strategic leverage for capital cost, recruiting, credibility, and personal balance-sheet exposure.",
            "Vertical integration can turn impossible unit economics into hard but solvable engineering systems.",
            "Technology bottlenecks create valuation before mature profits appear, especially in batteries, rockets, AI compute, and tunneling.",
            "Narrative can be productive capital when it recruits talent and financing, but it also creates reputational and governance risk.",
            "Platform control compounds optionality across data, identity, distribution, payments, vehicles, satellites, and AI.",
        ],
    )
    add_confidence_note(doc, "Medium", "Lessons are strategic synthesis and should not be treated as deterministic rules.", "[T1][T2][S2][A2]")

    doc.add_heading("14. Risks and counter-thesis", level=1)
    add_bullets(
        doc,
        [
            "Tesla multiple compression would directly reduce the public-equity component of wealth.",
            "Pledged shares and personal leverage can magnify downside, though the full debt picture is not public.",
            "Private valuation opacity affects SpaceX, xAI, X, Neuralink, and Boring.",
            "Regulation touches autonomous driving, launches, spectrum, medical implants, tunneling, payments, and AI.",
            "Brand and political spillover can affect demand, advertising, hiring, and regulatory posture across assets.",
        ],
    )
    add_confidence_note(doc, "Medium-high", "Risk categories are supported by filings and source limitations, but exact personal exposure is incomplete.", "[T1][T2][S1][X2][N1][B1]")

    doc.add_heading("15. Comparable billionaire patterns", level=1)
    add_para(
        doc,
        "Musk is closest to a hybrid of founder/operator public equity and founder/operator private company. Compared with pure software/platform founders, he has more exposure to capital-intensive physical systems. "
        "Compared with luxury or inherited wealth, his fortune is less brand-dividend stable and more tied to public/private capital-market belief, technical execution, and regulatory permission.",
    )
    add_confidence_note(doc, "Medium", "Peer comparison is analytical synthesis from the project wealth-engine taxonomy.", "[D1]")


def add_mark_specific_sections(doc: Document, context: EnrichedContext, charts: list[Path]) -> None:
    """Add a deep public-company platform/network-effects report for Mark Zuckerberg."""
    person = context.person
    metric = context.metrics
    economic_pct = meta_economic_ownership_pct()

    doc.add_heading("6. Meta public-company platform wealth engine", level=1)
    add_para(
        doc,
        "Zuckerberg's Forbes-level wealth is primarily a public-company platform wealth engine, not a salary or private-company mark. "
        "The simplified equation is: net worth is approximately economic exposure to Meta Platforms multiplied by Meta's public market value, "
        "plus or minus taxes, share sales, pledges, trusts, liquidity constraints, and other assets. The governance equation is different: "
        "Meta's dual-class stock gives Zuckerberg voting control far above his economic ownership [M2].",
    )
    add_kv_table(
        doc,
        [
            ("Forbes 2025 net worth estimate", money(person.get("net_worth_2025_usd_b"))),
            ("Class A shares beneficially owned", f"{META_FACTS['zuckerberg_class_a_2025']:,}"),
            ("Class B shares beneficially owned", f"{META_FACTS['zuckerberg_class_b_2025']:,}"),
            ("Approx. economic exposure from proxy counts", pct(economic_pct)),
            ("Class B ownership", pct(META_FACTS["zuckerberg_class_b_pct_2025"], already_percent=True)),
            ("Total voting power", pct(META_FACTS["zuckerberg_voting_power_pct_2025"], already_percent=True)),
            ("Class A / Class B voting rights", "Class A: one vote per share; Class B: ten votes per share"),
            ("Pledged shares disclosed", f"{META_FACTS['zuckerberg_pledged_class_b_2025']:,} Class B shares"),
        ],
    )
    if len(charts) > 1:
        add_picture_with_caption(doc, charts[1], "Illustrative Meta market-cap sensitivity to Zuckerberg's Meta-linked economic exposure")
    add_confidence_note(doc, "High for proxy ownership/voting; medium for personal net-worth bridge", "The proxy supports share and voting-control data. Taxes, trusts, debt, liquidity, and non-Meta assets remain incomplete.", "[F1][M2]")

    doc.add_heading("7. Meta-specific business empire map", level=1)
    add_para(
        doc,
        "Meta's empire is a platform bundle: Facebook, Instagram, WhatsApp, Messenger, and Threads aggregate social identity, communication, creator distribution, and attention. "
        "The advertising platform then monetizes that attention through targeting, ranking, measurement, and auction-based ad placement. Meta AI and infrastructure investments are intended to support recommendations, ad tools, product creation, and developer/open-source positioning. "
        "Reality Labs is the long-duration option on VR, mixed reality, augmented reality, wearables, neural interfaces, and a possible next computing platform [M1][M4].",
    )
    add_matrix_table(
        doc,
        ["Surface", "Network effect", "Monetization role", "Switching cost / lock-in", "Key risk"],
        [
            ["Facebook", "Identity, groups, friends, events, feed", "Mature ad inventory and social graph signal", "History, groups, identity, and local/community ties", "Aging demographics, engagement shifts, privacy limits"],
            ["Instagram", "Creator/follower graph and visual discovery", "High-value attention surface for ads, Reels, commerce and creators", "Creator distribution, audience relationships, saved content, messaging", "Creator competition, short-form video pressure, monetization mix"],
            ["WhatsApp + Messenger", "Private messaging network and contact graph", "Lower direct ad monetization; business messaging and commerce option", "Contact-network utility and group chats; privacy expectations", "Encryption limits data use; monetization must preserve trust"],
            ["Threads", "Text/social graph extension", "Potential future public-conversation inventory; currently option-like", "Cross-app identity and creator portability from Instagram", "Unproven monetization and competitive attention market"],
            ["Ads auction + data", "More users and advertisers improve matching/liquidity", "Core auction/pricing/measurement engine", "Advertiser learning, campaign history, tooling, measurement workflows", "Regulation, OS/browser limits, signal loss, brand safety"],
            ["Meta AI + infrastructure", "Recommendation and creation tools can raise relevance", "Supports feed ranking, ad tools, creative generation, assistants", "Model quality, data, compute scale, developer ecosystem", "Capex, depreciation, model quality, legal/policy risk"],
            ["Reality Labs", "Potential future device/interface ecosystem", "Long-duration hardware/software/content option", "Hardware ecosystem and developer installed base if adoption scales", "Large losses, adoption uncertainty, platform timing"],
        ],
    )
    if len(charts) > 3:
        add_picture_with_caption(doc, charts[3], "Meta business empire map: apps, ad engine, AI infrastructure, and Reality Labs option value")
    add_confidence_note(doc, "High for product/segment description; medium for strategic synthesis", "Product and segment descriptions are filing-based; platform-interaction analysis is interpretive.", "[M1][M4]")

    doc.add_heading("8. Advertising auction and network-effects economics", level=1)
    add_para(
        doc,
        "Meta became valuable because it combined social graph, identity, attention, and advertiser demand at global scale. Users create content and signals; ranking and recommendation systems allocate attention; advertisers bid for outcomes; Meta's ad systems price and deliver placements across Facebook, Instagram, Messenger, and related surfaces. "
        "The 10-K states that substantially all revenue is generated from advertising placements, and that marketers purchase ads across Facebook, Instagram, Messenger, third-party apps, and websites [M1].",
    )
    add_para(
        doc,
        f"In 2024, Meta reported Family daily active people of {META_FACTS['family_daily_active_people_dec_2024_b']:.2f}B on average for December, ad impressions up {pct(META_FACTS['ad_impressions_growth_2024'])} year over year, and average price per ad up {pct(META_FACTS['average_price_per_ad_growth_2024'])} year over year. "
        "That combination shows both scale and pricing/auction recovery. The moat is not merely having apps; it is the feedback loop among identity, friends/followers, creator content, advertiser demand, measurement, and machine-learning ranking [M1][M4].",
    )
    add_para(
        doc,
        "The data advantage has limits. Meta's filings support reliance on user activity signals, targeting, measurement, and ad-delivery systems, but privacy regulation, mobile operating-system/browser changes, consent requirements, and product choices can reduce usable signals. "
        "WhatsApp also has a different trust and privacy model than the public feed surfaces, so its network value should not be treated as identical to Instagram or Facebook ad inventory [M4].",
    )
    add_confidence_note(doc, "High for reported metrics; medium for moat interpretation", "User/ad metrics are from the 10-K; network-effects interpretation is analytical.", "[M1][M4]")

    doc.add_heading("9. Financial statement and market-value linkage", level=1)
    free_cash_flow_proxy = META_FACTS["operating_cash_flow_2024_b"] - META_FACTS["capex_2024_b"]
    operating_margin = META_FACTS["operating_income_2024_b"] / META_FACTS["revenue_2024_b"]
    add_kv_table(
        doc,
        [
            ("2024 revenue", money(META_FACTS["revenue_2024_b"])),
            ("2024 operating income", money(META_FACTS["operating_income_2024_b"])),
            ("2024 operating margin", pct(operating_margin)),
            ("2024 net income", money(META_FACTS["net_income_2024_b"])),
            ("2024 operating cash flow", money(META_FACTS["operating_cash_flow_2024_b"])),
            ("2024 purchases of property and equipment", money(META_FACTS["capex_2024_b"])),
            ("2024 capex incl. finance-lease principal", money(META_FACTS["capex_including_finance_lease_principal_2024_b"])),
            ("2024 OCF less PPE purchases proxy", money(free_cash_flow_proxy)),
            ("2024 Meta-defined free cash flow", money(META_FACTS["meta_defined_free_cash_flow_2024_b"])),
            ("2024 cash used for Class A repurchases", money(META_FACTS["share_repurchases_2024_b"])),
            ("2024 repurchased and retired shares", money(META_FACTS["repurchased_and_retired_2024_b"])),
            ("2024 dividends paid", money(META_FACTS["dividends_2024_b"])),
        ],
    )
    add_para(
        doc,
        "The market-value linkage is direct: Meta's public equity value capitalizes its advertising cash machine, AI/recommendation investment, capital returns, and option-value projects. "
        "Zuckerberg's wealth therefore moves with Meta's public market capitalization, while his voting control affects governance and strategic time horizon rather than immediate economic ownership [M1][M2].",
    )
    add_confidence_note(doc, "High for financial statement values; medium for wealth sensitivity", "Financials are official 10-K values and checked against Meta's investor-relations release. Sensitivity is a derived illustration from proxy ownership counts.", "[M1][M2][M3]")

    doc.add_heading("10. Segment economics: Family of Apps versus Reality Labs", level=1)
    add_para(
        doc,
        f"Meta's 2024 segment table makes the wealth engine visible. Family of Apps generated {money(META_FACTS['family_revenue_2024_b'])} of revenue and {money(META_FACTS['family_operating_income_2024_b'])} of operating income. "
        f"Reality Labs generated {money(META_FACTS['reality_labs_revenue_2024_b'])} of revenue and an operating loss of {money(abs(META_FACTS['reality_labs_operating_loss_2024_b']))}. "
        "The Family of Apps funds the option-value portfolio; Reality Labs consumes capital to pursue the next computing platform [M5].",
    )
    if len(charts) > 2:
        add_picture_with_caption(doc, charts[2], "Meta 2024 segment economics: Family of Apps versus Reality Labs")
    add_confidence_note(doc, "High", "Segment revenue and operating income/loss are sourced from Meta's 2024 10-K segment table.", "[M5]")

    doc.add_heading("11. AI infrastructure and recommendation systems", level=1)
    add_para(
        doc,
        "Meta's AI investments are not separate from the ad engine; they are increasingly part of the ad engine. The 10-K says AI supports systems that rank content, the discovery engine, advertiser tools, generative AI experiences, and product-development efficiency. "
        "Meta also describes significant AI investments, including generative AI, content recommendations, ad delivery, targeting, measurement, new products, and Llama open-source models [M1][M4].",
    )
    add_para(
        doc,
        "The strategic question is whether AI spend expands the moat or compresses margins. If AI improves recommendations and ad return on investment, it can strengthen Meta's auction economics. If infrastructure and headcount spend outrun monetization, it lowers near-term free cash flow and can pressure the market multiple [M1][M4].",
    )
    add_confidence_note(doc, "High for filing language; medium for strategic inference", "The 10-K supports AI use cases and infrastructure risk; economic outcome remains uncertain.", "[M1][M4]")

    doc.add_heading("12. Reality Labs and metaverse option value", level=1)
    add_para(
        doc,
        "Reality Labs should be treated as long-duration option value, not the current cash engine. The 10-K describes RL as consumer hardware, software, content, virtual/mixed/augmented reality devices, social platforms, neural interfaces, and foundational next-computing-platform technologies. "
        "Meta also states that many products may only be fully realized in the next decade and that RL is expected to continue operating at a loss for the foreseeable future [M1][M5].",
    )
    add_para(
        doc,
        "Strategically, RL resembles a call option funded by the Family of Apps. If Meta controls the next interface layer, the option can be enormous. If VR/AR/wearables adoption remains niche or capital intensity stays high, the losses reduce current earnings power and investor patience [M5].",
    )
    add_confidence_note(doc, "High for losses and company description; medium-low for option-value payoff", "The operating loss is sourced; future platform value is uncertain.", "[M1][M5]")

    doc.add_heading("13. First-principles business analysis", level=1)
    add_para(
        doc,
        "The scarce resource Zuckerberg controls is not a conventional physical asset base. It is a global social-attention network with identity, relationships, creators, messages, content inventory, advertiser demand, data signals, and ranking infrastructure. "
        "Social graph plus attention plus identity creates inventory; advertiser demand creates monetization; measurement and AI improve allocation; scale creates operating leverage.",
    )
    add_para(
        doc,
        "This differs from physical-capital compounding. Meta's marginal distribution cost is digital, its operating leverage is software/ad-market driven, and its biggest constraints are attention, trust, privacy, regulation, platform policy, AI infrastructure, and competition rather than factory throughput or logistics capacity [M1][M4].",
    )
    add_confidence_note(doc, "Medium-high", "The resource-control thesis is synthesis anchored in Meta's disclosed products, ad model, AI systems, and risks.", "[M1][M4]")

    doc.add_heading("14. Forbes wealth history, drawdown, and rebound", level=1)
    add_para(
        doc,
        f"Project data show Zuckerberg first observed on the Forbes annual list in {int(metric['first_year_observed'])} at {money(metric['first_net_worth_usd_b'])}, reaching {money(person['net_worth_2025_usd_b'])} in 2025. "
        f"The observed multiple is {float(metric['wealth_multiple_first_to_2025']):.1f}x, nominal CAGR is {pct(metric['CAGR_nominal'])}, log-linear R^2 is {float(metric['exponential_fit_r2']):.3f}, and estimated doubling time is {float(metric['estimated_doubling_time_years']):.2f} years. "
        "The log-linear fit is descriptive only; R^2 of 0.852 is reasonably strong but visibly imperfect, not proof of true exponential growth. The processed history has no 2009 observation, and CAGR uses the 2008-2025 calendar span [D1].",
    )
    add_para(
        doc,
        "The history is not smooth. Forbes observations fell from $97.0B in 2021 to $67.3B in 2022 and $64.4B in 2023, then rebounded to $177.0B in 2024 and $216.0B in 2025. "
        "The 2022 drop was the largest one-year dollar loss in the processed series, while 2023 was the peak-to-trough drawdown low. Strategically, that drawdown/rebound matters because it is consistent with market concern about metaverse losses, privacy/targeting limits, and growth anxiety, followed by renewed confidence in cost discipline, ad recovery, AI/recommendation momentum, and operating leverage [D1][M1].",
    )
    add_confidence_note(doc, "High for Forbes observations; medium for strategic interpretation", "Annual wealth observations are project-derived; cause analysis is tied to Meta filings but not a single-factor proof.", "[D1][M1]")

    doc.add_heading("15. Risks and counter-thesis", level=1)
    add_bullets(
        doc,
        [
            "Privacy, data-use, data-combination, youth safety, competition, advertising, AI, and content-moderation regulation can reduce targeting, measurement, engagement, and ad demand [M4].",
            "Mobile operating-system and browser policy changes can limit data signals and ad effectiveness [M4].",
            "Reality Labs may continue to consume capital for years without creating a profitable next platform [M1][M5].",
            "AI infrastructure can raise capex, depreciation, and operating expense before monetization is proven [M1][M4].",
            "Dual-class control reduces outside Class A shareholder influence and concentrates key-person/governance risk [M2][M4].",
            "Zuckerberg pledge/tax/trust/liquidity details are incomplete, so the personal balance-sheet bridge remains approximate [M2].",
        ],
    )
    add_confidence_note(doc, "High for listed risk categories; medium for wealth impact", "Risk categories are filing-supported; magnitude and timing remain uncertain.", "[M2][M4]")

    doc.add_heading("16. Transferable business lessons", level=1)
    add_numbered(
        doc,
        [
            "A platform becomes a wealth engine when user identity, attention, content, and advertiser demand reinforce each other.",
            "Economic ownership and voting control are separate levers; Zuckerberg's voting power is much larger than his economic exposure.",
            "Network effects compound when ranking/recommendation systems increase relevance for users and ROI for advertisers.",
            "Operating leverage can produce enormous wealth rebounds when revenue growth returns after cost discipline.",
            "Long-duration options are affordable only when the core engine throws off enough cash to fund them.",
            "Regulation is not peripheral for social platforms; it can reshape targeting, measurement, data use, product design, and valuation.",
        ],
    )
    add_confidence_note(doc, "Medium-high", "Lessons are derived from cited platform economics, segment data, and ownership structure.", "[M1][M2][M5]")

    doc.add_heading("17. Comparable billionaire patterns", level=1)
    add_para(
        doc,
        "Zuckerberg is the clean second archetype in the project: a technology/platform monopoly and network-effects wealth engine. Compared with physical-asset founder fortunes, the asset is less industrial and more software/attention/ad-market driven. "
        "Compared with luxury or inherited fortunes, the engine is less stable brand pricing power and more platform scale, AI ranking, regulatory exposure, and public-equity multiple sensitivity. Compared with investor/capital allocator fortunes, the capital allocation is embedded inside Meta through buybacks, dividends, AI infrastructure, and Reality Labs option spending [D1][M1].",
    )
    add_confidence_note(doc, "Medium", "Comparable pattern is analytical synthesis from the project taxonomy and Meta evidence.", "[D1][M1][M5]")


def add_arnault_specific_sections(doc: Document, context: EnrichedContext, charts: list[Path]) -> None:
    """Add deep luxury/family-control sections for Bernard Arnault & family."""
    person = context.person
    metric = context.metrics
    fashion_margin = 15.230 / 41.060
    group_margin = LVMH_FACTS["profit_from_recurring_operations_2024_b"] / LVMH_FACTS["revenue_2024_b"]
    fashion_profit_share = 15.230 / LVMH_FACTS["profit_from_recurring_operations_2024_b"]
    family_exposure = lvmh_family_group_ownership_pct()

    doc.add_heading("6. Arnault wealth equation: family-controlled LVMH exposure", level=1)
    add_para(
        doc,
        f"The simplified wealth equation is: Forbes net worth for {person['name']} of {money(person.get('net_worth_2025_usd_b'))} is approximately family-controlled exposure to LVMH's listed equity value, plus or minus taxes, trusts, holding-company structure, dividends, liquidity discounts, debt, and other private assets. "
        f"LVMH's 2024 Universal Registration Document reports that the Arnault family group held {LVMH_FACTS['family_group_shares_2024']:,} shares, equal to {LVMH_FACTS['family_group_share_capital_pct_2024']:.2f}% of share capital and {LVMH_FACTS['family_group_exercisable_voting_rights_pct_2024']:.2f}% of voting rights exercisable at shareholder meetings. "
        "That is a family-group/control disclosure, not a complete personal balance-sheet model [F1][LVMH3].",
    )
    add_kv_table(
        doc,
        [
            ("Forbes 2025 estimate", money(person.get("net_worth_2025_usd_b"))),
            ("Visible core asset", "LVMH public equity and family-controlled holding structure"),
            ("Family group share-capital exposure", f"{LVMH_FACTS['family_group_share_capital_pct_2024']:.2f}%"),
            ("Family group voting rights exercisable at meetings", f"{LVMH_FACTS['family_group_exercisable_voting_rights_pct_2024']:.2f}%"),
            ("Christian Dior SE component", f"{LVMH_FACTS['christian_dior_share_capital_pct_2024']:.2f}% of LVMH share capital; {LVMH_FACTS['christian_dior_exercisable_voting_rights_pct_2024']:.2f}% of exercisable voting rights"),
            ("Not fully visible", "Personal taxes, trusts, debt, liquidity, dividends retained, holding-company discounts, and intra-family allocation."),
        ],
        font_size=8.0,
    )
    add_confidence_note(doc, "High for family group control; medium for personal wealth bridge", "The URD supports family-group capital/voting control but not a full personal balance sheet.", "[F1][LVMH3]")

    doc.add_heading("7. LVMH business empire map", level=1)
    add_para(
        doc,
        "The Arnault fortune is a listed luxury portfolio story, not a single-brand biography. LVMH describes itself as a family-run group with more than 75 Maisons rooted in six sectors. "
        "The operating system combines brand heritage, creative renewal, selective distribution, retail control, acquisition discipline, and group-level capital allocation across Fashion and Leather Goods, Selective Retailing, Wines and Spirits, Watches and Jewelry, Perfumes and Cosmetics, and Other activities [LVMH1][LVMH2][LVMH8].",
    )
    if len(charts) > 3:
        add_picture_with_caption(doc, charts[3], "LVMH portfolio and Arnault family-control map")
    add_confidence_note(doc, "High for portfolio categories; medium for brand-level economics", "LVMH discloses business groups and Maisons, but not full brand-level revenue/profit for each Maison.", "[LVMH2][LVMH8]")

    doc.add_heading("8. Segment economics: Fashion and Leather Goods is the profit core", level=1)
    add_para(
        doc,
        f"LVMH's 2024 filing shows the economic center of gravity. Group revenue was {eur_b(LVMH_FACTS['revenue_2024_b'])}, and profit from recurring operations was {eur_b(LVMH_FACTS['profit_from_recurring_operations_2024_b'])}, a recurring operating margin of {pct(group_margin)}. "
        f"Fashion and Leather Goods generated {eur_b(41.060)} of revenue and {eur_b(15.230)} of profit from recurring operations, implying a {pct(fashion_margin)} operating margin and roughly {pct(fashion_profit_share)} of group profit from recurring operations. "
        "This is why the wealth engine is luxury portfolio compounding, not generic retail scale [LVMH1][LVMH2][LVMH4].",
    )
    if len(charts) > 2:
        add_picture_with_caption(doc, charts[2], "LVMH 2024 segment revenue and profit from recurring operations")
    add_matrix_table(
        doc,
        ["Business group", "2024 revenue", "2024 profit from recurring operations", "Strategic interpretation"],
        [
            [name, eur_b(revenue), eur_b(profit), "High-margin luxury profit engine." if name == "Fashion and Leather Goods" else "Portfolio diversification and sector-specific cycle exposure."]
            for name, revenue, profit in LVMH_FACTS["segment_rows_2024"]
        ],
        font_size=7.8,
    )
    add_confidence_note(doc, "High", "Segment revenue and recurring operating profit are from the 2024 LVMH URD and official key-figures page.", "[LVMH1][LVMH2][LVMH5]")

    doc.add_heading("9. Fashion and Leather Goods operating model", level=1)
    add_para(
        doc,
        f"The Fashion and Leather Goods engine is operationally concrete. LVMH reports that this business group was {pct(LVMH_FACTS['fashion_retail_share_2024'])} retail by revenue in 2024 and had more than {LVMH_FACTS['fashion_stores_2024_min']:,} stores. "
        "The URD states that controlling distribution is a strategic priority because it lets the group retain retail margins, control brand image, control the sales environment, and maintain closer customer contact. "
        "Louis Vuitton's workshop footprint and supplier strategy show that the moat includes manufacturing know-how and quality control, not only advertising [LVMH4].",
    )
    add_bullets(
        doc,
        [
            "Pricing power comes from cultivated brand desirability, controlled supply, product excellence, and heritage.",
            "Distribution control protects brand image while retaining retail margin.",
            "Creative renewal refreshes heritage without commoditizing the Maison.",
            "Portfolio discipline reduces dependence on one product cycle while Fashion and Leather Goods remains the profit anchor.",
        ],
    )
    add_confidence_note(doc, "High for distribution and margin facts; medium for moat synthesis", "The operating facts are from the URD; the moat interpretation is first-principles analysis anchored to those facts.", "[LVMH4]")

    doc.add_heading("10. Ownership, control, and succession structure", level=1)
    add_para(
        doc,
        f"The key control fact is voting power. The Arnault family group held {LVMH_FACTS['family_group_share_capital_pct_2024']:.2f}% of share capital but {LVMH_FACTS['family_group_exercisable_voting_rights_pct_2024']:.2f}% of voting rights exercisable at shareholder meetings as of December 31, 2024. "
        "That voting differential supports strategic continuity even when LVMH has a large public float. The URD also states that, aside from Christian Dior SE, the Arnault family and companies owned by it directly or indirectly held 7.13% of share capital and 8.04% of exercisable voting rights [LVMH3].",
    )
    add_para(
        doc,
        "Succession should be treated as governance risk and continuity mechanism, not as gossip. Official evidence supports family-controlled governance; it does not fully disclose future family allocation, trusts, or private succession mechanics [LVMH3].",
    )
    add_confidence_note(doc, "High for disclosed control percentages; medium-low for future succession mechanics", "Control percentages are official. Future private arrangements remain unknown.", "[LVMH3]")

    doc.add_heading("11. Financial statement linkage", level=1)
    add_para(
        doc,
        f"LVMH's financial statements connect to Arnault wealth because public markets capitalize luxury earnings, cash generation, and control of scarce brands. In 2024, LVMH reported {eur_b(LVMH_FACTS['revenue_2024_b'])} revenue, {eur_b(LVMH_FACTS['gross_margin_2024_b'])} gross margin, {eur_b(LVMH_FACTS['profit_from_recurring_operations_2024_b'])} profit from recurring operations, {eur_b(LVMH_FACTS['group_share_net_profit_2024_b'])} group-share net profit, {eur_b(LVMH_FACTS['net_cash_from_operating_activities_2024_b'])} net cash from operating activities, and {eur_b(LVMH_FACTS['operating_free_cash_flow_2024_b'])} operating free cash flow. "
        f"Operating investments were {eur_b(LVMH_FACTS['operating_investments_2024_b'])}, which shows that the business is capital-intensive in stores, workshops, supply chain, hospitality, and brand experience but still cash-generative [LVMH1].",
    )
    add_kv_table(
        doc,
        [
            ("2024 revenue", eur_b(LVMH_FACTS["revenue_2024_b"])),
            ("2024 profit from recurring operations", eur_b(LVMH_FACTS["profit_from_recurring_operations_2024_b"])),
            ("2024 recurring operating margin", pct(group_margin)),
            ("2024 operating free cash flow", eur_b(LVMH_FACTS["operating_free_cash_flow_2024_b"])),
            ("2024 operating investments", eur_b(LVMH_FACTS["operating_investments_2024_b"])),
            ("2024 net financial debt", eur_b(LVMH_FACTS["net_financial_debt_2024_b"])),
            ("2024 equity", eur_b(LVMH_FACTS["equity_2024_b"])),
        ],
        font_size=8.2,
    )
    add_confidence_note(doc, "High", "Values are from LVMH 2024 URD financial highlights and key financial tables.", "[LVMH1][LVMH5]")

    doc.add_heading("12. LVMH public-market sensitivity", level=1)
    add_para(
        doc,
        f"Using the family-group share-capital exposure of {LVMH_FACTS['family_group_share_capital_pct_2024']:.2f}%, a EUR 100B change in LVMH market capitalization would imply about EUR {100 * family_exposure:,.1f}B of look-through change to family-group LVMH equity value before taxes, trusts, debt, liquidity, holding-company discounts, and intra-family allocation. "
        "This is an illustration, not a Forbes valuation model [LVMH3].",
    )
    if len(charts) > 1:
        add_picture_with_caption(doc, charts[1], "Illustrative LVMH market-cap sensitivity to family group value")
    add_confidence_note(doc, "Medium", "The ownership percentage is official; the linear sensitivity is a derived illustration and not a complete personal net-worth model.", "[LVMH3]")

    doc.add_heading("13. Forbes wealth history, CAGR, and exponential-style fit", level=1)
    add_para(
        doc,
        f"The Forbes annual series has {int(metric.get('years_observed'))} observations from {int(metric.get('first_year_observed'))} to 2025. Arnault moved from {money(metric.get('first_net_worth_usd_b'))} to {money(person.get('net_worth_2025_usd_b'))}, a {float(metric.get('wealth_multiple_first_to_2025')):,.1f}x multiple and nominal endpoint CAGR of {pct(metric.get('CAGR_nominal'))}. "
        f"The log-linear fit has slope {float(metric.get('log_linear_growth_slope')):,.4f}, R^2 {float(metric.get('exponential_fit_r2')):,.3f}, and estimated doubling time of {float(metric.get('estimated_doubling_time_years')):,.2f} years. The fit is strong as a descriptive annual-list model, but it is not proof of true exponential growth [D1].",
    )
    add_para(
        doc,
        f"The cycle risk is visible. The processed metrics show peak net worth of {money(metric.get('peak_net_worth_usd_b'))}, max one-year gain of {money(metric.get('max_one_year_gain_usd_b'))}, max one-year loss of {money(metric.get('max_one_year_loss_usd_b'))}, annual growth volatility of {pct(metric.get('volatility_of_annual_growth'))}, and largest drawdown of {pct(metric.get('largest_drawdown_pct'))}. "
        "The 2024-to-2025 drop in the Forbes series should be read as equity-market and luxury-cycle sensitivity, not audited realized cash loss [D1].",
    )
    add_confidence_note(doc, "High for metrics; medium for cycle interpretation", "Metrics are computed from annual Forbes observations; cause attribution needs share-price/event data for a full event study.", "[D1][LVMH1]")

    doc.add_heading("14. First-principles luxury compounding analysis", level=1)
    add_para(
        doc,
        "The scarce resource Arnault controls is organized desirability: historic brands, creative direction, craftsmanship, selective distribution, premium retail real estate, supply-chain know-how, and family governance that can invest through cycles. "
        "Luxury differs from software platforms because scarcity, brand memory, cultural status, and physical retail experience are part of the product. It differs from ordinary retail because the goal is not maximum unit volume; the goal is controlled demand at premium economics [LVMH4][LVMH8].",
    )
    add_para(
        doc,
        "The constraint LVMH attacks is fragmentation. Luxury brands can have strong heritage but limited capital, distribution, and global operating discipline. LVMH aggregates brands into a portfolio that can professionalize retail, expand internationally, invest in artisanship and supply, and preserve brand distinction [LVMH1][LVMH2][LVMH4].",
    )
    add_confidence_note(doc, "Medium-high", "The thesis is synthesis from official portfolio, segment, distribution, and control evidence.", "[LVMH1][LVMH2][LVMH4][LVMH8]")

    doc.add_heading("15. Capital allocation and portfolio compounding", level=1)
    add_para(
        doc,
        "The Arnault pattern is portfolio compounding: use control of a listed luxury group to acquire, own, professionalize, and scale scarce brands while preserving Maison-level identity. "
        "The 2024 URD supports the breadth of the Maison portfolio and business groups, but this report does not assign standalone valuations to Louis Vuitton, Dior, Tiffany, Sephora, or other Maisons because official filings do not provide complete brand-level valuation tables [LVMH2][LVMH8].",
    )
    add_bullets(
        doc,
        [
            "Cash-flow engine: the disclosed business groups produce revenue and recurring operating profit.",
            "Option-value engine: portfolio acquisitions, hospitality, media, and brand extensions may add value but require evidence before separate valuation.",
            "Governance engine: family control supports long time horizons but creates succession and key-person governance risk.",
        ],
    )
    add_confidence_note(doc, "Medium-high", "Portfolio and financial facts are official; individual acquisition economics and brand-level valuations need additional research.", "[LVMH1][LVMH2][LVMH3][LVMH8]")

    doc.add_heading("16. Industry structure and geographic context", level=1)
    add_para(
        doc,
        "LVMH's luxury exposure is global. The 2024 URD revenue mix by region includes Asia excluding Japan, the United States, Europe excluding France, Japan, France, and other markets. "
        "That geography matters because a luxury slowdown, currency move, tourist-flow change, or policy shock in a major region can affect revenue growth and valuation even if brand equity remains intact [LVMH1].",
    )
    add_matrix_table(doc, ["Region", "2024 revenue share"], [[region, f"{share}%"] for region, share in LVMH_FACTS["revenue_region_pct_2024"].items()], font_size=8.2)
    add_para(
        doc,
        "The Fashion and Leather Goods market is fragmented at the brand level but global scale favors groups that can combine creative talent, retail execution, supplier relationships, capital, and brand stewardship across markets [LVMH4].",
    )
    add_confidence_note(doc, "High for geography; medium for industry-structure synthesis", "Regional mix and competition/distribution evidence are official. Competitive position is an analytical synthesis.", "[LVMH1][LVMH4]")

    doc.add_heading("17. Risks and counter-thesis", level=1)
    add_bullets(
        doc,
        [
            "Brand dilution: if growth targets push volume too far, scarcity and desirability can weaken [LVMH4].",
            "Luxury-cycle sensitivity: Forbes wealth can fall when public markets mark down LVMH on slower demand, margin pressure, currency moves, or multiple compression [D1][LVMH1].",
            "Creative and talent risk: luxury houses rely on designers, artisans, executives, and retail execution to renew heritage without commoditizing it [LVMH4].",
            "Raw-material and supply risk: LVMH discloses exposure to rare materials, climate impacts, raw-material pricing, energy, labor, traceability, and certification risk [LVMH1].",
            "Governance and succession risk: family control supports continuity but concentrates control and leaves private succession arrangements incompletely visible [LVMH3].",
            "Personal balance-sheet opacity: taxes, trusts, debt, personal liquidity, dividends retained, and intra-family allocation are not fully disclosed [F1][LVMH3].",
        ],
    )
    add_confidence_note(doc, "High for disclosed risk categories; medium for wealth impact", "Risk categories are official or project-derived; the size and timing of wealth effects remain uncertain.", "[D1][LVMH1][LVMH3][LVMH4]")

    doc.add_heading("18. Comparable billionaire patterns", level=1)
    add_para(
        doc,
        "Arnault is the project's fourth archetype template: luxury/retail brand ownership plus family-controlled holding-company compounding. The closest pattern is not a software platform, commodity cycle, or single-founder private valuation. "
        "It is a controlled listed group that converts brand scarcity, distribution control, creative renewal, acquisitions, and margin durability into public-equity value [D1][LVMH1][LVMH3].",
    )
    add_confidence_note(doc, "Medium-high", "Comparable pattern is a project taxonomy synthesis anchored in LVMH financial and ownership evidence.", "[D1][LVMH1][LVMH3][LVMH4]")

    doc.add_heading("19. Transferable business lessons", level=1)
    add_numbered(
        doc,
        [
            "Scarcity can be an operating system when brand image, distribution, supply, price, and creative renewal are controlled together.",
            "A portfolio of premium brands can compound more durably than a single product line when each Maison retains identity but receives group-level capital and operating discipline.",
            "Family control can support long-horizon investment, but it must be assessed with succession and governance risk in mind.",
            "High margins in luxury are earned by avoiding commoditization, not by maximizing volume at all costs.",
            "Public-equity wealth can grow from intangible assets when filings show repeatable earnings, cash flow, and control of distribution.",
            "The best acquisition platforms preserve the acquired asset's scarcity rather than forcing one generic operating model across the portfolio.",
        ],
    )
    add_confidence_note(doc, "Medium-high", "Lessons derive from LVMH's disclosed business groups, segment economics, distribution strategy, and ownership/control structure.", "[LVMH1][LVMH2][LVMH3][LVMH4]")

    doc.add_heading("20. Bernard Arnault evidence gaps and confidence notes", level=1)
    add_matrix_table(
        doc,
        ["Report area", "Confidence", "Remaining gap"],
        [
            ["Forbes rank and net worth", "High", "Forbes estimate is canonical but not an audited balance sheet."],
            ["LVMH family group capital/voting control", "High", "Exact personal/intra-family allocation and private estate structures are not fully visible."],
            ["LVMH segment economics", "High", "Brand-level profitability and standalone Maison valuations are not disclosed."],
            ["Market-cap sensitivity", "Medium", "Derived from family group share-capital exposure; not a full Forbes model."],
            ["Private assets / holding-company details", "Low-medium", "Use only when official or high-quality evidence is added."],
            ["Succession and governance outlook", "Medium-low", "Official board/control facts exist; future private succession mechanics remain unknown."],
        ],
        font_size=8.0,
    )
    add_confidence_note(doc, "High", "The report separates official evidence from inference and missing private balance-sheet evidence.", "[F1][D1][LVMH1][LVMH3]")


def add_ellison_specific_sections(doc: Document, context: EnrichedContext, charts: list[Path]) -> None:
    """Add deep enterprise software/database/cloud sections for Larry Ellison."""
    person = context.person
    metric = context.metrics
    ownership_pct = oracle_ellison_ownership_pct()
    total_margin = ORACLE_FACTS["total_operating_margin_2024_b"] / ORACLE_FACTS["total_revenue_2024_b"]
    cloud_license_share = ORACLE_FACTS["cloud_and_license_revenue_2024_b"] / ORACLE_FACTS["total_revenue_2024_b"]
    cloud_services_support_share = ORACLE_FACTS["cloud_services_license_support_total_2024_b"] / ORACLE_FACTS["total_revenue_2024_b"]

    doc.add_heading("6. Ellison wealth equation: Oracle public equity first", level=1)
    add_para(
        doc,
        f"The simplified wealth equation is: Forbes net worth for {person['name']} of {money(person.get('net_worth_2025_usd_b'))} is approximately Ellison's economic exposure to Oracle public equity multiplied by Oracle's public market value, plus or minus taxes, debt, pledged-share terms, liquidity, trusts, private assets, and other holdings. "
        f"Oracle's 2024 proxy reports Ellison beneficial ownership of {ORACLE_FACTS['ellison_beneficial_shares_2024']:,} shares, or approximately {ORACLE_FACTS['ellison_beneficial_ownership_pct_2024']:.1f}% of outstanding Oracle common stock as of September 16, 2024. "
        "That proxy ownership percentage is the visible core of the Forbes wealth bridge, not a complete personal balance sheet [F1][ORCL4].",
    )
    add_kv_table(
        doc,
        [
            ("Forbes 2025 estimate", money(person.get("net_worth_2025_usd_b"))),
            ("Visible core asset", "Oracle Corporation public equity"),
            ("Proxy beneficial shares", f"{ORACLE_FACTS['ellison_beneficial_shares_2024']:,}"),
            ("Proxy beneficial ownership", f"{ORACLE_FACTS['ellison_beneficial_ownership_pct_2024']:.1f}%"),
            ("Currently exercisable options included", f"{ORACLE_FACTS['ellison_exercisable_options_included_2024']:,}"),
            ("Pledged shares disclosed", f"{ORACLE_FACTS['ellison_pledged_shares_2024']:,}"),
            ("Role evidence", "Founder, Chairman, CTO, director since 1977, and largest stockholder"),
            ("Not fully visible", "Taxes, debt balances/terms, trusts, liquidity discounts, private assets, family entities, and non-Oracle holdings."),
        ],
        font_size=8.0,
    )
    add_confidence_note(doc, "High for Oracle proxy ownership; medium for personal wealth bridge", "The proxy supports ownership, role, and pledging disclosures. Forbes supplies the estimated net worth. Private balance-sheet details remain incomplete.", "[F1][ORCL4][ORCL5]")

    doc.add_heading("7. Oracle business empire map", level=1)
    add_para(
        doc,
        "The Ellison wealth engine is enterprise software lock-in, not consumer attention, luxury scarcity, or retail logistics. Oracle sits inside mission-critical database, cloud, application, middleware, Java, hardware, and services ecosystems. "
        "The report treats private assets and outside investments conservatively: they appear only where the Oracle proxy discloses related-party transactions or where independent source rows exist; they are not used as a precise Forbes wealth bridge [ORCL2][ORCL3][ORCL4][ORCL5].",
    )
    if len(charts) > 3:
        add_picture_with_caption(doc, charts[3], "Oracle enterprise software, cloud, and founder public-equity wealth-engine map")
    add_confidence_note(doc, "High for Oracle business categories; medium for product-level valuation", "Oracle filing evidence supports business categories, but not standalone valuations for each product ecosystem.", "[ORCL2][ORCL3]")

    doc.add_heading("8. Revenue model and financial statement linkage", level=1)
    add_para(
        doc,
        f"Oracle's fiscal 2024 10-K shows a company built around recurring enterprise software economics. Total revenue was {money_precise_b(ORACLE_FACTS['total_revenue_2024_b'])}, total operating margin was {money_precise_b(ORACLE_FACTS['total_operating_margin_2024_b'])}, and total operating margin percentage was {pct(total_margin)}. "
        f"Cloud and license revenue was {money_precise_b(ORACLE_FACTS['cloud_and_license_revenue_2024_b'])}, or about {pct(cloud_license_share)} of total revenue. Cloud services and license support alone totaled {money_precise_b(ORACLE_FACTS['cloud_services_license_support_total_2024_b'])}, about {pct(cloud_services_support_share)} of total revenue [ORCL1][ORCL2].",
    )
    add_matrix_table(
        doc,
        ["Revenue category", "FY2024 revenue", "Business interpretation"],
        [
            ["Cloud services", money_precise_b(ORACLE_FACTS["cloud_services_revenue_2024_b"]), "Subscription cloud services, OCI/SaaS/PaaS/IaaS demand, and recurring cloud consumption."],
            ["License support", money_precise_b(ORACLE_FACTS["license_support_revenue_2024_b"]), "Recurring support/maintenance economics from installed enterprise software base."],
            ["Cloud license and on-premise license", money_precise_b(ORACLE_FACTS["cloud_license_on_prem_revenue_2024_b"]), "New license sales and on-premise license activity; more episodic than support."],
            ["Hardware", money_precise_b(ORACLE_FACTS["hardware_revenue_2024_b"]), "Engineered systems, servers, storage, and hardware support around the Oracle stack."],
            ["Services", money_precise_b(ORACLE_FACTS["services_revenue_2024_b"]), "Consulting and advanced customer support around implementation and operations."],
        ],
        font_size=7.8,
    )
    if len(charts) > 2:
        add_picture_with_caption(doc, charts[2], "Oracle FY2024 revenue and business margin by major category")
    add_confidence_note(doc, "High", "Revenue and margin values are extracted from Oracle's fiscal 2024 Form 10-K tables.", "[ORCL1][ORCL2][ORCL7]")

    doc.add_heading("9. Cloud services, license support, and enterprise lock-in", level=1)
    add_para(
        doc,
        "Oracle's lock-in is not merely brand loyalty. Enterprise databases and applications sit inside transaction processing, ERP, finance, HR, healthcare, analytics, identity, security, compliance, and data workflows. "
        "Changing those systems can require data migration, application rewrites, retraining, audit/compliance work, downtime risk, and new vendor integration. That switching-cost structure supports recurring license-support and cloud-services economics [ORCL2][ORCL3].",
    )
    add_para(
        doc,
        f"The key financial clue is the mix: cloud services and license support totaled {money_precise_b(ORACLE_FACTS['cloud_services_license_support_total_2024_b'])}, split between applications cloud services/license support at {money_precise_b(ORACLE_FACTS['applications_cloud_support_revenue_2024_b'])} and infrastructure cloud services/license support at {money_precise_b(ORACLE_FACTS['infrastructure_cloud_support_revenue_2024_b'])}. "
        "This is the installed-base monetization layer that makes Oracle different from a pure new-license software vendor [ORCL2].",
    )
    add_confidence_note(doc, "High for revenue mix; medium for switching-cost interpretation", "Filing data supports the revenue mix and renewal importance. Switching-cost analysis is first-principles synthesis from enterprise software economics.", "[ORCL2][ORCL3]")

    doc.add_heading("10. OCI, cloud transition, and AI infrastructure demand", level=1)
    add_para(
        doc,
        "Oracle's cloud transition changes the growth narrative. The legacy model was database/license/support economics; the emerging market narrative is database plus cloud infrastructure, SaaS applications, data-center capacity, and AI infrastructure demand. "
        "The 10-K risk factors explicitly discuss data-center capacity and AI product risks, including the need to build and support AI products and the risk that competitors' AI products achieve higher market acceptance [ORCL3].",
    )
    add_para(
        doc,
        "This report does not invent an OCI valuation or AI revenue split. The evidence supports a cloud/AI infrastructure thesis because Oracle reports infrastructure cloud services and license support, discusses data-center capacity risk, and frames AI as a significant product and investment area. Precise AI-specific revenue, margin, and customer concentration would require additional Oracle disclosures [ORCL2][ORCL3].",
    )
    add_confidence_note(doc, "Medium-high", "Cloud infrastructure revenue categories are official; AI-specific valuation attribution remains incomplete.", "[ORCL2][ORCL3]")

    doc.add_heading("11. Cash flow, capex, buybacks, and dividends", level=1)
    add_para(
        doc,
        f"Oracle's public-equity wealth engine also depends on cash generation and capital allocation. Fiscal 2024 operating cash flow was {money_precise_b(ORACLE_FACTS['operating_cash_flow_2024_b'])}, capital expenditures were {money_precise_b(ORACLE_FACTS['capital_expenditures_2024_b'])}, and free cash flow was {money_precise_b(ORACLE_FACTS['free_cash_flow_2024_b'])}. "
        f"Oracle also reported {money_precise_b(ORACLE_FACTS['common_stock_repurchases_2024_b'])} of payments for common-stock repurchases and {money_precise_b(ORACLE_FACTS['dividends_paid_2024_b'])} of dividends paid to stockholders [ORCL1].",
    )
    add_para(
        doc,
        "For Ellison, capital allocation compounds through two channels. First, cash flow supports Oracle's investment, debt service, dividends, and buybacks. Second, because he owns a very large percentage of Oracle, buybacks and valuation rerating can mechanically increase the value of his retained ownership, while dividends provide cash flow. The report does not claim how much dividend cash Ellison personally retained or how it was taxed [ORCL1][ORCL4].",
    )
    add_confidence_note(doc, "High for cash-flow and capital-return values; medium for personal compounding inference", "Financial values are official; personal tax/use of proceeds is unknown.", "[ORCL1][ORCL4]")

    doc.add_heading("12. Oracle market-cap sensitivity to Ellison wealth", level=1)
    add_para(
        doc,
        f"With {ORACLE_FACTS['ellison_beneficial_ownership_pct_2024']:.1f}% proxy beneficial ownership, a $100B change in Oracle market capitalization would mechanically imply about ${100 * ownership_pct:,.1f}B of change in Ellison's Oracle-linked beneficial-ownership value before taxes, debt, pledge terms, liquidity, trusts, and private assets. "
        "That makes Ellison's Forbes wealth especially sensitive to Oracle public-equity rerating around cloud, AI infrastructure, support durability, and capital allocation [ORCL4].",
    )
    if len(charts) > 1:
        add_picture_with_caption(doc, charts[1], "Illustrative Oracle market-cap sensitivity to Ellison Oracle-linked wealth")
    add_confidence_note(doc, "Medium", "The ownership percentage is sourced; the market-cap sensitivity is a derived illustration, not a sourced Forbes valuation model.", "[ORCL4]")

    doc.add_heading("13. Pledged shares, governance, and control limits", level=1)
    add_para(
        doc,
        f"The proxy discloses that as of September 16, 2024, Ellison had pledged {ORACLE_FACTS['ellison_pledged_shares_2024']:,} Oracle shares as collateral for certain personal indebtedness. It also states the pledged shares secured personal term loans, were not pledged for margin accounts, and were subject to Governance Committee review. "
        "This is important because pledged shares create a risk factor for public shareholders and for the personal wealth bridge, but the filing does not disclose all loan terms, balances, covenants, trusts, taxes, or liquidity arrangements [ORCL5].",
    )
    add_para(
        doc,
        "Ellison's influence is operating and ownership-based rather than dual-class legal control. The proxy identifies him as founder, chairman, CTO, director since 1977, and largest stockholder. It does not describe a Meta-style dual-class vote structure or an LVMH-style family-control chain [ORCL4][ORCL5].",
    )
    add_confidence_note(doc, "High for pledged-share disclosure; medium-low for full personal debt picture", "Pledged-share count and governance review are official. Full personal loan details are not public.", "[ORCL4][ORCL5]")

    doc.add_heading("14. Forbes wealth history, CAGR, and exponential-style fit", level=1)
    add_para(
        doc,
        f"The Forbes annual series has {int(metric.get('years_observed'))} observations from {int(metric.get('first_year_observed'))} to 2025. Ellison moved from {money(metric.get('first_net_worth_usd_b'))} to {money(person.get('net_worth_2025_usd_b'))}, a {float(metric.get('wealth_multiple_first_to_2025')):,.1f}x multiple and nominal endpoint CAGR of {pct(metric.get('CAGR_nominal'))}. "
        f"The log-linear fit has slope {float(metric.get('log_linear_growth_slope')):,.4f}, R^2 {float(metric.get('exponential_fit_r2')):,.3f}, and estimated doubling time of {float(metric.get('estimated_doubling_time_years')):,.2f} years. This is a strong descriptive fit to annual Forbes observations, not proof of true exponential growth or a forecast [D1].",
    )
    add_para(
        doc,
        f"The processed metrics show peak net worth of {money(metric.get('peak_net_worth_usd_b'))}, max one-year gain of {money(metric.get('max_one_year_gain_usd_b'))}, max one-year loss of {money(metric.get('max_one_year_loss_usd_b'))}, annual growth volatility of {pct(metric.get('volatility_of_annual_growth'))}, and largest drawdown of {pct(metric.get('largest_drawdown_pct'))}. "
        "The 2025 Forbes value is the series peak, consistent with public-market rerating sensitivity rather than a smooth operating-cash-flow path [D1][ORCL4].",
    )
    add_confidence_note(doc, "High for metrics; medium for market interpretation", "Metrics are computed from annual Forbes observations; causal attribution requires Oracle share-price and event evidence.", "[D1][ORCL4]")

    doc.add_heading("15. First-principles enterprise software analysis", level=1)
    add_para(
        doc,
        "The scarce resource Ellison helped control is a mission-critical data layer inside enterprises. Databases and enterprise applications become valuable when they are reliable, deeply integrated, supported for years, and expensive to replace. "
        "Unlike Meta's social graph, Oracle's network is not consumer attention. Unlike Amazon, the moat is less logistics and commerce frequency. Unlike LVMH, the value is not luxury scarcity. Oracle's moat is enterprise dependency: data gravity, integration cost, renewal economics, uptime risk, and long procurement cycles [ORCL2][ORCL3].",
    )
    add_para(
        doc,
        "This explains why Oracle can generate durable support cash flow even when technology delivery models change. The cloud transition is a threat if customers migrate away, but it is also an opportunity if Oracle converts the installed base into OCI, cloud applications, database services, and AI infrastructure demand [ORCL2][ORCL3].",
    )
    add_confidence_note(doc, "Medium-high", "First-principles claims synthesize official revenue mix, renewal/risk language, and cloud/AI risk disclosures.", "[ORCL2][ORCL3]")

    doc.add_heading("16. Country and macro/regulatory context", level=1)
    add_para(
        doc,
        f"Forbes lists Ellison's country/territory as {clean(person.get('country_or_territory'))}, while Oracle is a U.S.-listed enterprise software and cloud infrastructure company with global revenue exposure. "
        f"Oracle's fiscal 2024 geographic revenue was {money_precise_b(ORACLE_FACTS['americas_revenue_2024_b'])} in the Americas, {money_precise_b(ORACLE_FACTS['emea_revenue_2024_b'])} in Europe/Middle East/Africa, and {money_precise_b(ORACLE_FACTS['asia_pacific_revenue_2024_b'])} in Asia Pacific. "
        "That makes Ellison's public-equity wealth exposed to U.S. capital markets and global enterprise IT regulation, including data sovereignty, cybersecurity, AI, cloud procurement, competition, and customer data-center requirements [F1][ORCL2][ORCL3][ORCL6].",
    )
    add_kv_table(
        doc,
        [
            ("Forbes country/territory", clean(person.get("country_or_territory"))),
            ("Company listing / disclosure anchor", "Oracle Corporation SEC filings and Oracle Investor Relations"),
            ("Americas FY2024 revenue", money_precise_b(ORACLE_FACTS["americas_revenue_2024_b"])),
            ("EMEA FY2024 revenue", money_precise_b(ORACLE_FACTS["emea_revenue_2024_b"])),
            ("Asia Pacific FY2024 revenue", money_precise_b(ORACLE_FACTS["asia_pacific_revenue_2024_b"])),
            ("Regulatory exposure", "Cloud/data-center capacity, cybersecurity, privacy/data sovereignty, competition, AI product risk, tax and cross-border enterprise procurement."),
        ],
        font_size=8.0,
    )
    add_confidence_note(doc, "High for geography and company-disclosure anchor; medium for macro interpretation", "Geographic revenue and risk-factor categories are official; macro/regulatory impact on Ellison's personal wealth is interpreted through Oracle public-equity exposure.", "[F1][ORCL2][ORCL3][ORCL6]")

    doc.add_heading("17. Risks and counter-thesis", level=1)
    add_bullets(
        doc,
        [
            "Cloud transition risk: customers may shift workloads to competing cloud platforms if Oracle cannot deliver differentiated cloud economics and service levels [ORCL3].",
            "AI risk: Oracle may fail to recoup AI investments or competitors may achieve higher market acceptance [ORCL3].",
            "Data-center capacity risk: Oracle may be unable to increase existing capacity or establish new data centers fast enough to meet demand [ORCL3].",
            "Renewal risk: renewals of license support, hardware support, and cloud subscriptions are important to future success [ORCL3].",
            "Pledged-share risk: Ellison's pledged shares are disclosed and monitored, but loan balances and terms are not fully public [ORCL5].",
            "Personal balance-sheet opacity: taxes, trusts, debt terms, private assets, liquidity discounts, and non-Oracle holdings are not fully disclosed [F1][ORCL4][ORCL5].",
        ],
    )
    add_confidence_note(doc, "High for public-company risks; medium for personal wealth impact", "Risk categories are official; magnitude and timing are uncertain.", "[ORCL3][ORCL5]")

    doc.add_heading("18. Peer comparison and comparable billionaire patterns", level=1)
    add_para(
        doc,
        "Ellison is the fifth archetype template in the project: enterprise software/database lock-in plus founder public-equity wealth. He differs from Zuckerberg because Oracle's lock-in is enterprise workflow and data dependency rather than consumer social network effects. "
        "He differs from Bezos because Oracle's physical infrastructure is mainly data centers and enterprise systems, not consumer logistics. He differs from Arnault because the moat is operational switching cost and support renewal, not luxury brand scarcity [D1][ORCL2][ORCL3].",
    )
    add_confidence_note(doc, "Medium-high", "Comparable pattern is a project taxonomy synthesis anchored in Oracle filings and Forbes wealth metrics.", "[D1][ORCL2][ORCL4]")

    doc.add_heading("19. Transferable business lessons", level=1)
    add_numbered(
        doc,
        [
            "Mission-critical software can compound for decades when replacement risk is high and support economics are recurring.",
            "A founder can create enormous wealth without dual-class control when retained public-equity ownership is very large.",
            "Installed-base economics can fund transition into cloud and AI, but the transition must protect renewal economics.",
            "Database lock-in is not a slogan; it comes from data gravity, uptime risk, compliance, integrations, and retraining cost.",
            "Buybacks and dividends matter more when a founder retains a high ownership percentage.",
            "Pledged shares are part of the risk analysis; they should be disclosed separately from economic ownership.",
        ],
    )
    add_confidence_note(doc, "Medium-high", "Lessons derive from Oracle ownership, financial, risk, and revenue-mix evidence.", "[ORCL1][ORCL2][ORCL4][ORCL5]")

    doc.add_heading("20. Larry Ellison evidence gaps and confidence notes", level=1)
    add_matrix_table(
        doc,
        ["Report area", "Confidence", "Remaining gap"],
        [
            ["Forbes rank and net worth", "High", "Forbes estimate is canonical but not an audited balance sheet."],
            ["Oracle beneficial ownership", "High", "Proxy ownership is official; personal tax and trust structures are incomplete."],
            ["Pledged shares", "High for count; medium-low for debt terms", "Proxy discloses 277M pledged shares but not complete personal loan economics."],
            ["Oracle revenue/cash flow", "High", "FY2024 financial values are from Oracle's Form 10-K."],
            ["OCI/AI valuation attribution", "Medium-low", "Cloud/AI evidence exists, but precise AI revenue/margin/valuation attribution is not fully disclosed."],
            ["Private assets and outside investments", "Low-medium", "Only include if supported by proxy related-party rows or future primary evidence."],
        ],
        font_size=8.0,
    )
    add_confidence_note(doc, "High", "The report separates official filing evidence from derived sensitivity and missing private-balance-sheet evidence.", "[F1][D1][ORCL1][ORCL4][ORCL5]")


def add_generic_future_sections(doc: Document, context: EnrichedContext, evidence_keys: str) -> None:
    """Add reusable sections for non-Elon reports or evidence-pack-driven drafts."""
    person = context.person
    metric = context.metrics
    entities = context.evidence_pack["entity_or_asset"].dropna().astype(str).str.strip()
    entities = [item for item in dict.fromkeys(entities.tolist()) if item]
    if not entities:
        entities = [clean(person.get("primary_company_or_asset")) or clean(metric.get("key_asset_or_company")) or "Core asset"]

    doc.add_heading("6. Core operating companies/assets", level=1)
    add_para(
        doc,
        f"The current structured dataset identifies {clean(person.get('primary_company_or_asset')) or 'the primary asset'} as the key asset/company and {clean(person.get('source_of_wealth'))} as the source of wealth. "
        "Person-specific evidence-pack rows should expand this into ownership, operating economics, valuation bridge, and risk factors.",
    )
    add_bullets(doc, [f"Evidence target: {entity}" for entity in entities[:8]])
    add_confidence_note(doc, "Medium-low", "This section is reusable but should be upgraded with person-specific filings or evidence-pack rows.", evidence_keys)

    pack_rows = evidence_pack_rows_for_body(context.evidence_pack)
    if pack_rows:
        doc.add_heading("6A. Person-specific evidence collected for this draft", level=2)
        add_para(
            doc,
            "These rows are structured evidence inputs for the enriched draft. They anchor the report in primary or high-quality sources while keeping unsupported ownership, valuation, revenue, debt, trust, and liquidity claims out of the narrative.",
        )
        add_matrix_table(
            doc,
            ["Key", "Entity / asset", "Evidence type", "Claim supported", "Confidence"],
            pack_rows,
            font_size=7.6,
        )
        add_confidence_note(
            doc,
            "Medium-high where sources are primary; lower where only a locator exists",
            "Evidence-pack claims are included as supportable draft anchors, not as a complete personal balance-sheet bridge.",
            "".join(row[0] for row in pack_rows[:8]),
        )

    doc.add_heading("7. Financial statement and market value linkage", level=1)
    add_para(
        doc,
        "Financial statement linkage is intentionally conservative until primary filings, annual reports, prospectuses, or audited statements are supplied. "
        "For public-equity wealth engines, connect revenue, margins, cash flow, market cap, and ownership. For private companies, separate verified operating facts from private valuation marks.",
    )
    if pack_rows:
        filing_rows = [
            row
            for row in pack_rows
            if any(token in row[2].casefold() or token in row[3].casefold() for token in ["filing", "annual", "proxy", "financial", "investor", "valuation"])
        ]
        if filing_rows:
            add_para(
                doc,
                "The evidence pack contains filing or investor-source rows that should be used to upgrade this section from draft to person-specific analysis. Exact numbers should be copied only after reviewing the underlying source table and report date.",
            )
            add_matrix_table(
                doc,
                ["Key", "Entity / asset", "Evidence type", "Claim supported", "Confidence"],
                filing_rows[:6],
                font_size=7.6,
            )
    add_confidence_note(doc, "Low until evidence pack is filled", "No unsupported revenue, profit, or valuation numbers are generated by default.", evidence_keys)

    doc.add_heading("8. Industry structure and competitive position", level=1)
    add_para(
        doc,
        f"The project industry label is {clean(person.get('industry'))}. Final analysis should map profit pools, market structure, regulation, suppliers/customers, and scarcity sources using primary or high-quality secondary evidence.",
    )
    add_confidence_note(doc, "Medium-low", "Industry label is sourced, but industry mechanics need person-specific sources.", "[F1]")

    doc.add_heading("9. Moat, flywheel, and first-principles analysis", level=1)
    add_para(
        doc,
        f"Working evidence summary from the metrics table: {clean(metric.get('evidence_summary'))}. This is a classification clue, not a final moat conclusion. "
        "The reusable template asks: what scarce resource is controlled, what constraint was attacked, why customers or capital markets rewarded the result, and what could break the engine?",
    )
    add_confidence_note(doc, "Medium-low", "Moat analysis is a draft until backed by filings, economics, and competitive evidence.", "[D1]")

    doc.add_heading("10. Capital allocation pattern", level=1)
    add_para(doc, "Add evidence for reinvestment, acquisitions, buybacks, dividends, leverage, asset sales, family control, fund performance, or holding-company compounding. Unsupported claims remain excluded.")
    add_confidence_note(doc, "Low until evidence pack is filled", "The generic generator avoids inventing capital allocation history.", evidence_keys)

    doc.add_heading("11. Country and macro/regulatory context", level=1)
    add_para(
        doc,
        f"Country/territory in the Forbes record: {clean(person.get('country_or_territory'))}. Final reports should add tax, currency, industrial policy, concession, exchange-listing, sanction, or succession context where it materially affects the wealth engine.",
    )
    add_confidence_note(doc, "Medium-low", "Country field is sourced; macro/regulatory interpretation requires added sources.", "[F1]")

    doc.add_heading("12. Risks and counter-thesis", level=1)
    add_para(
        doc,
        f"Observed annual wealth volatility: {pct(metric.get('volatility_of_annual_growth'))}; largest Forbes-history drawdown: {pct(metric.get('largest_drawdown_pct'))}. "
        "Person-specific risks should be supported by filings, sector data, debt disclosures, regulatory history, or credible reporting.",
    )
    add_confidence_note(doc, "Medium", "Quantitative risk metrics are project-derived; qualitative risks require evidence-pack support.", "[D1]")

    doc.add_heading("13. Strategic patterns and transferable lessons", level=1)
    add_para(
        doc,
        "Generic lesson draft: durable billionaire wealth usually reflects concentrated ownership of scarce assets, favorable industry structure, capital allocation, and timing. "
        "The final version should translate the specific person's evidence into repeatable lessons without copying superficial biography.",
    )
    add_confidence_note(doc, "Medium-low", "Reusable synthesis requires person-specific upgrade before research-grade publication.", evidence_keys)


def add_appendices(doc: Document, evidence: list[EvidenceItem]) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Evidence appendix", level=1)
    add_para(
        doc,
        "Body text uses short keys. The overview table below is intentionally compact; detailed source cards provide the full locator, access date, evidence note, confidence, and limitation.",
    )
    add_confidence_note(doc, "High", "Appendix rows are generated from project data, curated evidence, and optional evidence-pack rows.", "[F1][D1]")
    add_matrix_table(
        doc,
        ["Key", "Category", "Source", "Publisher / date", "Confidence"],
        [[item.key, item.category, item.title, f"{item.publisher}; {item.date}", item.confidence] for item in evidence],
        font_size=8,
    )
    for category in EVIDENCE_CATEGORIES:
        rows = [item for item in evidence if item.category == category]
        if not rows:
            continue
        doc.add_heading(category, level=2)
        for item in rows:
            add_kv_table(
                doc,
                [
                    ("Key", item.key),
                    ("Source", item.title),
                    ("Publisher / date", f"{item.publisher}; {item.date}"),
                    ("Supports", item.supports),
                    ("Quality", item.reliability),
                    ("Full locator / URL", item.locator),
                    ("Accessed at", item.accessed_at),
                    ("Confidence", item.confidence),
                    ("Limitation", item.limitation),
                ],
                font_size=7.8,
            )

    doc.add_heading("Evidence gaps", level=1)
    add_bullets(
        doc,
        [
            "Exact ownership stakes, voting control, trusts, debt, and liquidity discounts where not publicly disclosed.",
            "Audited private-company revenue, profit, cash flow, debt, and capex for private assets.",
            "Private-company valuation share classes, liquidation preferences, and tender mechanics.",
            "Segment financials and unit economics where companies do not disclose enough detail.",
            "Regulatory context, country risk, litigation, and succession materials where not yet collected.",
        ],
    )
    add_confidence_note(doc, "High", "These are explicit limitations, not hidden assumptions.", "[F1][D1]")

    doc.add_heading("Confidence level by section", level=1)
    add_matrix_table(
        doc,
        ["Section type", "Default confidence", "Reason"],
        [
            ["Forbes rank/net worth", "High", "Canonical Forbes annual-list/project source."],
            ["Forbes annual wealth history", "High", "Derived from processed annual Forbes snapshots."],
            ["Public-company financial linkage", "High when based on filings", "Requires annual reports, 10-K/20-F filings, or proxy statements."],
            ["Private-company valuation", "Medium to low", "Tender/funding marks are episodic and usually unaudited."],
            ["Strategy synthesis", "Medium", "Useful but interpretive; must cite the facts it relies on."],
            ["Option-value ventures", "Low to medium", "Large upside but often incomplete ownership, revenue, and regulatory evidence."],
        ],
        font_size=8,
    )
    add_confidence_note(doc, "High", "Confidence framework is part of the reusable report template.", "[F1][D1]")

    doc.add_heading("Claims not final without more evidence", level=1)
    add_bullets(
        doc,
        [
            "Precise sum-of-the-parts personal net worth bridge.",
            "Specific private-company wealth attribution without sourced stake and valuation data.",
            "Revenue, margin, or cash-flow claims without filings, audited statements, or credible primary evidence.",
            "Regulatory, litigation, or country-risk conclusions without official or high-quality sources.",
            "Strategic lessons that are not traceable to specific evidence rows.",
        ],
    )
    add_confidence_note(doc, "High", "This section prevents unsupported claims from being promoted to final conclusions.", "[F1][D1]")

    doc.add_heading("Data limitations", level=1)
    add_para(
        doc,
        "This report is a research memo, not an audited wealth statement or investment recommendation. Forbes values are estimates; annual snapshots miss intra-year volatility; private-company values are opaque; and personal debt, taxes, trusts, family entities, and liquidity discounts are rarely fully visible."
    )
    add_confidence_note(doc, "High", "Limitations are explicit and should remain in every enriched report.", "[F1][D1]")


def create_enriched_report(
    context: EnrichedContext,
    *,
    output_dir: Path | None = None,
    charts_dir: Path | None = None,
    variant: str = "enriched_draft",
    overwrite: bool = False,
    year: int = DEFAULT_TARGET_YEAR,
) -> Path:
    """Create one enriched DOCX report from reusable structured inputs."""
    config = get_year_config(year)
    ensure_year_dirs(config)
    output_dir = output_dir or config.people_reports_dir
    charts_dir = charts_dir or config.charts_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    evidence = build_evidence_items(context)
    evidence_keys = "[" + "][".join(item.key for item in evidence[:4]) + "]"
    charts = create_charts(context, evidence, charts_dir, variant)
    output_path = output_dir / output_filename(context.person, variant, config)
    if output_path.exists() and not overwrite:
        raise FileExistsError(f"Refusing to overwrite existing report without --overwrite: {output_path}")
    if output_path.exists() and overwrite:
        try:
            with output_path.open("a+b"):
                pass
        except OSError as exc:
            raise RuntimeError(f"Target DOCX appears open or locked: {output_path}") from exc

    doc = Document()
    style_document(doc)
    add_title_and_reader_guide(doc, context, variant)
    add_generic_core_sections(doc, context, evidence_keys, charts)
    forbes_uri = clean(context.person.get("forbes_uri"))
    if config.legacy_layout and forbes_uri == "jeff-bezos":
        add_bezos_specific_sections(doc, context, charts)
    elif config.legacy_layout and forbes_uri == "elon-musk":
        add_elon_specific_sections(doc, context, charts)
    elif config.legacy_layout and forbes_uri == "mark-zuckerberg":
        add_mark_specific_sections(doc, context, charts)
    elif config.legacy_layout and forbes_uri == "bernard-arnault":
        add_arnault_specific_sections(doc, context, charts)
    elif config.legacy_layout and forbes_uri == "larry-ellison":
        add_ellison_specific_sections(doc, context, charts)
    else:
        add_generic_future_sections(doc, context, evidence_keys)
    if charts and (not config.legacy_layout or forbes_uri not in {"jeff-bezos", "mark-zuckerberg", "bernard-arnault", "larry-ellison"}):
        last_chart = charts[-1]
        doc.add_heading("Business empire map", level=1)
        add_picture_with_caption(doc, last_chart, "Business empire map generated from structured data and evidence-pack entities")
        add_confidence_note(doc, "Medium", "The map includes only evidence-visible assets and intentionally excludes unsourced holdings.", evidence_keys)
    add_appendices(doc, evidence)

    doc.core_properties.title = f"{context.person['name']} Business Empire Analysis - {variant}"
    doc.core_properties.subject = f"Forbes Top 100 Billionaires {config.year} enriched business empire report"
    doc.core_properties.author = "Forbes research pipeline"
    doc.save(output_path)
    return output_path


def extract_docx_validation(
    path: Path,
    *,
    charts_dir: Path = CHARTS_DIR,
    expected_rank: int | None = None,
    expected_name: str | None = None,
    report_status: str | None = None,
    evidence_registry_path: Path | None = None,
    leakage_terms: list[str] | None = None,
    allowed_leakage_terms: list[str] | None = None,
) -> dict[str, Any]:
    """Validate a DOCX and return report QA metrics."""
    with zipfile.ZipFile(path) as archive:
        bad_member = archive.testzip()
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    doc = Document(path)
    headings = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.style.name.startswith("Heading") and paragraph.text.strip()]
    heading1 = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.style.name == "Heading 1" and paragraph.text.strip()]
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    all_text = "\n".join(paragraphs + [cell.text for table in doc.tables for row in table.rows for cell in row.cells])
    body_parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Evidence appendix":
            break
        body_parts.append(paragraph.text)
    body_text = "\n".join(body_parts)
    body_keys = sorted(set(re.findall(r"\[([A-Z][0-9A-Z]*)\]", body_text)))
    appendix_keys: set[str] = set()
    evidence_rows = 0
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if headers and headers[0] == "key":
            for row in table.rows[1:]:
                key = row.cells[0].text.strip()
                if key:
                    appendix_keys.add(key)
                    evidence_rows += 1
    referenced_chart_names = sorted(set(re.findall(r"([A-Za-z0-9_-]+\.png)", all_text)))
    missing_charts = [name for name in referenced_chart_names if not (charts_dir / name).exists()]

    section_results: dict[str, bool] = {}
    current_heading = ""
    current_text: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name.startswith("Heading 1") and text:
            if current_heading:
                joined = "\n".join(current_text)
                section_results[current_heading] = bool(re.search(r"\[[A-Z][0-9A-Z]*\]", joined) or "Confidence note:" in joined)
            current_heading = text
            current_text = []
        elif current_heading:
            current_text.append(text)
    if current_heading:
        joined = "\n".join(current_text)
        section_results[current_heading] = bool(re.search(r"\[[A-Z][0-9A-Z]*\]", joined) or "Confidence note:" in joined)

    status = report_status or infer_report_status(path)
    valid_report_status = status in REPORT_STATUS_VALUES
    expected_name_match: bool | None = None
    expected_rank_match: bool | None = None
    if expected_name:
        expected_name_match = expected_name.casefold() in all_text.casefold()
    if expected_rank is not None:
        rank_patterns = [
            rf"\brank\s+{expected_rank}\b",
            rf"\brank\s*[:#-]?\s*{expected_rank}\b",
            rf"\b{expected_rank:03d}_",
        ]
        expected_rank_match = any(re.search(pattern, all_text, flags=re.IGNORECASE) for pattern in rank_patterns)

    allowed_terms = {term.casefold() for term in (allowed_leakage_terms or [])}
    leakage_hits: dict[str, int] = {}
    for term in leakage_terms or []:
        clean_term = term.strip()
        if not clean_term or clean_term.casefold() in allowed_terms:
            continue
        pattern = rf"(?<![A-Za-z]){re.escape(clean_term)}(?![A-Za-z])"
        count = len(re.findall(pattern, body_text, flags=re.IGNORECASE))
        if count:
            leakage_hits[clean_term] = count

    evidence_registry_rows: int | None = None
    evidence_registry_has_rows: bool | None = None
    if evidence_registry_path is not None:
        if evidence_registry_path.exists():
            registry = pd.read_csv(evidence_registry_path)
            if "used_in_report_file" in registry.columns:
                report_name = path.name
                evidence_registry_rows = int(
                    registry["used_in_report_file"].astype(str).str.replace("\\", "/", regex=False).str.endswith(report_name).sum()
                )
            else:
                evidence_registry_rows = 0
        else:
            evidence_registry_rows = 0
        evidence_registry_has_rows = evidence_registry_rows > 0

    word_count = len(re.findall(r"\b\w+\b", all_text))
    optional_checks_pass = (
        (expected_name_match is not False)
        and (expected_rank_match is not False)
        and not leakage_hits
        and valid_report_status
        and (evidence_registry_has_rows is not False)
    )
    return {
        "path": str(path),
        "valid_docx_zip": bad_member is None,
        "bad_zip_member": bad_member,
        "opens_with_python_docx": True,
        "report_status": status,
        "valid_report_status": valid_report_status,
        "expected_name_match": expected_name_match,
        "expected_rank_match": expected_rank_match,
        "section_count": len(heading1),
        "headings": headings,
        "body_citation_key_count": len(body_keys),
        "body_citation_keys": body_keys,
        "appendix_evidence_key_count": len(appendix_keys),
        "appendix_evidence_keys": sorted(appendix_keys),
        "evidence_rows": evidence_rows,
        "missing_body_keys_from_appendix": sorted(set(body_keys) - appendix_keys),
        "referenced_chart_count": len(referenced_chart_names),
        "referenced_charts": referenced_chart_names,
        "embedded_media_count": len(media),
        "missing_referenced_charts": missing_charts,
        "sections_without_citation_or_confidence": [heading for heading, passed in section_results.items() if not passed],
        "contains_raw_urls_in_body": "http://" in body_text or "https://" in body_text,
        "previous_person_leakage_hits": leakage_hits,
        "evidence_registry_rows": evidence_registry_rows,
        "evidence_registry_has_rows": evidence_registry_has_rows,
        "word_count": word_count,
        "approx_pages": max(1, round(word_count / 450 + len(media) * 0.6)),
        "passed": (
            bad_member is None
            and not (set(body_keys) - appendix_keys)
            and not missing_charts
            and not [heading for heading, passed in section_results.items() if not passed]
            and not ("http://" in body_text or "https://" in body_text)
            and optional_checks_pass
        ),
    }


def generate_enriched_report(
    *,
    rank: int | None = None,
    name: str | None = None,
    uri: str | None = None,
    processed_dir: Path | None = None,
    output_dir: Path | None = None,
    charts_dir: Path | None = None,
    evidence_pack_path: Path | None = None,
    variant: str = "enriched_draft",
    overwrite: bool = False,
    year: int = DEFAULT_TARGET_YEAR,
) -> tuple[Path, dict[str, Any]]:
    """Generate and validate a single enriched report."""
    config = get_year_config(year)
    processed_dir = processed_dir or config.processed_dir
    charts_dir = charts_dir or config.charts_dir
    context = load_context(
        rank=rank,
        name=name,
        uri=uri,
        processed_dir=processed_dir,
        evidence_pack_path=evidence_pack_path,
        year=year,
    )
    path = create_enriched_report(
        context,
        output_dir=output_dir,
        charts_dir=charts_dir,
        variant=variant,
        overwrite=overwrite,
        year=year,
    )
    validation = extract_docx_validation(path, charts_dir=charts_dir)
    if not validation["passed"]:
        raise RuntimeError(f"Enriched DOCX validation failed: {validation}")
    return path, validation


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate one enriched Forbes billionaire DOCX report.")
    parser.add_argument("--year", type=int, default=DEFAULT_TARGET_YEAR, help="Canonical annual-list year to use.")
    selector = parser.add_mutually_exclusive_group(required=True)
    selector.add_argument("--rank", type=int, help="Forbes annual-list rank to generate.")
    selector.add_argument("--name", help="Exact or partial person name to generate.")
    selector.add_argument("--uri", help="Forbes profile URI to generate.")
    parser.add_argument("--processed-dir", type=Path, default=None)
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--charts-dir", type=Path, default=None)
    parser.add_argument("--evidence-pack", type=Path, default=None)
    parser.add_argument("--variant", default="enriched_draft", help="Filename variant, for example enriched_v2 or enriched_draft.")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite the target report if it exists and is not locked.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    path, validation = generate_enriched_report(
        rank=args.rank,
        name=args.name,
        uri=args.uri,
        processed_dir=args.processed_dir,
        output_dir=args.output_dir,
        charts_dir=args.charts_dir,
        evidence_pack_path=args.evidence_pack,
        variant=args.variant,
        overwrite=args.overwrite,
        year=args.year,
    )
    print(f"Generated enriched DOCX report: {path}")
    print(f"Sections: {validation['section_count']}")
    print(f"Body citation keys: {validation['body_citation_key_count']}")
    print(f"Evidence rows: {validation['evidence_rows']}")
    print(f"Referenced charts: {validation['referenced_chart_count']}")
    print(f"Approx pages: {validation['approx_pages']}")


if __name__ == "__main__":
    main()
