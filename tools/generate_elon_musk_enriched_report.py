"""Generate the enriched rank-1 Elon Musk business empire DOCX report.

This is intentionally rank-specific for Phase 3A. It does not alter the
baseline DOCX or batch-generate reports for the other 99 people.
"""

from __future__ import annotations

import math
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
PROCESSED_DIR = BASE_DIR / "data" / "processed"
CHARTS_DIR = BASE_DIR / "reports" / "charts"
PEOPLE_DIR = BASE_DIR / "reports" / "people"
BASELINE = PEOPLE_DIR / "001_elon-musk_business_analysis.docx"
OUTPUT = PEOPLE_DIR / "001_elon-musk_business_analysis_enriched.docx"

ACCESS_DATE = "2026-06-11"


@dataclass(frozen=True)
class Source:
    key: str
    title: str
    publisher: str
    date: str
    reliability: str
    supports: str
    locator: str
    limitation: str


SOURCES = [
    Source(
        "F1",
        "Forbes World's Billionaires 2025 annual list cached API record",
        "Forbes",
        "2025-04-01",
        "Canonical ranking source",
        "Rank 1, $342B 2025 net worth, age, country, and source of wealth.",
        "forbes.com/forbesapi/person/billionaires/2025/... and local raw cache",
        "Forbes net worth is an estimate and not an audited personal balance sheet.",
    ),
    Source(
        "D1",
        "Project processed Forbes annual wealth history",
        "Local project data derived from Forbes annual files",
        "2001-2025 snapshots",
        "Canonical project-derived dataset",
        "Musk observed wealth history, CAGR, log-linear fit, drawdown, and volatility.",
        "data/processed/billionaire_wealth_history_long.csv and billionaire_growth_metrics.csv",
        "Annual snapshots can miss intra-year volatility and private valuation marks.",
    ),
    Source(
        "T1",
        "Tesla Form 10-K for fiscal year 2024",
        "U.S. SEC / Tesla, Inc.",
        "2025-01-30",
        "Primary filing",
        "Tesla revenue, gross profit, operating income, net income, operating cash flow, capex, cash, risk factors, and market sensitivity.",
        "sec.gov/Archives/edgar/data/1318605/000162828025003063/tsla-20241231.htm",
        "Tesla reports company results, not Musk's personal balance sheet.",
    ),
    Source(
        "T2",
        "Tesla 2024 definitive proxy statement",
        "U.S. SEC / Tesla, Inc.",
        "2024-04-29",
        "Primary filing",
        "Musk beneficial ownership of 715,022,706 Tesla shares, 20.5% beneficial ownership, pledged-share disclosure, and 2018 CEO award context.",
        "sec.gov/Archives/edgar/data/1318605/000110465924053333/tm2326076d15_def14a.htm",
        "Ownership table is as of March 31, 2024, before the Forbes 2025 list date.",
    ),
    Source(
        "T3",
        "Tesla 2025 definitive proxy statement",
        "U.S. SEC / Tesla, Inc.",
        "2025-09-17",
        "Primary filing",
        "Later ownership context: 717,323,438 beneficial shares and 19.8% ownership as of September 15, 2025, including exercisable 2018 award options.",
        "sec.gov/Archives/edgar/data/1318605/000110465925090866/tm252289-12_def14a.htm",
        "Post-dates Forbes 2025 annual list; used only as later context.",
    ),
    Source(
        "S1",
        "SpaceX valued at $350bn as company agrees to buy shares from employees",
        "The Guardian, citing Bloomberg/internal tender reporting",
        "2024-12-11",
        "Reputable secondary",
        "SpaceX private valuation estimate around $350B, tender offer mechanics, and reported Musk stake context.",
        "theguardian.com/science/2024/dec/11/spacex-valued-at-350bn-as-company-agrees-to-buy-shares-from-employees",
        "Private valuation and ownership figures are secondary estimates, not company-filed financials.",
    ),
    Source(
        "S2",
        "NASA Commercial Crew Program",
        "NASA",
        "accessed 2026-06-11",
        "Primary government source",
        "Commercial crew public-private model and NASA reliance on SpaceX Dragon for ISS transportation.",
        "nasa.gov/humans-in-space/commercial-space/commercial-crew-program/",
        "Program page describes missions, not SpaceX economics or profit margins.",
    ),
    Source(
        "S3",
        "NASA picks SpaceX to land next Americans on Moon",
        "NASA",
        "2021-04-16",
        "Primary government source",
        "SpaceX HLS Starship award, $2.89B firm-fixed-price milestone-based contract, reusable Starship architecture.",
        "nasa.gov/news-release/as-artemis-moves-forward-nasa-picks-spacex-to-land-next-americans-on-moon/",
        "Contract award value does not disclose SpaceX margin or internal development cost.",
    ),
    Source(
        "X1",
        "Elon Musk to Acquire Twitter",
        "Twitter, Inc. press release via PR Newswire",
        "2022-04-25",
        "Company announcement",
        "$44B Twitter acquisition price, $54.20 per share, debt and equity financing commitments.",
        "prnewswire.com/news-releases/elon-musk-to-acquire-twitter-301532245.html",
        "Announcement is pre-close and does not report later X operating performance.",
    ),
    Source(
        "X2",
        "Elon Musk says his company xAI just bought his other company X for $33 billion",
        "Business Insider",
        "2025-03-28",
        "Reputable secondary",
        "Reported xAI-X all-stock transaction values: xAI at $80B and X at $33B equity value.",
        "businessinsider.com/elon-musk-says-xai-acquired-x-in-all-stock-deal-2025-3",
        "Secondary report based on Musk announcement; exact ownership allocation and minority investor economics remain undisclosed.",
    ),
    Source(
        "A1",
        "Series B funding round",
        "xAI",
        "2024-05-26",
        "Company announcement",
        "xAI $6B Series B, investors, product infrastructure, and mission.",
        "x.ai/news/series-b",
        "No valuation table, revenue, profit, or Musk ownership percentage disclosed by xAI.",
    ),
    Source(
        "A2",
        "xAI raises $6B Series C",
        "xAI",
        "2024-12-23",
        "Company announcement",
        "xAI $6B Series C, Colossus 100,000 Nvidia Hopper GPUs, planned 200,000 GPU expansion, Grok/X product link.",
        "x.ai/news/series-c",
        "Company announcement is strategic and does not disclose audited financials.",
    ),
    Source(
        "A3",
        "Grok 3 Beta - The Age of Reasoning Agents",
        "xAI",
        "2025-02-19",
        "Company announcement",
        "Grok 3 model positioning, reasoning emphasis, Colossus training infrastructure, benchmark claims.",
        "x.ai/news/grok-3",
        "Benchmarks are company-reported and should be independently verified before investment-grade conclusions.",
    ),
    Source(
        "N1",
        "Neuralink's first implant partly detached from patient's brain",
        "The Guardian",
        "2024-05-09",
        "Reputable secondary",
        "Neuralink first human implant issue, restored functionality claim, and approximate $5B valuation context.",
        "theguardian.com/technology/article/2024/may/09/neuralink-brain-chip-implant",
        "Valuation and clinical details are incomplete; not a regulatory filing or audited company disclosure.",
    ),
    Source(
        "B1",
        "The Boring Company homepage and Prufrock/Vegas Loop pages",
        "The Boring Company",
        "accessed 2026-06-11",
        "Company source",
        "TBC mission, Loop operating status, Vegas Loop approvals, Prufrock target speed and cost claims, vertical integration.",
        "boringcompany.com, /vegas-loop, /prufrock",
        "Company claims are not independently audited and do not disclose revenue or Musk ownership.",
    ),
]


TESLA_FACTS = {
    "revenue_2024_b": 97.690,
    "gross_profit_2024_b": 17.450,
    "operating_income_2024_b": 7.076,
    "net_income_2024_b": 7.091,
    "operating_cash_flow_2024_b": 14.923,
    "capex_2024_b": 11.339,
    "cash_2024_b": 16.139,
    "assets_2024_b": 122.070,
    "musk_beneficial_shares_2024_m": 715.022706,
    "musk_beneficial_ownership_pct_2024": 20.5,
    "musk_beneficial_ownership_pct_2025_later": 19.8,
    "musk_pledged_shares_2024_m": 238.441261,
}


def money(value: float | int | None, suffix: str = "B") -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "not disclosed"
    return f"${value:,.1f}{suffix}"


def pct(value: float | None) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "not disclosed"
    return f"{value * 100:.1f}%"


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = paragraph._p.add_hyperlink(r_id)
    new_run = hyperlink.add_r()
    r_pr = new_run.get_or_add_rPr()
    new_run.text = text


def add_para(doc: Document, text: str = "", style: str | None = None):
    para = doc.add_paragraph(style=style)
    if text:
        para.add_run(text)
    return para


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths: tuple[float, float] = (2.2, 4.8)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    for row in table.rows:
        row.cells[0].width = Inches(widths[0])
        row.cells[1].width = Inches(widths[1])


def add_matrix_table(doc: Document, columns: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)


def load_project_data() -> tuple[pd.Series, pd.DataFrame, pd.Series]:
    top = pd.read_csv(PROCESSED_DIR / "top100_2025.csv")
    hist = pd.read_csv(PROCESSED_DIR / "billionaire_wealth_history_long.csv")
    metrics = pd.read_csv(PROCESSED_DIR / "billionaire_growth_metrics.csv")
    person = top.loc[top["rank_2025"].eq(1)].iloc[0]
    history = hist.loc[hist["forbes_uri"].eq("elon-musk")].sort_values("year")
    metric = metrics.loc[metrics["forbes_uri"].eq("elon-musk")].iloc[0]
    return person, history, metric


def build_charts(history: pd.DataFrame) -> list[Path]:
    CHARTS_DIR.mkdir(parents=True, exist_ok=True)
    created: list[Path] = []

    wealth_path = CHARTS_DIR / "elon_musk_wealth_history_enriched.png"
    fig, ax = plt.subplots(figsize=(8.4, 4.6), dpi=180)
    ax.plot(history["year"], history["net_worth_usd_b"], marker="o", linewidth=2.4, color="#1f77b4")
    ax.fill_between(history["year"], history["net_worth_usd_b"], color="#1f77b4", alpha=0.12)
    ax.set_title("Elon Musk Forbes Annual Net Worth History")
    ax.set_xlabel("Forbes annual list year")
    ax.set_ylabel("Net worth, USD billions")
    ax.grid(True, alpha=0.25)
    ax.annotate(
        "$342B in 2025",
        xy=(2025, float(history.loc[history["year"].eq(2025), "net_worth_usd_b"].iloc[0])),
        xytext=(2021.6, 320),
        arrowprops={"arrowstyle": "->", "color": "#333333"},
        fontsize=9,
    )
    fig.tight_layout()
    fig.savefig(wealth_path, bbox_inches="tight")
    plt.close(fig)
    created.append(wealth_path)

    sens_path = CHARTS_DIR / "elon_musk_tesla_sensitivity_enriched.png"
    delta = list(range(-500, 501, 100))
    y_2024 = [x * TESLA_FACTS["musk_beneficial_ownership_pct_2024"] / 100 for x in delta]
    y_2025 = [x * TESLA_FACTS["musk_beneficial_ownership_pct_2025_later"] / 100 for x in delta]
    fig, ax = plt.subplots(figsize=(8.4, 4.6), dpi=180)
    ax.axhline(0, color="#222222", linewidth=0.8)
    ax.axvline(0, color="#222222", linewidth=0.8)
    ax.plot(delta, y_2024, marker="o", label="20.5% 2024 proxy exposure", color="#d62728")
    ax.plot(delta, y_2025, marker="s", label="19.8% later 2025 proxy exposure", color="#9467bd")
    ax.set_title("Illustrative Tesla Market-Cap Sensitivity to Musk Wealth")
    ax.set_xlabel("Change in Tesla market capitalization, USD billions")
    ax.set_ylabel("Approx. change in Musk Tesla-linked wealth, USD billions")
    ax.grid(True, alpha=0.25)
    ax.legend(frameon=False, fontsize=8)
    ax.text(
        -490,
        max(y_2024) * 0.78,
        "Illustrative only: excludes taxes, pledge terms,\noption exercise prices, discounts, debt, and liquidity.",
        fontsize=8,
        bbox={"boxstyle": "round,pad=0.35", "facecolor": "#f7f7f7", "edgecolor": "#bbbbbb"},
    )
    fig.tight_layout()
    fig.savefig(sens_path, bbox_inches="tight")
    plt.close(fig)
    created.append(sens_path)

    map_path = CHARTS_DIR / "elon_musk_business_empire_map_enriched.png"
    fig, ax = plt.subplots(figsize=(10, 6), dpi=180)
    ax.axis("off")
    nodes = {
        "Elon Musk\nControl / ownership\nattention / capital access": (0.5, 0.78, "#111827"),
        "Tesla\nPublic equity engine": (0.18, 0.52, "#e11d48"),
        "SpaceX + Starlink\nPrivate valuation engine": (0.50, 0.52, "#2563eb"),
        "xAI + X data\nAI/platform optionality": (0.82, 0.52, "#7c3aed"),
        "Neuralink\nRegulated neurotech option": (0.32, 0.24, "#059669"),
        "The Boring Co.\nInfrastructure option": (0.68, 0.24, "#d97706"),
    }
    for label, (x, y, color) in nodes.items():
        ax.text(
            x,
            y,
            label,
            ha="center",
            va="center",
            color="white",
            fontsize=10,
            fontweight="bold",
            bbox={"boxstyle": "round,pad=0.55", "facecolor": color, "edgecolor": color, "alpha": 0.95},
        )
    center = nodes["Elon Musk\nControl / ownership\nattention / capital access"]
    for label, (x, y, _color) in nodes.items():
        if label.startswith("Elon"):
            continue
        ax.annotate(
            "",
            xy=(x, y + 0.065),
            xytext=(center[0], center[1] - 0.065),
            arrowprops={"arrowstyle": "->", "color": "#4b5563", "lw": 1.4},
        )
    ax.text(
        0.5,
        0.06,
        "Wealth architecture: public equity + private aerospace + AI/social data + long-duration options",
        ha="center",
        fontsize=10,
        color="#111827",
    )
    fig.tight_layout()
    fig.savefig(map_path, bbox_inches="tight")
    plt.close(fig)
    created.append(map_path)

    return created


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.bold = True
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 3"].font.size = Pt(10.5)


def add_title(doc: Document, person: pd.Series) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Elon Musk: Business Empire and Wealth Engine Analysis")
    run.bold = True
    run.font.size = Pt(20)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        f"Rank {int(person.rank_2025)} on Forbes World's Billionaires 2025 annual list | "
        f"{money(float(person.net_worth_2025_usd_b))} estimated net worth | Enriched Phase 3A sample"
    )
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "Body citations use short keys such as [F1] and [T1]. Full source details and limitations are in the evidence appendix."
    ).italic = True


def add_intro_tables(doc: Document, person: pd.Series, metric: pd.Series) -> None:
    doc.add_heading("1. Executive thesis", level=1)
    add_para(
        doc,
        "Musk's 2025 Forbes wealth is not a single-company biography. It is a portfolio of convex control positions: "
        "Tesla supplies the liquid public-equity beta, SpaceX supplies the private aerospace/Starlink valuation mark, "
        "and X, xAI, Neuralink, and The Boring Company add option value whose economics are much less disclosed. "
        "The core wealth equation is: net worth roughly equals ownership stakes multiplied by the value of major assets, "
        "less debt, taxes, discounts for illiquidity, pledge constraints, and uncertainty in private-company marks [F1][T2][S1].",
    )
    add_para(
        doc,
        "The important business lesson is that Musk compounded wealth by retaining concentrated ownership in companies attacking "
        "large physical or digital bottlenecks: electric-vehicle scale, reusable launch, satellite broadband, AI compute and data, "
        "brain-computer interfaces, and tunneling cost. That concentration creates huge upside, but it also concentrates risk: "
        "public-market sentiment, regulatory approvals, product execution, governance, and private valuation marks all flow into the same personal balance sheet.",
    )
    add_kv_table(
        doc,
        [
            ("Forbes 2025 rank", str(int(person.rank_2025))),
            ("Forbes 2025 net worth", money(float(person.net_worth_2025_usd_b))),
            ("Country", str(person.country_or_territory)),
            ("Forbes source of wealth", str(person.source_of_wealth)),
            ("Primary 2025 public wealth engine", "Tesla public equity and related option exposure [T2]"),
            ("Primary 2025 private wealth engine", "SpaceX private valuation, including Starlink and launch economics [S1][S2][S3]"),
            ("First Forbes annual observation", f"{int(metric.first_year_observed)} at {money(float(metric.first_net_worth_usd_b))}"),
            ("Observed wealth multiple", f"{float(metric.wealth_multiple_first_to_2025):.1f}x from first observation to 2025 [D1]"),
        ],
    )


def add_wealth_equation(doc: Document, history: pd.DataFrame, metric: pd.Series, charts: list[Path]) -> None:
    doc.add_heading("2. Wealth equation and asset map", level=1)
    add_para(
        doc,
        "The 2025 equation can be read as a layered asset stack rather than a smooth compounding account. "
        "Tesla is marked continuously by public markets. SpaceX is marked episodically by private tender rounds. "
        "X/xAI, Neuralink, and The Boring Company are option-like because exact ownership, operating profit, debt load, and private valuation marks are incomplete. "
        "That makes the right analytical form an attribution bridge, not a precise sum-of-the-parts audit.",
    )
    add_matrix_table(
        doc,
        ["Asset", "Wealth role", "Evidence", "Main uncertainty"],
        [
            [
                "Tesla",
                "Liquid public equity engine; Forbes-level wealth moves with stock price and option value.",
                "Musk beneficially owned 715.0M Tesla shares, 20.5% beneficial ownership as of Mar. 31, 2024 [T2].",
                "Tax, pledge, exercise-price, and liquidity effects mean stock exposure is not identical to spendable wealth.",
            ],
            [
                "SpaceX / Starlink",
                "Private aerospace and satellite-broadband valuation engine.",
                "Tender reports valued SpaceX around $350B in Dec. 2024; NASA confirms commercial crew and HLS relationships [S1][S2][S3].",
                "Private valuation, revenue, profit, and Musk ownership percentage are not public audited facts.",
            ],
            [
                "X / Twitter",
                "Strategic data, distribution, and payments/social platform optionality.",
                "Twitter agreed to be acquired for about $44B in 2022; xAI-X transaction values were later reported at $80B/$33B [X1][X2].",
                "Debt, revenue quality, ad demand, and exact ownership economics remain opaque.",
            ],
            [
                "xAI",
                "AI compute, model, and X data optionality.",
                "xAI announced $6B Series B and $6B Series C, with Colossus at 100,000 Nvidia Hopper GPUs [A1][A2].",
                "No audited revenue, profit, cap table, or Musk stake disclosed.",
            ],
            [
                "Neuralink",
                "Long-horizon regulated neurotechnology option.",
                "Human implant progress and technical setback are reported; approximate valuation is secondary [N1].",
                "Clinical, regulatory, reimbursement, and product-market outcomes remain early and uncertain.",
            ],
            [
                "The Boring Company",
                "Infrastructure-cost optionality and city/regulatory execution test.",
                "Vegas Loop operating/approval claims and Prufrock cost/speed targets are company-disclosed [B1].",
                "No public financials, cap table, or independent cost validation.",
            ],
        ],
    )
    doc.add_paragraph()
    doc.add_picture(str(charts[2]), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 1. Business empire map. Source: author's synthesis from evidence keys [F1][T1][T2][S1][A1][A2][B1].").italic = True

    doc.add_heading("3. Wealth history and exponential-style fit", level=1)
    add_para(
        doc,
        f"Forbes annual snapshots show Musk moving from {money(float(metric.first_net_worth_usd_b))} in {int(metric.first_year_observed)} "
        f"to {money(342.0)} in 2025 across {int(metric.years_observed)} observations. The nominal CAGR is {pct(float(metric.CAGR_nominal))}; "
        f"the log-linear slope is {float(metric.log_linear_growth_slope):.3f}, R^2 is {float(metric.exponential_fit_r2):.3f}, and the fitted doubling time is "
        f"{float(metric.estimated_doubling_time_years):.1f} years [D1]. This supports a strong exponential-style pattern in annual Forbes marks, "
        "but it should not be called true exponential growth: the underlying series is driven by discrete public-market reratings, private tender marks, option awards, and drawdowns.",
    )
    add_para(
        doc,
        f"The largest one-year gain in the project data is {money(float(metric.max_one_year_gain_usd_b))}; the largest one-year loss is "
        f"{money(abs(float(metric.max_one_year_loss_usd_b)))}; and the largest drawdown is {pct(abs(float(metric.largest_drawdown_pct)))} [D1]. "
        "That volatility is a feature of concentrated equity compounding, not a defect in the data.",
    )
    doc.add_picture(str(charts[0]), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figure 2. Forbes annual net worth history. Source: project wealth history derived from Forbes annual files [D1].").italic = True


def add_tesla_section(doc: Document, charts: list[Path]) -> None:
    doc.add_heading("4. Tesla public equity wealth engine", level=1)
    add_para(
        doc,
        "Tesla is Musk's most transparent wealth engine because it is public, liquid, and filed. The asset became valuable because Tesla converted "
        "EVs from a niche engineering project into a scaled automotive, battery, charging, software, and brand platform. Public markets then valued Tesla "
        "less like a conventional automaker and more like a high-growth technology/energy platform with autonomy and robotics optionality.",
    )
    add_para(
        doc,
        f"Financial-statement linkage is direct. In 2024 Tesla reported revenue of {money(TESLA_FACTS['revenue_2024_b'])}, gross profit of "
        f"{money(TESLA_FACTS['gross_profit_2024_b'])}, operating income of {money(TESLA_FACTS['operating_income_2024_b'])}, net income of "
        f"{money(TESLA_FACTS['net_income_2024_b'])}, operating cash flow of {money(TESLA_FACTS['operating_cash_flow_2024_b'])}, and "
        f"{money(TESLA_FACTS['capex_2024_b'])} of payments to acquire property and equipment [T1]. Those numbers show why Tesla can support a large public valuation, "
        "but they also show capital intensity: the equity story depends on factories, batteries, vehicles, energy storage, AI compute, and service infrastructure, not pure software margins.",
    )
    gross_margin = TESLA_FACTS["gross_profit_2024_b"] / TESLA_FACTS["revenue_2024_b"]
    op_margin = TESLA_FACTS["operating_income_2024_b"] / TESLA_FACTS["revenue_2024_b"]
    ocf_margin = TESLA_FACTS["operating_cash_flow_2024_b"] / TESLA_FACTS["revenue_2024_b"]
    add_kv_table(
        doc,
        [
            ("Revenue model", "Vehicle sales, leasing, energy generation/storage, services, software and regulatory credits [T1]."),
            ("2024 gross margin", pct(gross_margin)),
            ("2024 operating margin", pct(op_margin)),
            ("2024 operating cash flow margin", pct(ocf_margin)),
            ("Capital intensity", f"{money(TESLA_FACTS['capex_2024_b'])} PP&E capex in 2024 [T1]."),
            ("Musk ownership bridge", "715.0M beneficial shares / 20.5% beneficial ownership as of Mar. 31, 2024 [T2]."),
            ("Pledge / liquidity caveat", "238.4M shares disclosed as pledged collateral in the 2024 proxy [T2]."),
        ],
    )
    add_para(
        doc,
        "Moat analysis: Tesla's advantages are not one thing. The defensible system includes manufacturing learning curves, battery supply and integration, "
        "charging infrastructure, brand, over-the-air software, fleet data, and the capital-market trust to finance large physical expansion. The same system is fragile if EV demand slows, "
        "price cuts compress margins, autonomy claims disappoint, regulators constrain FSD/robotaxi ambitions, or customer sentiment turns against the brand [T1].",
    )
    doc.add_picture(str(charts[1]), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "Figure 3. Tesla market-cap sensitivity. This is an ownership-exposure explanation using proxy percentages, not a forecast or full net-worth model [T2][T3].",
    ).italic = True


def add_spacex_section(doc: Document) -> None:
    doc.add_heading("5. SpaceX and Starlink private-company valuation engine", level=1)
    add_para(
        doc,
        "SpaceX is the private-company engine that makes Musk's Forbes wealth less transparent than a pure public-equity fortune. Its value rests on two linked businesses: "
        "launch services with reusable rockets, and Starlink satellite broadband. The company attacks the cost and reliability constraint in access to orbit, then uses that lower-cost launch base "
        "to deploy a satellite communications network. That vertical loop is why the private valuation can rise even when audited public financials are unavailable.",
    )
    add_para(
        doc,
        "NASA evidence supports the strategic importance, not the full valuation. NASA's Commercial Crew Program describes a public-private model for safe, reliable and cost-effective human transport "
        "to the ISS, and SpaceX Crew Dragon is a central operating system in that architecture [S2]. NASA also awarded SpaceX a $2.89B firm-fixed-price, milestone-based Human Landing System contract "
        "for Artemis, linking Starship development to a government-backed lunar mission architecture [S3].",
    )
    add_para(
        doc,
        "The private valuation evidence is secondary: tender reports in late 2024 put SpaceX around $350B [S1]. That mark matters for Musk's wealth because a large private stake multiplied by a large "
        "private valuation can rival or exceed the public-equity contribution from Tesla. But the report should not treat the $350B mark as audited intrinsic value. It is an observed private-market price signal, "
        "likely influenced by Starlink growth expectations, launch cadence, government demand, and investor scarcity for private SpaceX exposure.",
    )
    add_matrix_table(
        doc,
        ["Mechanism", "Why it matters", "Risk / uncertainty"],
        [
            ["Reusable launch", "Lower launch cost and higher cadence can widen the addressable market for orbital services.", "Execution risk, launch failures, regulatory limits, and maintenance economics."],
            ["Starlink", "Turns launch capacity into recurring broadband/service potential and strategic communications infrastructure.", "Spectrum, orbital debris, customer economics, competition, geopolitics."],
            ["Government contracts", "NASA and defense demand can fund frontier technology and validate reliability.", "Political/regulatory dependence and contract concentration."],
            ["Private valuation", "Large tender marks directly affect Forbes-style wealth estimates.", "Private marks are episodic, opaque, and illiquid."],
        ],
    )


def add_optionality_sections(doc: Document) -> None:
    doc.add_heading("6. X / Twitter ownership and strategic optionality", level=1)
    add_para(
        doc,
        "X is not a clean cash-flow wealth engine in the evidence set. Twitter agreed to sell to an entity wholly owned by Musk for about $44B, with $25.5B of debt and margin loan financing "
        "and about $21B of equity commitment [X1]. After the acquisition, the strategic thesis shifted from public social-media valuation to private platform control: identity, distribution, payments, ads, creator economics, "
        "and data for AI products. The risk is equally clear: advertising trust, moderation controversy, debt burden, employee churn, regulatory scrutiny, and opaque financials.",
    )
    add_para(
        doc,
        "In 2025, xAI's reported all-stock acquisition of X valued xAI at $80B and X at $33B equity value [X2]. This is important optionality but not a final wealth number. Exact Musk ownership after investor rounds, "
        "X debt, minority interests, and valuation discounts are not public enough to convert into a precise Forbes bridge.",
    )

    doc.add_heading("7. xAI valuation and AI infrastructure optionality", level=1)
    add_para(
        doc,
        "xAI is a compute-and-data option, not yet an audited cash-flow engine in the evidence set. xAI announced a $6B Series B in May 2024 to take products to market, build infrastructure, and accelerate R&D [A1]. "
        "It then announced a $6B Series C in December 2024, naming major financial and strategic investors and describing Colossus as a 100,000 Nvidia Hopper GPU supercomputer with plans to double to 200,000 GPUs [A2].",
    )
    add_para(
        doc,
        "The business model can develop through consumer subscriptions, enterprise/API access, government/business use cases, and AI features tied to X distribution. The constraint attacked is scarce frontier AI compute plus real-time data. "
        "The moat is unproven but potentially includes capital access, fast infrastructure deployment, X distribution, and Musk's ability to recruit and finance frontier technical teams. The risk is severe: model performance, safety, inference cost, "
        "GPU supply, energy, legal disputes, and the challenge of converting benchmark claims into durable revenue [A2][A3].",
    )

    doc.add_heading("8. Neuralink long-horizon biotech/neurotechnology optionality", level=1)
    add_para(
        doc,
        "Neuralink should be treated as a long-duration clinical and regulatory option, not a near-term Forbes wealth engine. The company is attacking a scarce interface: high-bandwidth communication between human neural signals and digital systems. "
        "Evidence shows first-human progress but also technical setbacks, including a reported partial detachment/retraction issue with the first implant and later software adjustments restoring some functionality [N1].",
    )
    add_para(
        doc,
        "The upside case is extraordinary if brain-computer interfaces become approved medical devices, then broader assistive and augmentation platforms. But the business is capital intensive, medically regulated, ethically sensitive, and still early. "
        "Exact ownership, revenue, profit, reimbursement pathway, and clinical durability are not public enough for an audited valuation bridge.",
    )

    doc.add_heading("9. The Boring Company infrastructure optionality", level=1)
    add_para(
        doc,
        "The Boring Company is the least financially disclosed of the major assets, but strategically it follows the same constraint-removal pattern: reduce the cost and time of tunneling so cities can add transportation capacity without consuming scarce surface land. "
        "The company describes Loop as an all-electric underground public transportation system and says LVCC Loop is commercially operating [B1]. Vegas Loop materials state more than 3 million passengers transported through 8 stations, with 68 miles and 104 stations approved by Clark County and the City of Las Vegas [B1].",
    )
    add_para(
        doc,
        "Prufrock is the core operating bet: faster, standardized, vertically integrated tunnel boring, with the company targeting more than 1 mile per week and less than $8M/mile for Loop transportation tunnels [B1]. "
        "The wealth contribution is option-like: if those cost targets are independently validated at scale, infrastructure economics could rerate; if permitting, safety, utilization, or cost claims disappoint, the value contribution may remain small relative to Tesla and SpaceX.",
    )


def add_first_principles(doc: Document) -> None:
    doc.add_heading("10. First-principles business analysis", level=1)
    add_para(
        doc,
        "The scarce resource Musk controls is not simply capital. It is a bundle: founder control, technical recruiting power, public-market attention, private-market access, risk tolerance, and the ability to organize teams around problems that look uneconomic before scale. "
        "That bundle lets him attack constraints that incumbents often avoid because timelines are long and failure modes are public.",
    )
    add_matrix_table(
        doc,
        ["Company", "Constraint attacked", "Cash-flow vs option-value character", "Capital-market reason for reward"],
        [
            ["Tesla", "EV cost, battery scale, charging, software-defined vehicle adoption.", "Operating engine plus autonomy/robotics option.", "Public market rewarded growth, brand, margin potential, and founder-led optionality."],
            ["SpaceX", "Cost/reliability of launch and global broadband deployment.", "Private operating engine plus Starship/Mars/Starlink option.", "Private markets reward scarcity, technical lead, government validation, and Starlink TAM."],
            ["X", "Distribution/data layer for speech, social graph, identity, payments and AI inputs.", "Turnaround and data-platform option.", "Optionality if X becomes a payment/data/AI surface; punished if ads/debt dominate."],
            ["xAI", "Scarce AI compute, talent, data and product distribution.", "Mostly option-value at current evidence depth.", "Capital markets value frontier AI scarcity and infrastructure speed."],
            ["Neuralink", "High-bandwidth neural interface and paralysis assistive control.", "Long-duration regulated clinical option.", "Huge TAM narrative, but evidence remains early."],
            ["The Boring Co.", "Urban tunneling cost and surface-land scarcity.", "Infrastructure option.", "Potential rerating if standardized tunneling cost falls dramatically."],
        ],
    )
    add_para(
        doc,
        "The empire is therefore a barbell: Tesla and SpaceX are real operating engines with visible products and customers; xAI, X, Neuralink, and The Boring Company add increasingly uncertain option value. "
        "Capital markets rewarded the barbell because successful proof points in one hard domain made investors willing to underwrite adjacent hard-domain bets.",
    )


def add_lessons_and_risks(doc: Document) -> None:
    doc.add_heading("11. Transferable business lessons", level=1)
    add_numbered(
        doc,
        [
            "Ownership concentration is the compounding engine. Musk's Forbes-scale wealth comes from retaining large stakes through extreme volatility, not from salary.",
            "High-convexity bets can dominate a portfolio. Tesla, SpaceX, xAI, Neuralink, and Boring all have capped downside at invested capital but very large narrative upside if technical constraints fall.",
            "Public equity can become strategic leverage. Tesla's public valuation lowers capital cost, recruits talent, and makes Musk's personal wealth highly sensitive to market cap.",
            "Vertical integration can turn impossible unit economics into merely hard unit economics. Tesla integrates batteries/software/manufacturing; SpaceX integrates rocket manufacturing/launch/Starlink; Boring integrates TBMs/liners/operations.",
            "Technology bottlenecks create moats before profits show up. Battery scale, reusable rockets, AI compute, and tunneling speed are bottlenecks that can command valuation before mature margins arrive.",
            "Narrative is not decoration. In Musk's companies, mission narrative helps recruit, sell, finance, and sustain investor patience, but it also creates reputational and regulatory risk.",
            "Platform control compounds optionality. X can be an audience, data, payments, and AI distribution layer; Tesla can be a fleet/data/robotics platform; Starlink can be a connectivity platform.",
            "Execution and regulatory risk scale with ambition. The same concentration that creates upside also creates single-person, governance, safety, labor, political, and regulatory exposure.",
        ],
    )
    doc.add_heading("12. Risks and counter-thesis", level=1)
    add_bullets(
        doc,
        [
            "Tesla multiple compression: if Tesla is valued like an automaker rather than a technology/autonomy platform, Musk's liquid wealth bridge can fall sharply [T1][T2].",
            "Governance and key-person concentration: Tesla proxy materials themselves show the centrality of Musk's voting interest and leadership commitment [T3].",
            "Pledged shares and liquidity: pledged Tesla shares create a personal-balance-sheet caveat; pledge disclosure is not the same as disclosed borrowings, but it matters for downside mechanics [T2].",
            "Private valuation opacity: SpaceX, xAI, Neuralink, X, and Boring do not provide public audited segment financials or cap tables [S1][X2][A1][N1][B1].",
            "Regulatory chokepoints: autonomous driving, launch licenses, spectrum/orbital debris, medical implants, tunneling permits, payments, and AI policy are all external approval systems.",
            "Brand and political spillover: reputational shocks can affect Tesla demand, X ad demand, and regulatory posture across unrelated assets.",
        ],
    )

    doc.add_heading("13. Comparable billionaire patterns", level=1)
    add_para(
        doc,
        "Musk is closest to a hybrid of founder/operator public equity and founder/operator private company. Compared with Jeff Bezos or Mark Zuckerberg, Musk has more exposure to capital-intensive physical systems. "
        "Compared with Larry Ellison, Musk has less recurring enterprise-software cash flow and more technical/regulatory option value. Compared with Bernard Arnault, Musk has less stable luxury-brand pricing power and more dependence on public/private capital-market belief. "
        "The common pattern is concentrated founder ownership; the distinctive pattern is the stacking of several difficult, capital-intensive technology platforms under one reputation and talent magnet.",
    )


def add_appendices(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("14. Evidence appendix", level=1)
    add_para(
        doc,
        "Citation keys are intentionally short in the body. This appendix gives the source details, source quality, claim support, locator, and limitations.",
    )
    add_matrix_table(
        doc,
        ["Key", "Source", "Publisher / date", "Supports", "Quality", "Readable locator", "Limitation"],
        [[s.key, s.title, f"{s.publisher}; {s.date}", s.supports, s.reliability, s.locator, s.limitation] for s in SOURCES],
    )

    doc.add_heading("15. Evidence gaps", level=1)
    add_bullets(
        doc,
        [
            "Exact Musk ownership percentages for SpaceX, xAI, Neuralink, The Boring Company, and X after financing rounds and internal transactions.",
            "Audited revenue, gross margin, operating income, free cash flow, debt, and capex for SpaceX, Starlink, xAI, X, Neuralink, and The Boring Company.",
            "Detailed personal debt, margin loans, pledge terms, taxes, liquidity discounts, and family-office entity holdings.",
            "Independent verification of The Boring Company's tunneling cost/speed targets and utilization economics.",
            "Clinical trial durability, regulatory pathway, reimbursement economics, and adverse-event record for Neuralink.",
            "Private-company valuation dates and share classes used by Forbes to compute the 2025 annual-list estimate.",
        ],
    )

    doc.add_heading("16. Confidence level by section", level=1)
    add_matrix_table(
        doc,
        ["Section", "Confidence", "Reason"],
        [
            ["Forbes 2025 rank/net worth", "High", "Directly from canonical Forbes-derived local processed dataset [F1]."],
            ["Wealth history metrics", "High", "Computed from annual Forbes history with 14 observations [D1]."],
            ["Tesla financial linkage", "High", "SEC primary 10-K and proxy evidence [T1][T2]."],
            ["Tesla personal wealth sensitivity", "Medium-high", "Ownership percentage is sourced; actual personal wealth effect needs taxes, debt, options, discounts."],
            ["SpaceX valuation bridge", "Medium", "Valuation and stake are secondary/private estimates; NASA operating evidence is strong."],
            ["X/xAI optionality", "Medium-low", "Funding announcements are strong; valuation/ownership economics are opaque."],
            ["Neuralink and Boring optionality", "Low-medium", "Operating claims exist but financial and ownership evidence is incomplete."],
            ["Transferable lessons", "Medium", "Synthesis is evidence-grounded but interpretive."],
        ],
    )

    doc.add_heading("17. Claims not final without more evidence", level=1)
    add_bullets(
        doc,
        [
            "Any precise sum-of-the-parts estimate for Musk's personal net worth.",
            "Any claim that xAI, Neuralink, X, or The Boring Company currently contributes a specific dollar amount to Forbes 2025 net worth.",
            "Any claim that SpaceX's tender valuation equals intrinsic value or immediately realizable liquidity.",
            "Any claim that Tesla's autonomy or robotics option has already matured into recurring high-margin revenue.",
            "Any claim that Neuralink is clinically proven beyond early human feasibility evidence.",
        ],
    )

    doc.add_heading("18. Data limitations", level=1)
    add_para(
        doc,
        "This enriched sample is materially stronger than the baseline template because it replaces generic placeholders with sourced operating, financial, and strategic analysis. "
        "It is still not an audited wealth statement. Forbes values are estimates; private companies are opaque; Musk's personal leverage, tax basis, trusts, and family-office structures are not fully disclosed; and several ventures are valued primarily as options. "
        "The report should be reviewed as a strategy and wealth-engine memo, not as investment advice or a legal valuation opinion.",
    )


def create_report() -> dict[str, object]:
    if not BASELINE.exists():
        raise FileNotFoundError(f"Baseline file missing: {BASELINE}")
    if OUTPUT.exists():
        try:
            with OUTPUT.open("a+b"):
                pass
        except OSError as exc:
            raise RuntimeError(f"Target DOCX appears open or locked: {OUTPUT}") from exc
        OUTPUT.unlink()

    person, history, metric = load_project_data()
    charts = build_charts(history)

    doc = Document()
    style_document(doc)
    add_title(doc, person)
    add_intro_tables(doc, person, metric)
    add_wealth_equation(doc, history, metric, charts)
    add_tesla_section(doc, charts)
    add_spacex_section(doc)
    add_optionality_sections(doc)
    add_first_principles(doc)
    add_lessons_and_risks(doc)
    add_appendices(doc)
    doc.core_properties.title = "Elon Musk Business Empire Analysis - Enriched"
    doc.core_properties.subject = "Forbes Top 100 Billionaires 2025 enriched sample report"
    doc.core_properties.author = "Forbes research pipeline"
    PEOPLE_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    with zipfile.ZipFile(OUTPUT) as zf:
        bad = zf.testzip()
        if bad:
            raise RuntimeError(f"Invalid DOCX zip member: {bad}")
        if "word/document.xml" not in zf.namelist():
            raise RuntimeError("DOCX is missing word/document.xml")

    check_doc = Document(OUTPUT)
    headings = [p.text.strip() for p in check_doc.paragraphs if p.style.name.startswith("Heading") and p.text.strip()]
    all_text = "\n".join([p.text for p in check_doc.paragraphs] + [c.text for t in check_doc.tables for r in t.rows for c in r.cells])
    citation_keys = sorted(set(re.findall(r"\[([A-Z][0-9])\]", all_text)))
    words = len(re.findall(r"\b\w+\b", all_text))
    approx_pages = max(1, round(words / 450 + len(charts) * 0.6))
    return {
        "output": str(OUTPUT),
        "charts": [str(path) for path in charts],
        "sections": len(headings),
        "headings": headings,
        "tables": len(check_doc.tables),
        "citation_keys": citation_keys,
        "evidence_rows": len(SOURCES),
        "word_count": words,
        "approx_pages": approx_pages,
        "size_bytes": OUTPUT.stat().st_size,
    }


if __name__ == "__main__":
    result = create_report()
    for key, value in result.items():
        print(f"{key}: {value}")
